
# Shared helper utilities extracted from original file (verbatim bodies)
import asyncio
from linecache import cache
from flask import jsonify, Response
from datetime import datetime, timezone, timedelta
import os, json, time, re, requests, uuid
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy import and_
from app.extensions import logger
import threading
import traceback  
from app.models.db import Candidate, SessionLocal
from flask_cors import cross_origin
from flask import Blueprint, jsonify, request, Response
try:
    from app.extensions import executor
except Exception:
    executor = None

helpers_bp = Blueprint('helpers', __name__)

def extract_skills_from_resume(resume_content):
    """Extract technical skills from resume"""
    skills = []
    resume_lower = resume_content.lower()
    tech_skills = [
        'python', 'javascript', 'java', 'c++', 'c#', 'react', 'angular', 'vue',
        'node.js', 'django', 'flask', 'spring', 'sql', 'nosql', 'mongodb',
        'postgresql', 'mysql', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
        'git', 'ci/cd', 'machine learning', 'data science', 'api', 'rest',
        'graphql', 'typescript', 'golang', 'rust', 'swift', 'kotlin'
    ]
    for skill in tech_skills:
        if skill in resume_lower:
            skills.append(skill.title())
    return skills[:10]



def extract_projects_from_resume(resume_content):
    """Extract project names from resume"""
    projects = []
    lines = resume_content.split('\n')
    for i, line in enumerate(lines):
        if 'project' in line.lower():
            if ':' in line:
                project_name = line.split(':', 1)[1].strip()[:50]
                if project_name:
                    projects.append(project_name)
    return projects[:5]



def extract_experience_years(resume_content):
    """Extract years of experience from resume"""
    import re
    pattern = r'(\d+)\+?\s*years?\s*(?:of\s*)?experience'
    match = re.search(pattern, resume_content.lower())
    if match:
        return f"{match.group(1)}+ years"
    year_pattern = r'20\d{2}'
    years = re.findall(year_pattern, resume_content)
    if len(years) >= 2:
        min_year = min(int(y) for y in years)
        max_year = max(int(y) for y in years)
        experience = max_year - min_year
        if experience > 0:
            return f"{experience}+ years"
    return "Not specified"

def create_structured_interview_kb(candidate_name, position, company, resume_content, job_description):
    """Create a highly structured knowledge base for professional interviews"""
    
    # Extract key information
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    experience = extract_experience_years(resume_content) if resume_content else "Not specified"
    projects = extract_projects_from_resume(resume_content) if resume_content else []
    
    # Build the structured interview content
    structured_content = f"""
STRICT INTERVIEW PROTOCOL - FOLLOW EXACTLY:

YOU ARE: A professional AI interviewer conducting a structured technical interview.
YOUR BEHAVIOR: Professional, clear, structured. NO casual conversation.

CANDIDATE DETAILS:
- Name: {candidate_name}
- Position: {position}
- Company: {company}
- Experience: {experience}
- Key Skills: {', '.join(skills[:5]) if skills else 'General skills'}

INTERVIEW STRUCTURE - ASK THESE QUESTIONS IN EXACT ORDER:

=== QUESTION 1: SELF INTRODUCTION (ALWAYS ASK FIRST) ===
"Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background and what led you to apply for this role?"

=== QUESTION 2: TECHNICAL EXPERIENCE ===
{"I see from your resume that you have experience with " + skills[0] + ". Can you tell me about a specific project where you used " + skills[0] + " and what challenges you faced?" if skills else "Can you tell me about your most significant technical project and the technologies you used?"}

=== QUESTION 3: PROBLEM SOLVING ===
"That's interesting. Now, let me ask you about problem-solving. Can you describe a time when you encountered a complex technical problem and walk me through how you approached and solved it?"

=== QUESTION 4: TEAMWORK ===
"Great. Let's talk about teamwork. Tell me about a time when you had to collaborate with other team members on a challenging project. How did you handle any conflicts or disagreements?"

=== QUESTION 5: SPECIFIC SKILL DEEP DIVE ===
{"I noticed you also have experience with " + skills[1] + ". What's the most complex thing you've built using " + skills[1] + "? Please be specific about the technical details." if len(skills) > 1 else "What would you say is your strongest technical skill, and can you give me a detailed example of how you've applied it?"}

=== QUESTION 6: LEARNING ABILITY ===
"Technology evolves rapidly. Can you tell me about a time when you had to quickly learn a new technology or framework for a project? How did you approach the learning process?"

=== QUESTION 7: ROLE-SPECIFIC ===
"Let's talk specifically about this {position} role. Based on your understanding of the position, how do you see your skills and experience contributing to our team in the first 90 days?"

=== QUESTION 8: CHALLENGES ===
"What do you think would be the biggest challenge for you in this role, and how would you address it?"

=== QUESTION 9: CAREER GOALS ===
"Where do you see your career heading in the next 3-5 years, and how does this {position} role fit into those plans?"

=== QUESTION 10: CANDIDATE QUESTIONS ===
"Thank you for your answers. Now, do you have any questions for me about the role, the team, or {company}?"

CRITICAL RULES:
1. When you receive "INIT_INTERVIEW", IMMEDIATELY ask Question 1
2. Ask ONE question at a time
3. Wait for complete answer before next question
4. After each answer say: "Thank you for sharing that. [Next question]"
5. Stay professional - NO casual chat
6. If candidate says just "Hello", respond with Question 1
7. Track which question you're on to avoid repetition

FORBIDDEN BEHAVIORS:
- Do NOT say: "hey", "cool", "chat", "Oh", or use casual language
- Do NOT have conversations outside these 10 questions
- Do NOT ask random questions
- Do NOT give short one-word responses

RESUME CONTEXT:
{resume_content[:2000] if resume_content else 'No resume provided'}

JOB REQUIREMENTS:
{job_description[:1000] if job_description else f'Standard requirements for {position}'}

Remember: You are conducting a PROFESSIONAL STRUCTURED INTERVIEW. Stay on script!
"""
    
    return structured_content

def create_enhanced_kb_content(candidate_name, position, company, resume_content):
    """Create enhanced knowledge base content WITHOUT putting backslashes inside f-string expressions."""
    # Build any strings that contain backslashes/newlines first:
    resume_highlights = (
        f"Resume Highlights:\n{resume_content[:2000]}..."
        if resume_content else
        "No resume content available - focus on standard interview questions"
    )
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    experience_years = extract_experience_years(resume_content) if resume_content else "Not specified"

    skills_line = ", ".join(skills) if skills else "General software engineering skills"
    # Now the f-string only injects already-built variables (safe):
    return (
        "INTERVIEW SYSTEM CONFIGURATION\n"
        "==============================\n"
        f"MISSION: Conduct a professional, comprehensive interview for {candidate_name}\n"
        f"POSITION: {position}\n"
        f"COMPANY: {company}\n"
        "â±DURATION: 30-45 minutes\n"
        "MODE: AI-Powered Structured Interview\n\n"
        "CANDIDATE BACKGROUND\n"
        "====================\n"
        f"{resume_highlights}\n"
        f"Experience: {experience_years}\n"
        f"Key Skills: {skills_line}\n\n"
        "INTERVIEW FLOW\n"
        "==============\n"
        "1) Intro & Warm-up\n"
        "2) Skills Deep Dive\n"
        "3) Problem Solving\n"
        "4) Behavioral\n"
        "5) Wrap-up\n"
    )

def format_conversation_transcript(conversation_data, candidate_name):
    """Format conversation data into readable transcript"""
    transcript = f"Interview Transcript - {candidate_name}\n"
    transcript += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    transcript += "=" * 70 + "\n\n"
    
    for entry in conversation_data:
        try:
            timestamp = datetime.fromisoformat(entry['timestamp'].replace('Z', '+00:00'))
            time_str = timestamp.strftime('%H:%M:%S')
            speaker = 'AI Interviewer' if entry['speaker'] == 'avatar' else 'Candidate'
            content = entry['content'].strip()
            
            transcript += f"[{time_str}] {speaker}:\n{content}\n\n"
            
        except (KeyError, ValueError) as e:
            logger.warning(f"Error formatting entry: {e}")
            continue
    
    return transcript

def migrate_conversation_storage():
    """Migrate existing interview data to new conversation storage format"""
    from sqlalchemy import text
    
    session = SessionLocal()
    try:
        # Add new column if it doesn't exist
        try:
            session.execute(text("""
                ALTER TABLE candidates 
                ADD COLUMN interview_conversation_structured TEXT DEFAULT NULL
            """))
            session.commit()
            print("✅ Added interview_conversation_structured column")
        except Exception:
            print("ℹ️  Column interview_conversation_structured already exists")
        
        # Migrate existing data
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True,
            Candidate.interview_conversation_structured.is_(None)
        ).all()
        
        migrated = 0
        for candidate in candidates:
            try:
                # Try to reconstruct conversation from existing data
                qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
                conversation_data = []
                
                for i, qa in enumerate(qa_pairs):
                    if qa.get('question'):
                        conversation_data.append({
                            'id': f'q_{i}_{int(time.time())}',
                            'type': 'question',
                            'speaker': 'avatar',
                            'content': qa['question'],
                            'timestamp': qa.get('timestamp', datetime.now().isoformat()),
                            'sequence': len(conversation_data) + 1
                        })
                    
                    if qa.get('answer'):
                        conversation_data.append({
                            'id': f'a_{i}_{int(time.time())}',
                            'type': 'answer',
                            'speaker': 'candidate',
                            'content': qa['answer'],
                            'timestamp': qa.get('answered_at', datetime.now().isoformat()),
                            'sequence': len(conversation_data) + 1,
                            'linked_question_id': f'q_{i}_{int(time.time())}'
                        })
                
                if conversation_data:
                    candidate.interview_conversation_structured = json.dumps(conversation_data)
                    migrated += 1
            
            except Exception as e:
                print(f"❌ Failed to migrate candidate {candidate.id}: {e}")
        
        session.commit()
        print(f"✅ Migrated {migrated} candidates to new conversation storage")
        
    except Exception as e:
        print(f"❌ Migration failed: {e}")
        session.rollback()
    finally:
        session.close()

def periodic_interview_completion_check():
    """Periodically check for interviews that should be completed"""
    while True:
        try:
            time.sleep(300)  # Check every 5 minutes
            
            session = SessionLocal()
            try:
                # Find uncompleted interviews
                uncompleted = session.query(Candidate).filter(
                    Candidate.interview_started_at.isnot(None),
                    Candidate.interview_completed_at.is_(None)
                ).all()
                
                for candidate in uncompleted:
                    check_and_complete_interview(candidate.id)
                    
            finally:
                session.close()
                
        except Exception as e:
            logger.error(f"Error in periodic completion check: {e}")

# Start this when the app starts
completion_check_thread = threading.Thread(
    target=periodic_interview_completion_check,
    daemon=True
)
completion_check_thread.start()

def check_and_update_expired_interviews():
    """Automatically check and mark expired interviews"""
    session = SessionLocal()
    try:
        now = datetime.now()
        
        # Find expired interviews
        expired_candidates = session.query(Candidate).filter(
            Candidate.interview_expires_at < now,
            Candidate.interview_completed_at.is_(None),
            Candidate.interview_status != 'expired'
        ).all()
        
        for candidate in expired_candidates:
            candidate.interview_status = 'expired'
            candidate.final_status = 'Interview Link Expired'
            logger.info(f"Marked interview as expired for {candidate.name}")
        
        # Find abandoned interviews (no activity for 2 hours)
        two_hours_ago = now - timedelta(hours=2)
        abandoned_candidates = session.query(Candidate).filter(
            Candidate.interview_started_at.isnot(None),
            Candidate.interview_completed_at.is_(None),
            Candidate.interview_last_activity < two_hours_ago,
            Candidate.interview_status != 'abandoned'
        ).all()
        
        for candidate in abandoned_candidates:
            candidate.interview_status = 'abandoned'
            candidate.final_status = 'Interview Abandoned'
            logger.info(f"Marked interview as abandoned for {candidate.name}")
        
        session.commit()
        
    except Exception as e:
        logger.error(f"Error checking expired interviews: {e}")
        session.rollback()
    finally:
        session.close()
def analyze_answer_quality(answer_text):
    """Analyze individual answer quality"""
    score = 50  # Base score
    
    # Length bonus
    word_count = len(answer_text.split())
    if word_count > 100:
        score += 20
    elif word_count > 50:
        score += 10
    elif word_count < 10:
        score -= 20
    
    # Technical keywords bonus
    tech_keywords = ['implemented', 'developed', 'designed', 'architecture', 
                    'framework', 'database', 'algorithm', 'optimization']
    for keyword in tech_keywords:
        if keyword.lower() in answer_text.lower():
            score += 5
    
    # STAR method indicators
    star_keywords = ['situation', 'task', 'action', 'result', 'challenge', 'solution']
    for keyword in star_keywords:
        if keyword.lower() in answer_text.lower():
            score += 3
    
    return min(100, max(0, score))

# Add this improved scoring function to backend.py

def analyze_interview_with_ai_enhanced(qa_pairs, candidate):
    """Enhanced AI analysis with actual content evaluation"""
    try:
        # Check if OpenAI is available
        openai_key = os.getenv('OPENAI_API_KEY')
        if openai_key:
            import openai
            openai.api_key = openai_key
            
            # Prepare Q&A for analysis
            qa_text = "\n".join([
                f"Q: {qa.get('question', '')}\nA: {qa.get('answer', '')}" 
                for qa in qa_pairs if qa.get('answer')
            ])
            
            # Get AI analysis
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{
                    "role": "system",
                    "content": f"You are evaluating interview responses for {candidate.job_title} position. Score each category 0-100."
                }, {
                    "role": "user", 
                    "content": f"""Analyze these interview Q&As and provide scores:
                    
{qa_text}

Return JSON with:
- technical_score: (0-100)
- communication_score: (0-100)  
- problem_solving_score: (0-100)
- cultural_fit_score: (0-100)
- overall_score: (0-100)
- feedback: (detailed text feedback)
- strengths: (list of strengths)
- weaknesses: (list of areas to improve)"""
                }],
                temperature=0.3
            )
            
            result = json.loads(response.choices[0].message.content)
            return result
            
        else:
            # Fallback to enhanced rule-based scoring
            return analyze_with_enhanced_rules(qa_pairs, candidate)
            
    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        return analyze_with_enhanced_rules(qa_pairs, candidate)

def analyze_with_enhanced_rules(qa_pairs, candidate):
    """Enhanced rule-based scoring with better metrics"""
    if not qa_pairs:
        return {
            'technical_score': 0,
            'communication_score': 0,
            'problem_solving_score': 0,
            'cultural_fit_score': 0,
            'overall_score': 0,
            'confidence': 0,
            'feedback': 'No interview data available'
        }
    
    # Calculate metrics
    answered = sum(1 for qa in qa_pairs if qa.get('answer'))
    total = len(qa_pairs)
    completion_rate = (answered / total * 100) if total > 0 else 0
    
    # Calculate answer quality metrics
    answer_lengths = [len(qa.get('answer', '')) for qa in qa_pairs if qa.get('answer')]
    avg_length = sum(answer_lengths) / len(answer_lengths) if answer_lengths else 0
    
    # Response times
    response_times = [qa.get('response_time', 0) for qa in qa_pairs if qa.get('response_time')]
    avg_response_time = sum(response_times) / len(response_times) if response_times else 0
    
    # Calculate scores with better logic
    technical_score = min(100, max(0, 
        50 + # Base score
        (completion_rate * 0.2) + # 20% from completion
        (min(avg_length / 10, 30)) # Up to 30 points for detailed answers
    ))
    
    communication_score = min(100, max(0,
        40 + # Base score
        (completion_rate * 0.3) + # 30% from completion
        (30 if avg_response_time < 30 else 15) # Quick responses
    ))
    
    problem_solving_score = min(100, max(0,
        45 + # Base score
        (completion_rate * 0.25) + # 25% from completion
        (min(avg_length / 8, 30)) # Detail in answers
    ))
    
    cultural_fit_score = min(100, max(0,
        50 + # Base score
        (completion_rate * 0.3) + # 30% from completion
        (20 if answered >= total * 0.8 else 0) # Bonus for high completion
    ))
    
    overall_score = (
        technical_score * 0.35 +
        communication_score * 0.25 +
        problem_solving_score * 0.25 +
        cultural_fit_score * 0.15
    )

     # Generate strengths and weaknesses
    strengths = []
    weaknesses = []
    
    if completion_rate >= 80:
        strengths.append("High completion rate")
    else:
        weaknesses.append("Low completion rate")
    
    if avg_length > 150:
        strengths.append("Detailed responses")
    elif avg_length < 50:
        weaknesses.append("Brief responses")
    
    if avg_response_time < 15:
        strengths.append("Quick response time")
    elif avg_response_time > 30:
        weaknesses.append("Slow response time")
    
    return {
        'technical_score': round(technical_score),
        'communication_score': round(communication_score),
        'problem_solving_score': round(problem_solving_score),
        'cultural_fit_score': round(cultural_fit_score),
        'overall_score': round(overall_score),
        'confidence': 0.75,
        'feedback': f"""
Interview Analysis (Rule-Based):
- Completion Rate: {completion_rate:.1f}%
- Questions Answered: {answered}/{total}
- Average Response Detail: {'Detailed' if avg_length > 150 else 'Brief' if avg_length > 50 else 'Very Brief'}
- Response Speed: {'Quick' if avg_response_time < 15 else 'Moderate' if avg_response_time < 30 else 'Thoughtful'}

Overall: {'Strong candidate' if overall_score >= 70 else 'Potential candidate' if overall_score >= 50 else 'Needs improvement'}
"""
    }

# Update the trigger_auto_scoring function in backend.py:

def trigger_auto_scoring(candidate_id):
    """Automatically trigger AI scoring when interview completes"""
    def run_scoring():
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                logger.error(f"Candidate {candidate_id} not found for scoring")
                return
            
            # Update status to processing
            candidate.interview_ai_analysis_status = 'processing'
            session.commit()
            
            # Parse Q&A data
            qa_pairs = []
            try:
                # Try to parse Q&A pairs
                if candidate.interview_qa_pairs:
                    qa_pairs = json.loads(candidate.interview_qa_pairs)
                elif candidate.interview_questions_asked and candidate.interview_answers_given:
                    # Build Q&A pairs from separate fields
                    questions = json.loads(candidate.interview_questions_asked or '[]')
                    answers = json.loads(candidate.interview_answers_given or '[]')
                    
                    for i, question in enumerate(questions):
                        qa_pair = {
                            'question': question.get('text', ''),
                            'answer': answers[i].get('text', '') if i < len(answers) else '',
                            'timestamp': question.get('timestamp', None)
                        }
                        qa_pairs.append(qa_pair)
            except Exception as e:
                logger.error(f"Failed to parse Q&A data: {e}")
                qa_pairs = []
            
            # Perform analysis
            if len(qa_pairs) > 0:
                # Use your enhanced analysis function
                scores = analyze_interview_with_ai_enhanced(qa_pairs, candidate)
            else:
                # Default scores if no Q&A data
                scores = {
                    'technical_score': 0,
                    'communication_score': 0,
                    'problem_solving_score': 0,
                    'cultural_fit_score': 0,
                    'overall_score': 0,
                    'feedback': 'No interview data available for analysis',
                    'confidence': 0
                }
            
            # Update candidate with scores
            candidate.interview_ai_score = scores.get('overall_score', 0)
            candidate.interview_ai_technical_score = scores.get('technical_score', 0)
            candidate.interview_ai_communication_score = scores.get('communication_score', 0)
            candidate.interview_ai_problem_solving_score = scores.get('problem_solving_score', 0)
            candidate.interview_ai_cultural_fit_score = scores.get('cultural_fit_score', 0)
            candidate.interview_ai_overall_feedback = scores.get('feedback', '')
            
            # Extract strengths and weaknesses
            if 'strengths' in scores:
                candidate.interview_ai_strengths = json.dumps(scores['strengths'])
            if 'weaknesses' in scores:
                candidate.interview_ai_weaknesses = json.dumps(scores['weaknesses'])
            
            # Set final status based on score
            if candidate.interview_ai_score >= 70:
                candidate.interview_final_status = 'Passed'
                candidate.final_status = 'Interview Passed'
            else:
                candidate.interview_final_status = 'Failed'
                candidate.final_status = 'Interview Failed'
            
            # Mark analysis as complete
            candidate.interview_ai_analysis_status = 'completed'
            
            # Store completion timestamp
            if hasattr(candidate, 'interview_ai_analysis_completed_at'):
                candidate.interview_ai_analysis_completed_at = datetime.now()
            
            session.commit()
            
            # Clear cache to update frontend
            cache.delete_memoized(get_cached_candidates)
            
            logger.info(f"Auto-scoring completed for candidate {candidate_id}: {candidate.interview_ai_score}%")
            
            # Send notification if configured
            if hasattr(globals(), 'notify_scoring_complete'):
                notify_scoring_complete(candidate)
            
        except Exception as e:
            logger.error(f"Auto-scoring failed for candidate {candidate_id}: {e}", exc_info=True)
            if candidate:
                candidate.interview_ai_analysis_status = 'failed'
                session.commit()
        finally:
            session.close()
    
    # Run in background thread
    if 'executor' in globals():
        executor.submit(run_scoring)
    else:
        # Run directly if no executor
        run_scoring()

def trigger_ai_analysis(candidate_id):
    """Trigger REAL AI analysis for completed interview - NO RANDOM SCORES"""
    def run_analysis():
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                logger.error(f"Candidate {candidate_id} not found")
                return
            
            # Update status
            candidate.interview_ai_analysis_status = 'processing'
            session.commit()
            
            # Parse Q&A data
            questions = json.loads(candidate.interview_questions_asked or '[]')
            answers = json.loads(candidate.interview_answers_given or '[]')
            
            # Build Q&A pairs for analysis
            qa_pairs = []
            for i, question in enumerate(questions):
                qa_pair = {
                    'question': question.get('text', '') if isinstance(question, dict) else str(question),
                    'answer': ''
                }
                if i < len(answers):
                    answer = answers[i]
                    qa_pair['answer'] = answer.get('text', '') if isinstance(answer, dict) else str(answer)
                qa_pairs.append(qa_pair)
            
            # CHECK FOR INVALID/TEST RESPONSES
            has_invalid_responses = False
            invalid_patterns = ['INIT_INTERVIEW', 'TEST_RESPONSE', 'undefined', 'null']
            
            for qa in qa_pairs:
                answer_text = qa.get('answer', '').strip()
                if any(pattern in answer_text for pattern in invalid_patterns) or len(answer_text) < 5:
                    has_invalid_responses = True
                    logger.warning(f"Invalid response detected: {answer_text[:50]}")
                    break
            
            # DYNAMIC SCORING BASED ON ACTUAL CONTENT
            if has_invalid_responses or len(qa_pairs) == 0:
                # Failed interview - invalid or no responses
                candidate.interview_ai_score = 0
                candidate.interview_ai_technical_score = 0
                candidate.interview_ai_communication_score = 0
                candidate.interview_ai_problem_solving_score = 0
                candidate.interview_ai_cultural_fit_score = 0
                candidate.interview_ai_overall_feedback = """
Interview Analysis: FAILED
- No valid responses provided
- Interview appears to be incomplete or contains test data
- Candidate did not properly complete the interview

Recommendation: Schedule a new interview
"""
                candidate.interview_final_status = 'Failed - Invalid Response'
                
            else:
                # ACTUAL DYNAMIC ANALYSIS
                scores = analyze_interview_content(qa_pairs, candidate)
                
                candidate.interview_ai_score = scores['overall']
                candidate.interview_ai_technical_score = scores['technical']
                candidate.interview_ai_communication_score = scores['communication']
                candidate.interview_ai_problem_solving_score = scores['problem_solving']
                candidate.interview_ai_cultural_fit_score = scores['cultural_fit']
                candidate.interview_ai_overall_feedback = scores['feedback']
                
                # Set final status based on ACTUAL score
                if scores['overall'] >= 70:
                    candidate.interview_final_status = 'Recommended'
                elif scores['overall'] >= 50:
                    candidate.interview_final_status = 'Review Required'
                else:
                    candidate.interview_final_status = 'Not Recommended'
            
            candidate.interview_ai_analysis_status = 'completed'
            session.commit()
            
            logger.info(f"Dynamic analysis completed for {candidate.name}: {candidate.interview_ai_score}%")
            
            # Clear cache
            cache.delete_memoized(get_cached_candidates)
            
        except Exception as e:
            logger.error(f"AI analysis failed for candidate {candidate_id}: {e}", exc_info=True)
            if candidate:
                candidate.interview_ai_analysis_status = 'failed'
                candidate.interview_ai_overall_feedback = f"Analysis failed: {str(e)}"
                session.commit()
        finally:
            session.close()
    
    # Run in background thread
    executor.submit(run_analysis)

def analyze_interview_content(qa_pairs, candidate):
    """REAL content analysis - not random!"""
    
    # Initialize scores
    scores = {
        'technical': 0,
        'communication': 0,
        'problem_solving': 0,
        'cultural_fit': 0,
        'overall': 0
    }
    
    if not qa_pairs:
        return scores
    
    # Technical keywords to look for
    tech_keywords = [
        'implement', 'develop', 'design', 'architecture', 'algorithm',
        'database', 'api', 'framework', 'optimize', 'scale', 'debug',
        'testing', 'deployment', 'code', 'programming', 'software'
    ]
    
    soft_keywords = [
        'team', 'collaborate', 'communicate', 'lead', 'manage',
        'problem', 'solution', 'challenge', 'learn', 'adapt'
    ]
    
    total_score = 0
    answered_count = 0
    total_word_count = 0
    technical_mentions = 0
    soft_skill_mentions = 0
    
    for qa in qa_pairs:
        answer = qa.get('answer', '').lower()
        question = qa.get('question', '').lower()
        
        if not answer or len(answer) < 5:
            continue
            
        answered_count += 1
        words = answer.split()
        word_count = len(words)
        total_word_count += word_count
        
        # Score based on answer length (longer = more detailed)
        length_score = min(30, word_count / 3)  # Max 30 points for length
        
        # Check for technical keywords
        tech_found = sum(1 for kw in tech_keywords if kw in answer)
        technical_mentions += tech_found
        tech_score = min(40, tech_found * 10)  # Max 40 points for technical content
        
        # Check for soft skills
        soft_found = sum(1 for kw in soft_keywords if kw in answer)
        soft_skill_mentions += soft_found
        soft_score = min(30, soft_found * 10)  # Max 30 points for soft skills
        
        # Calculate answer score
        answer_score = length_score + tech_score + soft_score
        
        # Categorize score
        if 'technical' in question or 'code' in question or 'implement' in question:
            scores['technical'] += answer_score
        elif 'team' in question or 'collaborate' in question:
            scores['cultural_fit'] += answer_score
        elif 'problem' in question or 'challenge' in question:
            scores['problem_solving'] += answer_score
        else:
            scores['communication'] += answer_score
        
        total_score += answer_score
    
    # Calculate completion rate
    completion_rate = (answered_count / len(qa_pairs)) * 100 if qa_pairs else 0
    
    # Average word count per answer
    avg_word_count = total_word_count / answered_count if answered_count > 0 else 0
    
    # Normalize scores (0-100 scale)
    num_questions = len(qa_pairs)
    if num_questions > 0:
        for key in ['technical', 'communication', 'problem_solving', 'cultural_fit']:
            scores[key] = min(100, (scores[key] / num_questions) * 2)
    
    # Calculate overall score with penalties
    base_score = (scores['technical'] * 0.35 + 
                  scores['communication'] * 0.25 + 
                  scores['problem_solving'] * 0.25 + 
                  scores['cultural_fit'] * 0.15)
    
    # Apply penalties
    if completion_rate < 50:
        base_score *= 0.5  # 50% penalty for low completion
    elif completion_rate < 80:
        base_score *= 0.8  # 20% penalty
    
    if avg_word_count < 20:
        base_score *= 0.7  # 30% penalty for very brief answers
    
    scores['overall'] = min(100, max(0, base_score))
    
    # Generate detailed feedback
    scores['feedback'] = f"""
INTERVIEW ANALYSIS REPORT (Dynamic Scoring)
==========================================
Candidate: {candidate.name}
Position: {candidate.job_title}
Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}

SCORING METRICS:
- Completion Rate: {completion_rate:.1f}%
- Questions Answered: {answered_count}/{len(qa_pairs)}
- Average Response Length: {avg_word_count:.0f} words
- Technical Keywords Found: {technical_mentions}
- Soft Skill Keywords Found: {soft_skill_mentions}

DETAILED SCORES:
- Technical Skills: {scores['technical']:.1f}/100
- Communication: {scores['communication']:.1f}/100
- Problem Solving: {scores['problem_solving']:.1f}/100
- Cultural Fit: {scores['cultural_fit']:.1f}/100

OVERALL SCORE: {scores['overall']:.1f}/100

ASSESSMENT:
{get_assessment_text(scores['overall'], completion_rate, avg_word_count)}

RECOMMENDATION: {'Highly Recommended' if scores['overall'] >= 80 else 'Recommended' if scores['overall'] >= 70 else 'Consider for Further Review' if scores['overall'] >= 50 else 'Not Recommended'}
"""
    
    return scores

def get_assessment_text(overall_score, completion_rate, avg_word_count):
    """Generate assessment text based on metrics"""
    
    if overall_score >= 80:
        return """The candidate demonstrated excellent performance across all areas. 
Their responses were detailed, thoughtful, and showed strong technical knowledge 
combined with good communication skills. They are highly recommended for this position."""
    
    elif overall_score >= 70:
        return """The candidate showed good performance with solid responses to most questions. 
They demonstrated adequate technical knowledge and communication skills. 
Consider for the next round of interviews."""
    
    elif overall_score >= 50:
        return f"""The candidate showed moderate performance with some areas of concern.
Completion rate: {completion_rate:.0f}%
Average response detail: {avg_word_count:.0f} words
Additional evaluation may be needed to make a final decision."""
    
    else:
        return f"""The candidate's performance was below expectations.
Key concerns:
- Low completion rate: {completion_rate:.0f}%
- Brief responses: {avg_word_count:.0f} words average
- Limited demonstration of required skills
Not recommended for this position."""

def analyze_with_openai(qa_pairs, candidate):
    """Use OpenAI for more intelligent analysis"""
    try:
        import openai
        openai.api_key = os.getenv('OPENAI_API_KEY')
        
        if not openai.api_key:
            return None
        
        # Prepare Q&A text
        qa_text = "\n".join([
            f"Q: {qa['question']}\nA: {qa['answer']}" 
            for qa in qa_pairs if qa.get('answer')
        ])
        
        prompt = f"""
        Analyze this interview for {candidate.job_title} position.
        Score each aspect 0-100 based on ACTUAL content quality.
        
        Interview:
        {qa_text}
        
        Return JSON with: technical, communication, problem_solving, cultural_fit, overall scores and feedback.
        """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert interview evaluator."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        
        # Parse response and return scores
        import json
        result = json.loads(response.choices[0].message.content)
        return result
        
    except Exception as e:
        logger.error(f"OpenAI analysis failed: {e}")
        return None

def analyze_interview_with_ai(qa_pairs, candidate):
    """Analyze interview Q&A with AI (or rule-based fallback)"""
    try:
        # Try AI analysis first (you can integrate OpenAI, Claude, etc.)
        if os.getenv('OPENAI_API_KEY'):
            return analyze_with_openai(qa_pairs, candidate)
        else:
            # Fallback to rule-based scoring
            return analyze_with_rules(qa_pairs, candidate)
    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        return analyze_with_rules(qa_pairs, candidate)


def analyze_with_rules(qa_pairs, candidate):
    """Rule-based interview scoring"""
    scores = {
        'technical': 0,
        'communication': 0,
        'problem_solving': 0,
        'cultural_fit': 0,
        'overall': 0,
        'confidence': 0.85,
        'feedback': ''
    }
    
    if not qa_pairs:
        return scores
    
    # Calculate completion rate
    answered = sum(1 for qa in qa_pairs if qa.get('answer'))
    total = len(qa_pairs)
    completion_rate = (answered / total * 100) if total > 0 else 0
    
    # Calculate average response time
    response_times = [qa.get('response_time', 0) for qa in qa_pairs if qa.get('response_time')]
    avg_response_time = sum(response_times) / len(response_times) if response_times else 0
    
    # Score based on completion
    if completion_rate >= 90:
        scores['communication'] = 85
    elif completion_rate >= 70:
        scores['communication'] = 70
    else:
        scores['communication'] = 50
    
    # Score based on response times
    if avg_response_time < 10:  # Quick responses
        scores['problem_solving'] = 80
    elif avg_response_time < 30:
        scores['problem_solving'] = 70
    else:
        scores['problem_solving'] = 60
    
    # Analyze answer lengths
    answer_lengths = [len(qa.get('answer', '')) for qa in qa_pairs if qa.get('answer')]
    avg_length = sum(answer_lengths) / len(answer_lengths) if answer_lengths else 0
    
    if avg_length > 100:  # Detailed answers
        scores['technical'] = 75
        scores['cultural_fit'] = 70
    else:
        scores['technical'] = 60
        scores['cultural_fit'] = 60
    
    # Calculate overall score
    scores['overall'] = (
        scores['technical'] * 0.35 +
        scores['communication'] * 0.25 +
        scores['problem_solving'] * 0.25 +
        scores['cultural_fit'] * 0.15
    )
    
    # Generate feedback
    scores['feedback'] = f"""
Interview Analysis Summary:
- Completion Rate: {completion_rate:.1f}%
- Questions Answered: {answered}/{total}
- Average Response Time: {avg_response_time:.1f} seconds
- Average Answer Length: {avg_length:.0f} characters

Strengths:
{'- Good completion rate' if completion_rate >= 80 else ''}
{'- Quick response times' if avg_response_time < 15 else ''}
{'- Detailed answers provided' if avg_length > 100 else ''}

Areas for Improvement:
{'- Complete all questions' if completion_rate < 100 else ''}
{'- Provide more detailed responses' if avg_length < 50 else ''}

Overall Assessment: {'Recommended' if scores['overall'] >= 70 else 'Needs Further Evaluation'}
"""
    
    return scores

def get_interview_status(candidate):
    """Determine current interview status"""
    if candidate.interview_ai_analysis_status == 'completed':
        if candidate.interview_ai_score >= 70:
            return 'passed'
        else:
            return 'failed'
    elif candidate.interview_ai_analysis_status == 'processing':
        return 'analyzing'
    elif candidate.interview_completed_at:
        return 'completed'
    elif candidate.interview_started_at:
        if candidate.interview_last_activity:
            time_since = (datetime.now() - candidate.interview_last_activity).total_seconds()
            if time_since < 300:  # Active within last 5 minutes
                return 'in_progress'
            else:
                return 'inactive'
        return 'in_progress'
    elif candidate.interview_link_clicked:
        return 'link_clicked'
    elif candidate.interview_scheduled:
        return 'scheduled'
    else:
        return 'not_started'

# In backend.py, update the complete_interview_auto function:


def complete_interview_auto(candidate_id):
    """Auto-complete interview and trigger scoring when all questions are answered"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate or candidate.interview_completed_at:
            return
        
        # Mark interview as completed
        candidate.interview_completed_at = datetime.now()
        candidate.interview_progress_percentage = 100
        
        # Calculate duration
        if candidate.interview_started_at:
            duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
            candidate.interview_duration = int(duration)
        
        # Set analysis status to pending
        candidate.interview_ai_analysis_status = 'pending'
        
        session.commit()
        
        # Trigger auto-scoring immediately
        trigger_auto_scoring(candidate_id)
        
        logger.info(f"Interview auto-completed and scoring triggered for candidate {candidate_id}")
        
    except Exception as e:
        logger.error(f"Error auto-completing interview: {e}")
        session.rollback()
    finally:
        session.close()


def schedule_auto_scoring(candidate_id, delay_minutes=45):
    """Schedule automatic scoring if interview doesn't complete normally"""
    def check_and_score():
        time.sleep(delay_minutes * 60)
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            
            # Check if interview is still incomplete
            if candidate and candidate.interview_started_at and not candidate.interview_completed_at:
                # Check for recent activity
                if candidate.interview_last_activity:
                    time_since = (datetime.now() - candidate.interview_last_activity).total_seconds()
                    if time_since > 600:  # No activity for 10 minutes
                        # Force complete and score
                        complete_interview_auto(candidate_id)
        except Exception as e:
            logger.error(f"Error in scheduled scoring: {e}")
        finally:
            session.close()
    
    # Schedule in background
    threading.Thread(target=check_and_score, daemon=True).start()


def send_realtime_update(candidate_id, data):
    """Send real-time updates to frontend via WebSocket or polling"""
    # Store update for polling
    update_key = f"interview_update_{candidate_id}"
    cache.set(update_key, json.dumps(data), timeout=60)

def notify_scoring_complete(candidate):
    """Send notification when scoring is complete"""
    try:
        # Send email notification
        if candidate.email:
            subject = f"Interview Analysis Complete - {candidate.name}"
            body = f"""
            Interview analysis has been completed for {candidate.name}.
            
            Position: {candidate.job_title}
            Overall Score: {candidate.interview_ai_score:.1f}%
            Recommendation: {candidate.interview_final_status}
            
            View full results in the dashboard.
            """
            
            # Send to HR/Admin
            admin_email = os.getenv('ADMIN_EMAIL')
            if admin_email:
                send_email(admin_email, subject, body)
        
        logger.info(f"Scoring notification sent for candidate {candidate.id}")
        
    except Exception as e:
        logger.error(f"Error sending scoring notification: {e}")

def log_interview_activity(candidate_id, activity_type, data):
    """Log all interview activities for audit trail"""
    try:
        log_entry = {
            'candidate_id': candidate_id,
            'activity': activity_type,
            'timestamp': datetime.now().isoformat(),
            'data': data
        }
        
        # Store in logs directory
        log_dir = os.path.join('logs', 'interviews', str(candidate_id))
        os.makedirs(log_dir, exist_ok=True)
        
        log_file = os.path.join(log_dir, 'activity.jsonl')
        with open(log_file, 'a') as f:
            f.write(json.dumps(log_entry) + '\n')
            
    except Exception as e:
        logger.error(f"Error logging activity: {e}")

def calculate_time_difference(timestamp1, timestamp2):
    """Calculate time difference between two timestamps in seconds"""
    if not timestamp1 or not timestamp2:
        return None
    try:
        t1 = datetime.fromisoformat(timestamp1.replace('Z', '+00:00'))
        t2 = datetime.fromisoformat(timestamp2.replace('Z', '+00:00'))
        return abs((t2 - t1).total_seconds())
    except:
        return None

def upload_to_cloud_storage(local_path, filename):
    """Upload recording to cloud storage (implement based on your provider)"""
    try:
        # Example for S3
        if os.getenv('AWS_ACCESS_KEY_ID'):
            import boto3
            s3 = boto3.client('s3')
            bucket = os.getenv('S3_BUCKET_NAME', 'interview-recordings')
            key = f"interviews/{filename}"
            
            s3.upload_file(local_path, bucket, key)
            return f"https://{bucket}.s3.amazonaws.com/{key}"
       
        return None
        
    except Exception as e:
        logger.error(f"Cloud upload failed: {e}")
        return None

def _ok_preflight():
    resp = jsonify({})
    resp.status_code = 200
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    return resp

def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path


def _append_jsonl(path, obj):
    with open(path, 'a', encoding='utf-8') as f:
        f.write(json.dumps(obj, ensure_ascii=False) + '\n')

def extract_resume_content(resume_path):
    """Extract text content from resume with better error handling"""
    try:
        if not os.path.exists(resume_path):
            logger.error(f"Resume file not found: {resume_path}")
            return ""
        
        file_ext = os.path.splitext(resume_path)[1].lower()
        logger.info(f"Extracting resume: {resume_path} (type: {file_ext})")
        
        resume_text = ""
        
        if file_ext == '.pdf':
            # Try PyPDF2 first
            try:
                import PyPDF2
                with open(resume_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        resume_text += page.extract_text() + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with PyPDF2: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
            
            # Try pdfplumber as fallback
            try:
                import pdfplumber
                with pdfplumber.open(resume_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            resume_text += page_text + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with pdfplumber: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
                    
        elif file_ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(resume_path)
                resume_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                logger.info(f"DOCX extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                
        elif file_ext == '.txt':
            try:
                with open(resume_path, 'r', encoding='utf-8') as file:
                    resume_text = file.read()
                logger.info(f"TXT extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except UnicodeDecodeError:
                with open(resume_path, 'r', encoding='latin-1') as file:
                    resume_text = file.read()
                return resume_text.strip()
        
        # If we couldn't extract anything, return empty string
        if not resume_text:
            logger.error(f"Failed to extract any text from {resume_path}")
            
    except Exception as e:
        logger.error(f"Resume extraction failed: {e}", exc_info=True)
    
    return ""

def extract_pdf_content(pdf_path):
    """Extract content from PDF using multiple methods"""
    try:
        # Method 1: Try PyPDF2
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                if text.strip():
                    return text.strip()
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        # Method 2: Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # Method 3: Try pymupdf (fitz)
        try:
            import fitz  # pymupdf
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning(f"pymupdf extraction failed: {str(e)}")
            
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
    
    return ""



def extract_docx_content(docx_path):
    """Extract content from DOCX files"""
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text.strip()
        
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return ""


def extract_txt_content(txt_path):
    """Extract content from text files"""
    try:
        # Try UTF-8 first
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Fallback to latin-1
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
    except Exception as e:
        logger.error(f"TXT extraction failed: {str(e)}")
        return ""

def generate_kb_recommendations(checks):
    """Generate recommendations based on validation checks"""
    recommendations = []
    
    if not checks.get("heygen_api_configured"):
        recommendations.append("Set HEYGEN_API_KEY environment variable")
    
    if not checks.get("resume_file_exists"):
        recommendations.append("Upload candidate resume file")
    elif not checks.get("resume_extractable"):
        recommendations.append("Resume file exists but content extraction failed - check file format")
    
    if not checks.get("interview_token_exists"):
        recommendations.append("Generate interview token for candidate")
    
    if not checks.get("company_name_available"):
        recommendations.append("Set COMPANY_NAME environment variable or add company to candidate record")
    
    if checks.get("kb_already_exists"):
        recommendations.append("Knowledge base already exists - consider updating instead of recreating")
    
    return recommendations


def generate_custom_interview_prompt(candidate, resume_content, job_description):
    """Generate custom interview prompt based on resume and job"""
    return f"""
    You are interviewing {candidate.name} for {candidate.job_title} position.
    
    Based on their resume, ask questions about:
    1. Their experience with technologies mentioned in their resume
    2. Projects they've worked on
    3. Challenges they've faced
    4. Their approach to problem-solving
    
    Based on the job requirements, assess:
    1. Technical skills required for the role
    2. Soft skills and communication
    3. Cultural fit
    4. Career goals alignment
    
    Keep the interview conversational and professional.
    Ask follow-up questions based on their responses.
    """


# 2. Enhanced resume extraction function with better error handling

def extract_resume_content(resume_path):
    """Extract text content from resume with better error handling"""
    try:
        if not os.path.exists(resume_path):
            logger.error(f"Resume file not found: {resume_path}")
            return ""
        
        file_ext = os.path.splitext(resume_path)[1].lower()
        logger.info(f"Extracting resume: {resume_path} (type: {file_ext})")
        
        resume_text = ""
        
        if file_ext == '.pdf':
            # Try PyPDF2 first
            try:
                import PyPDF2
                with open(resume_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        resume_text += page.extract_text() + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with PyPDF2: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
            
            # Try pdfplumber as fallback
            try:
                import pdfplumber
                with pdfplumber.open(resume_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            resume_text += page_text + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with pdfplumber: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
                    
        elif file_ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(resume_path)
                resume_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                logger.info(f"DOCX extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                
        elif file_ext == '.txt':
            try:
                with open(resume_path, 'r', encoding='utf-8') as file:
                    resume_text = file.read()
                logger.info(f"TXT extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except UnicodeDecodeError:
                with open(resume_path, 'r', encoding='latin-1') as file:
                    resume_text = file.read()
                return resume_text.strip()
        
        # If we couldn't extract anything, return empty string
        if not resume_text:
            logger.error(f"Failed to extract any text from {resume_path}")
            
    except Exception as e:
        logger.error(f"Resume extraction failed: {e}", exc_info=True)
    
    return ""

def interview_auto_recovery_system():
    """Continuously monitor and fix incomplete interviews"""
    while True:
        time.sleep(120)  # Check every 2 minutes
        
        session = SessionLocal()
        try:
            from datetime import datetime, timedelta
            
            # Find potentially stuck interviews
            now = datetime.now()
            
            # Case 1: Started but not completed after 1 hour
            one_hour_ago = now - timedelta(hours=1)
            stuck_interviews = session.query(Candidate).filter(
                Candidate.interview_started_at.isnot(None),
                Candidate.interview_completed_at.is_(None),
                Candidate.interview_started_at < one_hour_ago
            ).all()
            
            for candidate in stuck_interviews:
                # Check if has Q&A data
                has_qa_data = False
                try:
                    qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
                    questions = json.loads(candidate.interview_questions_asked or '[]')
                    has_qa_data = len(qa_pairs) > 0 or len(questions) > 0
                except:
                    pass
                
                if has_qa_data or (now - candidate.interview_started_at).total_seconds() > 7200:
                    candidate.interview_completed_at = now
                    candidate.interview_status = 'completed'
                    candidate.interview_progress_percentage = 100
                    candidate.final_status = 'Interview Completed - Auto Recovery'
                    candidate.interview_ai_analysis_status = 'pending'
                    
                    if candidate.interview_started_at:
                        duration = (now - candidate.interview_started_at).total_seconds()
                        candidate.interview_duration = int(duration)
                    
                    logger.info(f"Auto-recovered interview for {candidate.name} (ID: {candidate.id})")
            
            # Case 2: Has 100% progress but no completion timestamp
            incomplete_100 = session.query(Candidate).filter(
                Candidate.interview_progress_percentage >= 100,
                Candidate.interview_completed_at.is_(None)
            ).all()
            
            for candidate in incomplete_100:
                candidate.interview_completed_at = now
                candidate.interview_status = 'completed'
                candidate.final_status = 'Interview Completed - Progress 100%'
                candidate.interview_ai_analysis_status = 'pending'
                logger.info(f"Completed interview at 100% progress for {candidate.name}")
            
            session.commit()
            
            # Clear cache after updates
            if stuck_interviews or incomplete_100:
                cache.clear()
            
        except Exception as e:
            logger.error(f"Auto-recovery error: {e}")
            session.rollback()
        finally:
            session.close()

def notify_admin(subject, message, error_details=None):
    """Send critical notifications to admin"""
    try:
        admin_email = os.getenv('ADMIN_EMAIL')
        if not admin_email:
            logger.warning("ADMIN_EMAIL not set, skipping notification")
            return
        
        from email_util import send_email
        
        body_html = f"""
        <html>
            <body>
                <h2>TalentFlow AI Alert: {subject}</h2>
                <p>{message}</p>
                {f'<pre>{error_details}</pre>' if error_details else ''}
                <p>Time: {datetime.now().isoformat()}</p>
            </body>
        </html>
        """
        
        send_email(admin_email, f"[TalentFlow Alert] {subject}", body_html)
        
    except Exception as e:
        logger.error(f"Failed to send admin notification: {e}")

def run_bulk_scraping_with_monitoring():
    """Wrapper to run bulk scraping with monitoring"""
    start_time = time.time()
    
    try:
        logger.info("Starting bulk scraping for all pending assessments")
        
        # Import and run the bulk scraping function
        try:
            from testlify_results_scraper import scrape_all_pending_assessments
        except ImportError as e:
            logger.error(f"Failed to import scraper: {e}")
            notify_admin(
                "Scraper Import Error",
                f"Could not import results scraper: {str(e)}. Please ensure testlify_results_scraper.py is available."
            )
            return
        
        # Run the async scraping function
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            results_summary = loop.run_until_complete(scrape_all_pending_assessments())
        finally:
            loop.close()
        
        duration = time.time() - start_time
        total_candidates = sum(results_summary.values()) if isinstance(results_summary, dict) else 0
        
        logger.info(f"Bulk scraping completed in {duration:.2f} seconds. Processed {len(results_summary)} assessments, {total_candidates} candidates.")
        
        # Send success notification
        if isinstance(results_summary, dict):
            summary_text = "\n".join([f"- {assessment}: {count} candidates" for assessment, count in results_summary.items()])
        else:
            summary_text = f"Processed {total_candidates} total candidates"
            
        notify_admin(
            "Bulk Assessment Results Scraping Completed",
            f"Assessments processed: {len(results_summary) if isinstance(results_summary, dict) else 'Unknown'}\nTotal candidates: {total_candidates}\nDuration: {duration:.2f} seconds\n\nBreakdown:\n{summary_text}"
        )
        
    except Exception as e:
        duration = time.time() - start_time
        error_msg = f"Bulk scraping failed after {duration:.2f} seconds"
        logger.error(error_msg, exc_info=True)
        
        # Send failure notification
        notify_admin(
            "Bulk Assessment Results Scraping Failed",
            error_msg,
            error_details=traceback.format_exc()
        )

def test_routes():
    """Test if routes are properly registered"""
    print("ðŸ“‹ Registered Routes:")
    for rule in app.url_map.iter_rules():
        print(f"  {rule.methods} {rule.rule} -> {rule.endpoint}")

class InterviewErrorRecovery:
    """Handle interview system errors and recovery"""
    
    @staticmethod
    def recover_incomplete_interviews():
        """Recover and complete incomplete interviews"""
        session = SessionLocal()
        try:
            # Find stuck interviews (started > 2 hours ago, not completed)
            cutoff_time = datetime.now() - timedelta(hours=2)
            stuck_interviews = session.query(Candidate).filter(
                Candidate.interview_started_at < cutoff_time,
                Candidate.interview_completed_at.is_(None)
            ).all()
            
            for candidate in stuck_interviews:
                # Check last activity
                if candidate.interview_last_activity:
                    time_since = (datetime.now() - candidate.interview_last_activity).total_seconds()
                    if time_since > 3600:  # No activity for 1 hour
                        # Force complete
                        candidate.interview_completed_at = datetime.now()
                        candidate.interview_ai_analysis_status = 'pending'
                        
                        # Trigger scoring with what we have
                        trigger_auto_scoring(candidate.id)
                        
                        logger.warning(f"Force completed stuck interview for candidate {candidate.id}")
            
            session.commit()
            
        except Exception as e:
            logger.error(f"Error recovering interviews: {e}")
        finally:
            session.close()
    
    @staticmethod
    def validate_interview_data(candidate_id):
        """Validate and fix interview data integrity"""
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return False
            
            issues_fixed = []
            
            # Fix Q&A pairs
            try:
                qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            except:
                qa_pairs = []
                candidate.interview_qa_pairs = '[]'
                issues_fixed.append('Reset invalid Q&A data')
            
            # Fix counts
            if candidate.interview_total_questions != len(qa_pairs):
                candidate.interview_total_questions = len(qa_pairs)
                issues_fixed.append('Fixed question count')
            
            answered = sum(1 for qa in qa_pairs if qa.get('answer'))
            if candidate.interview_answered_questions != answered:
                candidate.interview_answered_questions = answered
                issues_fixed.append('Fixed answer count')
            
            # Fix progress
            if candidate.interview_total_questions > 0:
                progress = (answered / candidate.interview_total_questions) * 100
                if abs(candidate.interview_progress_percentage - progress) > 1:
                    candidate.interview_progress_percentage = progress
                    issues_fixed.append('Fixed progress percentage')
            
            # Fix status inconsistencies
            if candidate.interview_completed_at and not candidate.interview_started_at:
                candidate.interview_started_at = candidate.interview_completed_at - timedelta(minutes=30)
                issues_fixed.append('Fixed missing start time')
            
            if issues_fixed:
                session.commit()
                logger.info(f"Fixed issues for candidate {candidate_id}: {', '.join(issues_fixed)}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating interview data: {e}")
            return False
        finally:
            session.close()

def start_health_monitoring():
    """Start periodic health monitoring"""
    def monitor():
        while True:
            time.sleep(300)  # Check every 5 minutes
            try:
                InterviewErrorRecovery.recover_incomplete_interviews()
            except Exception as e:
                logger.error(f"Health monitor error: {e}")
    
    thread = threading.Thread(target=monitor, daemon=True)
    thread.start()

class InterviewCompletionHandler:
    """Handles all interview completion logic with guaranteed database updates"""
    
    @staticmethod
    def complete_interview(token: str, trigger_source: str = "unknown") -> dict:
        """
        Complete an interview and ensure database is updated
        Returns: dict with success status and details
        """
        from datetime import datetime, timezone
        
        session = SessionLocal()
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                # Get candidate with lock to prevent concurrent updates
                candidate = session.query(Candidate).filter_by(
                    interview_token=token
                ).with_for_update().first()
                
                if not candidate:
                    return {"success": False, "error": "Candidate not found"}
                
                # Check if already completed
                if candidate.interview_completed_at:
                    logger.info(f"Interview already completed for {candidate.name}")
                    return {
                        "success": True,
                        "already_completed": True,
                        "completed_at": candidate.interview_completed_at.isoformat()
                    }
                
                # Set completion fields
                completion_time = datetime.now(timezone.utc)
                candidate.interview_completed_at = completion_time
                candidate.interview_status = 'completed'
                candidate.interview_progress_percentage = 100.0
                candidate.final_status = f'Interview Completed - {trigger_source}'
                
                # Calculate duration if started
                if candidate.interview_started_at:
                    # Handle timezone-aware comparison
                    start_time = candidate.interview_started_at
                    if start_time.tzinfo is None:
                        start_time = start_time.replace(tzinfo=timezone.utc)
                    duration = (completion_time - start_time).total_seconds()
                    candidate.interview_duration = int(duration)
                else:
                    candidate.interview_duration = 0
                
                # Set for AI analysis
                candidate.interview_ai_analysis_status = 'pending'
                candidate.interview_auto_score_triggered = False
                
                # Update Q&A stats if needed
                if not candidate.interview_total_questions:
                    candidate.interview_total_questions = 10  # Default
                if not candidate.interview_answered_questions:
                    candidate.interview_answered_questions = candidate.interview_total_questions
                
                # Force commit with explicit flush
                session.flush()
                session.commit()
                
                # Verify the update worked
                session.refresh(candidate)
                if candidate.interview_completed_at:
                    logger.info(f"Interview completed successfully for {candidate.name} at {completion_time}")
                    
                    # Trigger AI scoring in background
                    try:
                        executor.submit(trigger_auto_scoring, candidate.id)
                    except Exception as e:
                        logger.error(f"Failed to trigger scoring: {e}")
                    
                    return {
                        "success": True,
                        "completed_at": candidate.interview_completed_at.isoformat(),
                        "duration": candidate.interview_duration,
                        "trigger_source": trigger_source
                    }
                else:
                    raise Exception("Completion timestamp not saved")
                    
            except Exception as e:
                session.rollback()
                retry_count += 1
                logger.error(f"Attempt {retry_count} failed: {e}")
                
                if retry_count >= max_retries:
                    # Last resort: try direct SQL update
                    try:
                        session.close()
                        session = SessionLocal()
                        result = session.execute(
                            text("""
                                UPDATE candidates 
                                SET interview_completed_at = :completed_at,
                                    interview_status = 'completed',
                                    interview_progress_percentage = 100,
                                    final_status = :final_status,
                                    interview_ai_analysis_status = 'pending'
                                WHERE interview_token = :token
                            """),
                            {
                                "completed_at": datetime.now(),
                                "final_status": f"Interview Completed - {trigger_source}",
                                "token": token
                            }
                        )
                        session.commit()
                        
                        if result.rowcount > 0:
                            logger.info(f"Completed via direct SQL for token {token}")
                            return {"success": True, "method": "direct_sql"}
                    except Exception as sql_error:
                        logger.error(f"Direct SQL also failed: {sql_error}")
                    
                    return {"success": False, "error": str(e)}
                
                time.sleep(0.5)  # Brief delay before retry
            finally:
                if retry_count >= max_retries:
                    session.close()
        
        session.close()
        return {"success": False, "error": "Max retries exceeded"}
# Create global instance
completion_handler = InterviewCompletionHandler()

class AutomaticCompletionMonitor:
    """Monitor and automatically complete interviews"""
    
    def __init__(self):
        self.is_running = False
        self.check_interval = 30  # seconds
        self.thread = None
    
    def start(self):
        if self.is_running:
            return
        self.is_running = True
        self.thread = threading.Thread(target=self._monitor_loop, daemon=True)
        self.thread.start()
        logger.info("Automatic completion monitor started")
    
    def stop(self):
        self.is_running = False
        if self.thread:
            self.thread.join(timeout=5)
    
    def _monitor_loop(self):
        while self.is_running:
            try:
                self._check_all_interviews()
                time.sleep(self.check_interval)
            except Exception as e:
                logger.error(f"Monitor error: {e}")
                time.sleep(10)
    
    def _check_all_interviews(self):
        """Check all active interviews for completion conditions"""
        session = SessionLocal()
        try:
            # Find all active interviews
            active_interviews = session.query(Candidate).filter(
                Candidate.interview_started_at.isnot(None),
                Candidate.interview_completed_at.is_(None)
            ).all()
            
            for candidate in active_interviews:
                if self._should_complete(candidate):
                    logger.info(f"Auto-completing interview for {candidate.name}")
                    completion_handler.complete_interview(
                        candidate.interview_token,
                        "automatic_condition_met"
                    )
            
        except Exception as e:
            logger.error(f"Error checking interviews: {e}")
        finally:
            session.close()
    
    def _should_complete(self, candidate) -> bool:
        """Check if interview should be completed"""
        now = datetime.now()
        
        # Condition 1: Progress is 100%
        if candidate.interview_progress_percentage >= 100:
            return True
        
        # Condition 2: All questions answered
        if (candidate.interview_total_questions > 0 and 
            candidate.interview_answered_questions >= candidate.interview_total_questions):
            return True
        
        # Condition 3: Minimum threshold met (10 questions)
        if candidate.interview_answered_questions >= 10:
            return True
        
        # Condition 4: Interview running for over 45 minutes
        if candidate.interview_started_at:
            duration = (now - candidate.interview_started_at).total_seconds()
            if duration > 2700:  # 45 minutes
                return True
        
        # Condition 5: No activity for 15 minutes
        if candidate.interview_last_activity:
            inactive_time = (now - candidate.interview_last_activity).total_seconds()
            if inactive_time > 900 and candidate.interview_answered_questions >= 5:
                return True
        
        return False

# Create global instance
completion_monitor = AutomaticCompletionMonitor()

def process_pending_analyses():
    """Process any pending interview analyses"""
    session = SessionLocal()
    try:
        # Find all completed interviews without scores
        pending = session.query(Candidate).filter(
            Candidate.interview_completed_at.isnot(None),
            Candidate.interview_ai_score.is_(None),
            Candidate.interview_ai_analysis_status != 'processing'
        ).all()
        
        logger.info(f"Found {len(pending)} pending interview analyses")
        
        for candidate in pending:
            logger.info(f"Processing pending analysis for candidate {candidate.id}")
            trigger_auto_scoring(candidate.id)
            time.sleep(1)  # Small delay between processing
            
    except Exception as e:
        logger.error(f"Error processing pending analyses: {e}")
    finally:
        session.close()

# Add scheduled task to check every 5 minutes

def start_analysis_monitor():
    """Start monitoring for pending analyses"""

    def monitor():
        while True:
            time.sleep(300)  # Check every 5 minutes
            try:
                process_pending_analyses()
            except Exception as e:
                logger.error(f"Analysis monitor error: {e}")
    
    thread = threading.Thread(target=monitor, daemon=True)
    thread.start()

import atexit


def cleanup():
    """Cleanup resources on shutdown"""
    logger.info("Shutting down TalentFlow AI Backend...")
    executor.shutdown(wait=True)
    
atexit.register(cleanup)

# Helpers (kept verbatim)

def create_error_page(token, error):
    """Create enhanced error page with debugging info"""
    
    # Try to find if there are any scheduled interviews
    session = SessionLocal()
    try:
        recent_interviews = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).order_by(Candidate.interview_date.desc()).limit(5).all()
        
        debug_info = {
            "recent_interviews": [
                {"id": c.id, "name": c.name, "token": c.interview_token[:8] + "..."}
                for c in recent_interviews
            ] if recent_interviews else []
        }
    finally:
        session.close()
    
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview System Error</title>
  <style>
    body {{ 
      font-family: Arial, sans-serif; 
      text-align: center; 
      margin-top: 100px; 
      background: #f8f9fa; 
    }}
    .container {{ 
      max-width: 600px; 
      margin: 0 auto; 
      padding: 2rem; 
      background: white; 
      border-radius: 10px; 
      box-shadow: 0 2px 10px rgba(0,0,0,0.1); 
    }}
    .debug-info {{
      background: #f0f0f0;
      padding: 1rem;
      border-radius: 5px;
      margin-top: 1rem;
      text-align: left;
      font-size: 0.9em;
    }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;">Interview System Error</h1>
    <p>There was an error loading the interview.</p>
    <p><strong>Error:</strong> {error}</p>
    <p><strong>Token:</strong> {token}</p>
    
    <div class="debug-info">
      <strong>Debugging Steps:</strong>
      <ol>
        <li>Check if token exists: <code>GET /api/debug/check-token/{token}</code></li>
        <li>Create test interview: <code>POST /api/create-test-interview</code></li>
        <li>View recent interviews: <code>GET /api/candidates?interview_scheduled=true</code></li>
      </ol>
    </div>
    
    <p>The system has been notified. Please try again or contact support.</p>
    <button onclick="window.location.reload()" style="padding:10px 20px; background:#007bff; color:#fff; border:none; border-radius:5px; cursor:pointer; margin-top:1rem;">
      Retry
    </button>
  </div>
</body>
</html>"""

def create_expired_interview_page(token):
    """Create page for expired interviews"""
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview Expired</title>
  <style>
    body {{ font-family: Arial, sans-serif; text-align: center; margin-top: 100px; background: #f8f9fa; }}
    .container {{ max-width: 500px; margin: 0 auto; padding: 2rem; background: white; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;"> Interview Link Expired</h1>
    <p>This interview link has expired. Please contact HR for a new interview link.</p>
    <p><strong>Token:</strong> {token}</p>
  </div>
</body>
</html>""", 410



def create_error_page(token, error):
    """Create error page"""
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview Error</title>
  <style>
    body {{ font-family: Arial, sans-serif; text-align: center; margin-top: 100px; background: #f8f9fa; }}
    .container {{ max-width: 500px; margin: 0 auto; padding: 2rem; background: white; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;"> Interview System Error</h1>
    <p>There was an error loading the interview.</p>
    <p><strong>Error:</strong> {error}</p>
    <p><strong>Token:</strong> {token}</p>
    <p>The system has been notified. Please try again or contact support.</p>
    <button onclick="window.location.reload()" style="padding:10px 20px; background:#007bff; color:#fff; border:none; border-radius:5px; cursor:pointer;">
      Retry
    </button>
  </div>
</body>
</html>"""

