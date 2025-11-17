from flask import Flask, request, jsonify, redirect, render_template_string, send_from_directory
from flask_cors import CORS
from datetime import datetime, timedelta, timezone
from db import Candidate, SessionLocal
from flask_cors import cross_origin
import threading
import asyncio
import time
import traceback
import os
import json
from sqlalchemy import func, and_
import requests
import logging
from logging.handlers import RotatingFileHandler
from functools import wraps
from tenacity import retry, stop_after_attempt, wait_exponential
import sys
from flask_caching import Cache
import redis
from concurrent.futures import ThreadPoolExecutor
import uuid
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from interview_automation import start_interview_automation, stop_interview_automation
from werkzeug.utils import secure_filename
from interview_analysis_service_production import interview_analysis_service, AnalysisStatus
from flask_mail import Mail
from auth_routes import auth_bp
# from flask_caching import Cache
import atexit
# from dynamic_analysis import analyzer
if 'executor' not in globals():
    executor = ThreadPoolExecutor(max_workers=4)


# interview_analysis_service.cache = Cache(app, config=cache_config)

# Start service on app startup
print("Starting Production Interview Analysis Service...")
interview_analysis_service.start()
print("Interview Analysis Service running")
 
# Import your existing modules
try:
    from scraper import scrape_job
    from latest import create_programming_assessment
    from test_link import get_invite_link
    from clint_recruitment_system import run_recruitment_with_invite_link
    from email_util import send_assessment_email, send_assessment_reminder, send_interview_confirmation_email, send_interview_link_email, send_rejection_email
except ImportError as e:
    logging.error(f"Critical module import failed: {e}")
    raise

# Add after existing imports
try:
    from testlify_results_scraper import scrape_all_pending_assessments, scrape_assessment_results_by_name
except ImportError as e:
    logging.warning(f"Testlify scraper not available: {e}")

def _ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path

def _append_jsonl(path: str, obj: dict) -> None:
    _ensure_dir(os.path.dirname(path))
    with open(path, "a", encoding="utf-8") as fh:
        fh.write(json.dumps(obj, ensure_ascii=False) + "\n")

def _ok_preflight():
    # If you have CORS set globally, 200 is fine for OPTIONS
    return "", 200

ALLOWED_EXTS = {"webm", "mp4", "mkv", "mov"}

def _ext_from_filename(name: str, default_ext: str = "webm") -> str:
    ext = os.path.splitext(name or "")[1].lower().lstrip(".")
    return ext if ext in ALLOWED_EXTS else default_ext

# Setup proper logging
def setup_logging():
    """Configure logging for production"""
    if not os.path.exists('logs'):
        os.makedirs('logs')
    
    # Create formatter
    formatter = logging.Formatter(
        '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
    )
    
    # File handler
    file_handler = RotatingFileHandler(
        'logs/talentflow.log',
        maxBytes=10485760,  # 10MB
        backupCount=10,
        encoding='utf-8'
    )
    file_handler.setFormatter(formatter)
    file_handler.setLevel(logging.INFO)
    
    # Console handler  
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.DEBUG)
    console_handler.setFormatter(formatter)
    
    # Configure root logger
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    # Clear existing handlers
    for handler in logger.handlers[:]:
        logger.removeHandler(handler)
    
    # Add our handlers
    logger.addHandler(file_handler)
    
    # Only add console handler in development
    if os.getenv('FLASK_ENV') == 'development':
        logger.addHandler(console_handler)
    
    return logger

logger = setup_logging()

# Configuration from environment
ASSESSMENT_CONFIG = {
    'EXPIRY_HOURS': int(os.getenv('ASSESSMENT_EXPIRY_HOURS', '48')),
    'REMINDER_HOURS': int(os.getenv('ASSESSMENT_REMINDER_HOURS', '24')),
    'INTERVIEW_DELAY_DAYS': int(os.getenv('INTERVIEW_DELAY_DAYS', '3')),
    'ATS_THRESHOLD': float(os.getenv('ATS_THRESHOLD', '70')),
    'MAX_RETRIES': int(os.getenv('MAX_RETRIES', '3')),
    'RETRY_DELAY': int(os.getenv('RETRY_DELAY', '2'))
}

# Create Flask app
app = Flask(__name__)

# Setup caching
cache_config = {
    'CACHE_TYPE': 'simple',  # Use Redis in production
    'CACHE_DEFAULT_TIMEOUT': 300  # 5 minutes
}

if os.getenv('REDIS_URL'):
    cache_config = {
        'CACHE_TYPE': 'redis',
        'CACHE_REDIS_URL': os.getenv('REDIS_URL'),
        'CACHE_DEFAULT_TIMEOUT': 300
    }

cache = Cache(app, config=cache_config)

# Enhanced CORS Configuration
CORS(app, 
     origins=["http://localhost:3000", "http://127.0.0.1:3000", "https://yourfrontenddomain.com","http://127.0.0.1:3001"],
     allow_headers=["Content-Type", "Authorization", "X-Requested-With", "Accept", "Cache-Control","X-Api-Key"],
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
     supports_credentials=True,
     expose_headers=["Content-Type", "Authorization"]),

# NEXTJS_URL = os.getenv('NEXTJS_URL', 'http://localhost:3001')
HEYGEN_API_URL = "https://api.heygen.com/v1/your-avatar-endpoint"  # Change this!

HEYGEN_API_KEY = os.getenv('HEYGEN_API_KEY', '')
# Thread pool for background tasks
executor = ThreadPoolExecutor(max_workers=4)

# Pipeline status tracking
pipeline_status = {}
pipeline_lock = threading.Lock()

# Performance monitoring
request_metrics = {
    'total_requests': 0,
    'avg_response_time': 0,
    'slow_requests': 0
}


#  Admin notification function
def notify_admin(subject, message, error_details=None):
    """Send critical notifications to admin"""
    try:
        admin_email = os.getenv('ADMIN_EMAIL')
        if not admin_email:
            logger.warning("ADMIN_EMAIL not set, skipping notification")
            return
        
        from email_util import send_email
        
        body_html = f"""
        <html>
            <body>
                <h2>TalentFlow AI Alert: {subject}</h2>
                <p>{message}</p>
                {f'<pre>{error_details}</pre>' if error_details else ''}
                <p>Time: {datetime.now().isoformat()}</p>
            </body>
        </html>
        """
        
        send_email(admin_email, f"[TalentFlow Alert] {subject}", body_html)
        
    except Exception as e:
        logger.error(f"Failed to send admin notification: {e}")

# Add request timing middleware
@app.before_request
def before_request():
    request.start_time = time.time()
    request_metrics['total_requests'] += 1

@app.after_request
def after_request(response):
    if hasattr(request, 'start_time'):
        duration = time.time() - request.start_time
        
        # Update metrics
        if request_metrics['avg_response_time'] == 0:
            request_metrics['avg_response_time'] = duration
        else:
            request_metrics['avg_response_time'] = (request_metrics['avg_response_time'] + duration) / 2
        
        if duration > 5.0:  # Slow request threshold
            request_metrics['slow_requests'] += 1
            logger.warning(f"Slow request: {request.method} {request.path} took {duration:.2f}s")
        
        # Add performance headers
        response.headers['X-Response-Time'] = f"{duration:.3f}s"
        response.headers['X-Request-ID'] = getattr(request, 'request_id', 'unknown')
    
    return response

# Add request logging for debugging
@app.before_request
def log_request_info():
    """Log incoming requests for debugging"""
    request.request_id = str(uuid.uuid4())[:8]
    if request.endpoint and (request.endpoint.startswith('api_') or 'api' in request.path):
        logger.info(f"ðŸŒ [{request.request_id}] {request.method} {request.path} from {request.remote_addr}")
        if request.method == 'OPTIONS':
            logger.info(f"ðŸ”§ [{request.request_id}] CORS preflight for {request.path}")

# Rate limiting decorator with better performance
def rate_limit(max_calls=10, time_window=60):
    """Enhanced rate limiting decorator with memory optimization"""
    calls = {}
    cleanup_counter = 0
    
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            nonlocal cleanup_counter
            
            # Skip rate limiting for OPTIONS requests (CORS preflight)
            if request.method == 'OPTIONS':
                return func(*args, **kwargs)
            
            now = time.time()
            key = request.remote_addr
            
            # Periodic cleanup to prevent memory leaks
            cleanup_counter += 1
            if cleanup_counter % 100 == 0:
                cutoff = now - time_window * 2
                for ip in list(calls.keys()):
                    calls[ip] = [call_time for call_time in calls.get(ip, []) if call_time > cutoff]
                    if not calls[ip]:
                        del calls[ip]
            
            if key not in calls:
                calls[key] = []
            
            # Remove old calls
            calls[key] = [call_time for call_time in calls[key] if now - call_time < time_window]
            
            if len(calls[key]) >= max_calls:
                logger.warning(f"Rate limit exceeded for {key}")
                return jsonify({"error": "Rate limit exceeded"}), 429
            
            calls[key].append(now)
            return func(*args, **kwargs)
        return wrapper
    return decorator

# Pipeline management functions
def update_pipeline_status(job_id, status, message, progress=None):
    """Thread-safe pipeline status updates"""
    with pipeline_lock:
        pipeline_status[str(job_id)] = {
            'status': status,
            'message': message,
            'progress': progress,
            'timestamp': datetime.now().isoformat(),
            'job_id': str(job_id)
        }
    logger.info(f"Pipeline {job_id}: {status} - {message}")

def get_pipeline_status(job_id=None):
    """Get pipeline status (thread-safe)"""
    with pipeline_lock:
        if job_id:
            return pipeline_status.get(str(job_id))
        return dict(pipeline_status)

# Optimized data fetching with caching
@cache.memoize(timeout=300)
def get_cached_jobs():
    """Cached job fetching"""
    try:
        API_KEY = os.getenv("BAMBOOHR_API_KEY")
        SUBDOMAIN = os.getenv("BAMBOOHR_SUBDOMAIN")
        
        if not API_KEY or not SUBDOMAIN:
            raise ValueError("BambooHR credentials not configured")
            
        auth = (API_KEY, "x")
        headers = {"Accept": "application/json", "Content-Type": "application/json"}
        url = f"https://api.bamboohr.com/api/gateway.php/{SUBDOMAIN}/v1/applicant_tracking/jobs/"
        
        resp = requests.get(url, auth=auth, headers=headers, timeout=10)
        resp.raise_for_status()
        
        jobs = resp.json()
        open_jobs = []
        
        session = SessionLocal()
        try:
            for job in jobs:
                if job.get("status", {}).get("label", "").lower() == "open":
                    # Get candidate count for this job
                    candidate_count = session.query(Candidate).filter_by(job_id=str(job["id"])).count()
                    
                    open_jobs.append({
                        "id": job["id"],
                        "title": job.get("title", {}).get("label", ""),
                        "location": job.get("location", {}).get("label", ""),
                        "department": job.get("department", {}).get("label", ""),
                        "postingUrl": job.get("postingUrl", ""),
                        "applications": candidate_count,
                        "status": "Active",
                        "description": job.get("description", "")
                    })
        finally:
            session.close()
        
        return open_jobs
        
    except Exception as e:
        logger.error(f"BambooHR API error: {e}")
        # Fallback to database
        return get_jobs_from_database()

def get_jobs_from_database():
    """Fallback job fetching from database"""
    session = SessionLocal()
    try:
        jobs_data = session.query(
            Candidate.job_id,
            Candidate.job_title,
            func.count(Candidate.id).label('applications')
        ).filter(
            Candidate.job_id.isnot(None),
            Candidate.job_title.isnot(None)
        ).group_by(
            Candidate.job_id, 
            Candidate.job_title
        ).all()
        
        jobs = []
        for job_id, job_title, app_count in jobs_data:
            jobs.append({
                'id': str(job_id),
                'title': job_title,
                'department': 'Engineering',
                'location': 'Remote',
                'applications': app_count,
                'status': 'Active',
                'description': f'Job description for {job_title}',
                'postingUrl': ''
            })
        
        return jobs
    finally:
        session.close()

@cache.memoize(timeout=180)  # 3 minutes cache
def get_cached_candidates(job_id=None, status_filter=None):
    """Cached candidate fetching with optimized queries"""
    session = SessionLocal()
    try:
        query = session.query(Candidate)
        
        if job_id:
            query = query.filter_by(job_id=str(job_id))
        
        if status_filter:
            query = query.filter_by(status=status_filter)
        
        candidates = query.all()
        
        result = []
        for c in candidates:
            try:
                # Calculate time remaining for assessment
                time_remaining = None
                link_expired = False
                
                if c.exam_link_sent_date and not c.exam_completed:
                    deadline = c.exam_link_sent_date + timedelta(hours=ASSESSMENT_CONFIG['EXPIRY_HOURS'])
                    if datetime.now() < deadline:
                        time_remaining = (deadline - datetime.now()).total_seconds() / 3600
                    else:
                        link_expired = True
                
                candidate_data = {
                    "id": c.id,
                    "name": c.name or "Unknown",
                    "email": c.email or "",
                    "job_id": c.job_id,
                    "job_title": c.job_title or "Unknown Position",
                    "status": c.status,
                    "ats_score": float(c.ats_score) if c.ats_score else 0.0,
                    "linkedin": c.linkedin,
                    "github": c.github,
                    "phone": getattr(c, 'phone', None),
                    "resume_path": c.resume_path,
                    "resume_url": c.resume_path,  # Add this for frontend compatibility
                    "processed_date": c.processed_date.isoformat() if c.processed_date else None,
                    "score_reasoning": c.score_reasoning,
                    
                    # Assessment fields
                    "assessment_invite_link": c.assessment_invite_link,
                    "exam_link_sent": bool(c.exam_link_sent),
                    "exam_link_sent_date": c.exam_link_sent_date.isoformat() if c.exam_link_sent_date else None,
                    "exam_completed": bool(c.exam_completed),
                    "exam_completed_date": c.exam_completed_date.isoformat() if c.exam_completed_date else None,
                    "link_expired": link_expired,
                    "time_remaining_hours": time_remaining,
                    "exam_percentage": float(c.exam_percentage) if c.exam_percentage else None,
                    
                    # Interview scheduling fields
                    "interview_scheduled": bool(c.interview_scheduled),
                    "interview_date": c.interview_date.isoformat() if c.interview_date else None,
                    "interview_link": c.interview_link,
                    "interview_token": c.interview_token,
                    
                    # Interview progress fields
                    "interview_started_at": c.interview_started_at.isoformat() if c.interview_started_at else None,
                    "interview_completed_at": c.interview_completed_at.isoformat() if c.interview_completed_at else None,
                    "interview_duration": c.interview_duration or 0,
                    "interview_progress": c.interview_progress_percentage or 0,
                    "interview_questions_answered": c.interview_answered_questions or 0,
                    "interview_total_questions": c.interview_total_questions or 0,
                    
                    # Interview AI analysis fields
                    "interview_ai_score": c.interview_ai_score,
                    "interview_ai_technical_score": c.interview_ai_technical_score,
                    "interview_ai_communication_score": c.interview_ai_communication_score,
                    "interview_ai_problem_solving_score": c.interview_ai_problem_solving_score,
                    "interview_ai_cultural_fit_score": c.interview_ai_cultural_fit_score,
                    "interview_ai_overall_feedback": c.interview_ai_overall_feedback,
                    "interview_ai_analysis_status": c.interview_ai_analysis_status,
                    "interview_final_status": c.interview_final_status,
                    
                    # Interview insights
                    "strengths": json.loads(c.interview_ai_strengths or '[]') if c.interview_ai_strengths else [],
                    "weaknesses": json.loads(c.interview_ai_weaknesses or '[]') if c.interview_ai_weaknesses else [],
                    "recommendations": json.loads(c.interview_recommendations or '[]') if hasattr(c, 'interview_recommendations') and c.interview_recommendations else [],
                    
                    # Interview recording
                    "interview_recording_url": c.interview_recording_url,
                    
                    # Status fields
                    "final_status": c.final_status,
                }
                
                result.append(candidate_data)
                
            except Exception as e:
                logger.error(f"Error processing candidate {c.id}: {e}")
                continue
        
        return result
    finally:
        session.close()

@app.route('/', methods=['GET'])
def home():
    """Enhanced root endpoint with comprehensive API information"""
    try:
        # Get system statistics
        session = SessionLocal()
        try:
            stats = {
                "total_candidates": session.query(Candidate).count(),
                "total_jobs": session.query(Candidate.job_id).distinct().count(),
                "shortlisted_candidates": session.query(Candidate).filter_by(status='Shortlisted').count(),
                "completed_assessments": session.query(Candidate).filter_by(exam_completed=True).count(),
                "scheduled_interviews": session.query(Candidate).filter_by(interview_scheduled=True).count(),
            }
        except Exception:
            stats = {"error": "Could not fetch statistics"}
        finally:
            session.close()
        
        # Check system health
        system_health = {
            "database": "healthy",
            "cache": "healthy",
            "interview_automation": "running" if hasattr(app, 'interview_automation') else "stopped"
        }
        
        # API documentation with examples
        api_docs = {
            "endpoints": {
                "GET /api/jobs": {
                    "description": "Get all job postings",
                    "example": f"{request.host_url}api/jobs"
                },
                "GET /api/candidates": {
                    "description": "Get candidates (filterable by job_id, status)",
                    "example": f"{request.host_url}api/candidates?job_id=123&status=Shortlisted"
                },
                "POST /api/run_full_pipeline": {
                    "description": "Start recruitment pipeline for a job",
                    "example_payload": {
                        "job_id": "123",
                        "job_title": "Software Engineer",
                        "job_desc": "Job description here"
                    }
                },
                "GET /api/pipeline_status/<job_id>": {
                    "description": "Check pipeline status for specific job",
                    "example": f"{request.host_url}api/pipeline_status/123"
                }
            }
        }
        
        return jsonify({
            "message": "ðŸš€ TalentFlow AI Backend API",
            "tagline": "Intelligent Recruitment Automation Platform",
            "version": "2.1.0",
            "status": "operational",
            "timestamp": datetime.now().isoformat(),
            "uptime": "System started successfully",
            
            # System Statistics
            "statistics": stats,
            "system_health": system_health,
            
            # Quick Links
            "quick_links": {
                "health_check": f"{request.host_url}health",
                "api_documentation": f"{request.host_url}api/routes",
                "frontend_dashboard": "http://localhost:3000"
            },
            
            # API Information
            "api_info": api_docs,
            
            # Contact & Support
            "support": {
                "company": os.getenv('COMPANY_NAME', 'TalentFlow AI'),
                "admin_email": os.getenv('ADMIN_EMAIL', 'admin@talentflow.ai'),
                "documentation": "https://docs.talentflow.ai"
            },
            
            # Features Highlight
            "features": [
                "AI-Powered Resume Screening",
                "Automated Assessment Creation", 
                "Smart Email Automation",
                "Real-time Analytics Dashboard",
                "AI Avatar Interviews",
                "Pipeline Automation"
            ]
        }), 200
        
    except Exception as e:
        logger.error(f"Error in enhanced home route: {e}")
        return jsonify({
            "message": "TalentFlow AI Backend API",
            "version": "2.1.0",
            "status": "running",
            "error": "Partial system information available",
            "basic_endpoints": ["/api/jobs", "/api/candidates", "/health"]
        }), 200

# Enhanced API endpoints
@app.route('/api/jobs', methods=['GET', 'OPTIONS'])
@rate_limit(max_calls=30, time_window=60)
def api_jobs():
    """Enhanced API endpoint to get jobs with caching"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        jobs = get_cached_jobs()
        return jsonify(jobs), 200
    except Exception as e:
        logger.error(f"Error in api_jobs: {e}", exc_info=True)
        return jsonify({"error": "Failed to fetch jobs", "message": str(e)}), 500

@app.route('/api/candidates', methods=['GET', 'OPTIONS'])
@rate_limit(max_calls=60, time_window=60)
def api_candidates():
    """Enhanced API endpoint to get candidates with caching"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        job_id = request.args.get('job_id')
        status_filter = request.args.get('status')
        
        candidates = get_cached_candidates(job_id, status_filter)
        return jsonify(candidates), 200
        
    except Exception as e:
        logger.error(f"Error in api_candidates: {e}", exc_info=True)
        return jsonify({"error": "Failed to fetch candidates", "message": str(e)}), 500

# @app.route('/api/run_full_pipeline', methods=['POST', 'OPTIONS'])
# @rate_limit(max_calls=5, time_window=300)
# def api_run_full_pipeline():
#     """Enhanced pipeline API with status tracking"""
#     if request.method == 'OPTIONS':
#         return '', 200
    
#     try:
#         data = request.json
#         job_id = data.get('job_id')
#         job_title = data.get('job_title')
#         job_desc = data.get('job_desc', "")
        
#         logger.info(f"[{request.request_id}] Pipeline request: job_id={job_id}, job_title={job_title}")
        
#         if not job_id or not job_title:
#             return jsonify({"success": False, "message": "job_id and job_title are required"}), 400
        
#         # Check if pipeline is already running for this job
#         current_status = get_pipeline_status(job_id)
#         if current_status and current_status.get('status') == 'running':
#             return jsonify({
#                 "success": False,
#                 "message": f"Pipeline already running for {job_title}",
#                 "status": current_status
#             }), 409
        
#         # # Update status to starting
#         # update_pipeline_status(job_id, 'starting', f'Initializing pipeline for {job_title}', 0)
        
# # Update status to starting
#         update_pipeline_status(job_id, 'starting', f'Initializing pipeline for {job_title}', 0)
        
#         # Start the pipeline in background thread
#         future = executor.submit(run_pipeline_with_monitoring, job_id, job_title, job_desc)
        
#         # Store future for tracking
#         with pipeline_lock:
#             pipeline_status[str(job_id)]['future'] = future
        
#         return jsonify({
#             "success": True, 
#             "message": f"Pipeline started for {job_title}",
#             "job_id": job_id,
#             "estimated_time": "5-10 minutes",
#             "status_endpoint": f"/api/pipeline_status/{job_id}"
#         }), 200
        
#     except Exception as e:
#         logger.error(f"Error in run_full_pipeline: {e}", exc_info=True)
#         return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/pipeline_status', methods=['GET', 'OPTIONS'])
@app.route('/api/pipeline_status/<job_id>', methods=['GET', 'OPTIONS'])
def api_pipeline_status(job_id=None):
    """Get pipeline status for specific job or all jobs"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        if job_id:
            status = get_pipeline_status(job_id)
            if not status:
                return jsonify({"success": False, "message": "Pipeline not found"}), 404
            
            # Clean up the status (remove future object for JSON serialization)
            clean_status = {k: v for k, v in status.items() if k != 'future'}
            return jsonify({"success": True, "status": clean_status}), 200
        else:
            all_status = get_pipeline_status()
            # Clean up all statuses
            clean_statuses = {k: {sk: sv for sk, sv in v.items() if sk != 'future'} 
                            for k, v in all_status.items()}
            return jsonify({"success": True, "pipelines": clean_statuses}), 200
            
    except Exception as e:
        logger.error(f"Error in pipeline_status: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

# In backend.py, modify the pipeline endpoint
@app.route('/api/run_full_pipeline', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=5, time_window=300)
def api_run_full_pipeline():
    """Enhanced pipeline API with optional assessment creation"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        job_id = data.get('job_id')
        job_title = data.get('job_title')
        job_desc = data.get('job_desc', "")
        create_assessment = data.get('create_assessment', False)  # New parameter
        
        logger.info(f"Pipeline request: job_id={job_id}, create_assessment={create_assessment}")
        
        if not job_id or not job_title:
            return jsonify({"success": False, "message": "job_id and job_title are required"}), 400
        
        # Check if pipeline is already running
        current_status = get_pipeline_status(job_id)
        if current_status and current_status.get('status') == 'running':
            return jsonify({
                "success": False,
                "message": f"Pipeline already running for {job_title}",
                "status": current_status
            }), 409
        
        # Update status
        update_pipeline_status(job_id, 'starting', f'Initializing pipeline for {job_title}', 0)
        
        # Start pipeline with assessment flag
        future = executor.submit(run_pipeline_with_monitoring, job_id, job_title, job_desc, create_assessment)
        
        with pipeline_lock:
            pipeline_status[str(job_id)]['future'] = future
        
        return jsonify({
            "success": True, 
            "message": f"Pipeline started for {job_title}",
            "job_id": job_id,
            "create_assessment": create_assessment,
            "estimated_time": "3-5 minutes" if not create_assessment else "5-10 minutes"
        }), 200
        
    except Exception as e:
        logger.error(f"Error in run_full_pipeline: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

def run_pipeline_with_monitoring(job_id, job_title, job_desc, create_assessment=False):
    """Enhanced pipeline runner with optional assessment"""
    start_time = time.time()
    
    try:
        logger.info(f"Starting pipeline for job_id={job_id}, create_assessment={create_assessment}")
        update_pipeline_status(job_id, 'running', 'Pipeline started', 10)
        
        # Clear caches
        cache.delete_memoized(get_cached_candidates)
        cache.delete_memoized(get_cached_jobs)
        
        # Run modified pipeline
        full_recruitment_pipeline(job_id, job_title, job_desc, create_assessment)
        
        duration = time.time() - start_time
        update_pipeline_status(job_id, 'completed', f'Pipeline completed in {duration:.1f}s', 100)
        logger.info(f"Pipeline completed in {duration:.2f} seconds")
        
    except Exception as e:
        duration = time.time() - start_time
        error_msg = f"Pipeline failed after {duration:.2f} seconds: {str(e)}"
        update_pipeline_status(job_id, 'error', error_msg, None)
        logger.error(error_msg, exc_info=True)

def full_recruitment_pipeline(job_id, job_title, job_desc, create_assessment=False):
    """Modified pipeline with optional assessment creation"""
    try:
        logger.info(f"Starting recruitment pipeline for job_id={job_id}")
        
        # STEP 1: Scraping (40% progress if no assessment, 25% if assessment)
        progress_step1 = 40 if not create_assessment else 25
        try:
            update_pipeline_status(job_id, 'running', 'Scraping resumes...', progress_step1)
            logger.info(f"STEP 1: Scraping resumes for job_id={job_id}")
            asyncio.run(scrape_job(job_id))
            logger.info("Scraping completed successfully")
        except Exception as e:
            logger.error(f"Scraping failed: {str(e)}", exc_info=True)
        
        invite_link = None
        
        # STEP 2 & 3: Only if assessment creation is requested
        if create_assessment:
            # Create assessment (50% progress)
            try:
                update_pipeline_status(job_id, 'running', 'Creating assessment in Testlify...', 50)
                logger.info(f"STEP 2: Creating assessment for '{job_title}'")
                create_programming_assessment(job_title, job_desc)
                logger.info("Assessment created successfully")
            except Exception as e:
                logger.error(f"Assessment creation failed: {str(e)}", exc_info=True)
            
            # Get invite link (70% progress) - REMOVED as requested
            # We'll use a default or skip this step
            invite_link = f"https://candidate.testlify.com/assessment/{job_id}"
            update_pipeline_status(job_id, 'running', 'Assessment created', 70)
        
        # STEP 4: Run AI screening (100% progress)
        final_progress = 100
        try:
            update_pipeline_status(job_id, 'running', 'Running AI-powered screening...', final_progress - 10)
            logger.info("Running AI-powered screening...")
            run_recruitment_with_invite_link(
                job_id=job_id, 
                job_title=job_title, 
                job_desc=job_desc, 
                invite_link=invite_link
            )
            logger.info("AI screening completed successfully")
        except Exception as e:
            logger.error(f"AI screening failed: {str(e)}", exc_info=True)
            raise
        
        # Clear caches
        cache.delete_memoized(get_cached_candidates)
        cache.delete_memoized(get_cached_jobs)
        
        logger.info("Recruitment pipeline finished successfully")
            
    except Exception as e:
        logger.error(f"Fatal pipeline error: {e}", exc_info=True)
        raise

@app.route('/api/recruitment-stats', methods=['GET', 'OPTIONS'])
@rate_limit(max_calls=20, time_window=60)
@cache.memoize(timeout=600)  # 10 minute cache
def api_recruitment_stats():
    """Cached recruitment statistics"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        stats = []
        current_date = datetime.now()
        
        # Get last 6 months of data efficiently
        for i in range(6):
            try:
                month_date = current_date - timedelta(days=30*i)
                month_name = month_date.strftime('%b')
                
                # Calculate month boundaries
                month_start = month_date.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                if month_start.month == 12:
                    month_end = month_start.replace(year=month_start.year + 1, month=1, day=1) - timedelta(seconds=1)
                else:
                    month_end = month_start.replace(month=month_start.month + 1, day=1) - timedelta(seconds=1)
                
                # Single query for all stats
                applications = session.query(func.count(Candidate.id)).filter(
                    and_(
                        Candidate.processed_date >= month_start,
                        Candidate.processed_date <= month_end
                    )
                ).scalar() or 0
                
                interviews = session.query(func.count(Candidate.id)).filter(
                    and_(
                        Candidate.interview_scheduled == True,
                        Candidate.interview_date >= month_start,
                        Candidate.interview_date <= month_end
                    )
                ).scalar() or 0
                
                hires = session.query(func.count(Candidate.id)).filter(
                    and_(
                        Candidate.final_status == "Hired",
                        Candidate.processed_date >= month_start,
                        Candidate.processed_date <= month_end
                    )
                ).scalar() or 0
                
                stats.append({
                    "month": month_name,
                    "applications": applications,
                    "interviews": interviews,
                    "hires": hires
                })
                
            except Exception as e:
                logger.error(f"Error calculating stats for month {i}: {e}")
                stats.append({
                    "month": (current_date - timedelta(days=30*i)).strftime('%b'),
                    "applications": 0,
                    "interviews": 0,
                    "hires": 0
                })
        
        # Reverse to get chronological order
        stats.reverse()
        
        logger.info(f"Generated recruitment stats for {len(stats)} months")
        return jsonify(stats), 200
        
    except Exception as e:
        logger.error(f"Error in api_recruitment_stats: {e}", exc_info=True)
        return jsonify({"error": "Failed to get statistics", "message": str(e)}), 500
    finally:
        session.close()

@app.route('/api/send_reminder/<int:candidate_id>', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=10, time_window=60)
def api_send_reminder(candidate_id):
    """Send reminder to specific candidate"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return jsonify({"success": False, "message": "Candidate not found"}), 404
            
            # Check if candidate is eligible for reminder
            if not candidate.exam_link_sent or candidate.exam_completed:
                return jsonify({"success": False, "message": "Candidate not eligible for reminder"}), 400
            
            # Calculate hours remaining
            hours_remaining = 24  # Default
            if candidate.exam_link_sent_date:
                deadline = candidate.exam_link_sent_date + timedelta(hours=48)
                hours_remaining = max(0, int((deadline - datetime.now()).total_seconds() / 3600))
            
            # Send reminder email
            send_assessment_reminder(candidate, hours_remaining)
            
            # Update reminder tracking
            candidate.reminder_sent = True
            candidate.reminder_sent_date = datetime.now()
            session.commit()
            
            # Clear cache
            cache.delete_memoized(get_cached_candidates)
            
            return jsonify({
                "success": True,
                "message": f"Reminder sent to {candidate.name}"
            }), 200
            
        finally:
            session.close()
        
    except Exception as e:
        logger.error(f"Error in send_reminder: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/assessments', methods=['GET'])
def api_assessments():
    return jsonify([]), 200

@app.route('/secure-interview/<token>', methods=['GET'])
def secure_interview_page(token):
    """Serve the interview page ONLY if interview is not completed"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return create_error_page(token, "Interview not found"), 404
        
        # BLOCK if completed
        if getattr(candidate, 'interview_completed_at', None):
            completed_time = candidate.interview_completed_at.strftime('%B %d, %Y at %I:%M %p')
            return f"""<!DOCTYPE html>
            <html>
            <head>
                <title>Interview Completed</title>
                <style>
                    body {{ 
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        min-height: 100vh;
                        display: flex;
                        align-items: center;
                        justify-content: center;
                        margin: 0;
                    }}
                    .container {{
                        background: white;
                        padding: 3rem;
                        border-radius: 15px;
                        box-shadow: 0 20px 40px rgba(0,0,0,0.1);
                        text-align: center;
                        max-width: 500px;
                    }}
                    h1 {{ color: #28a745; }}
                    .icon {{ font-size: 4rem; margin-bottom: 1rem; }}
                    .details {{ 
                        background: #f8f9fa;
                        padding: 1rem;
                        border-radius: 8px;
                        margin: 1rem 0;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="icon"></div>
                    <h1>Interview Already Completed</h1>
                    <div class="details">
                        <p><strong>Candidate:</strong> {candidate.name}</p>
                        <p><strong>Position:</strong> {candidate.job_title}</p>
                        <p><strong>Completed on:</strong> {completed_time}</p>
                    </div>
                    <p>Thank you for completing your interview. Your responses have been submitted for review.</p>
                    <p>The interview link cannot be used again for security reasons.</p>
                    <p style="color: #666; font-size: 0.9em; margin-top: 2rem;">
                        If you have any questions, please contact our HR department.
                    </p>
                </div>
            </body>
            </html>""", 200
        
        # BLOCK if expired
        expires_at = getattr(candidate, 'interview_expires_at', None)
        if expires_at and expires_at < datetime.now():
            return create_expired_interview_page(token), 403
        
        # NOW ADD THE NORMAL INTERVIEW PAGE CODE (from your first code)
        # This only runs if NOT completed and NOT expired
        
        is_reconnection = bool(getattr(candidate, 'interview_started_at', None))
        
        company_name = os.getenv('COMPANY_NAME', 'Our Company')
        if getattr(candidate, 'company_name', None):
            company_name = candidate.company_name
        
        job_description = getattr(candidate, 'job_description', f'Interview for {candidate.job_title} position')
        
        interview_data = {
            'token': token,
            'candidateId': candidate.id,
            'candidateName': candidate.name,
            'candidateEmail': candidate.email,
            'position': candidate.job_title,
            'company': company_name,
            # 'knowledgeBaseId': getattr(candidate, 'knowledge_base_id', None),
            'sessionId': getattr(candidate, 'interview_session_id', None),
            'status': 'active',
            'jobDescription': job_description,
            'atsScore': candidate.ats_score,
            'resumePath': candidate.resume_path,
            'isReconnection': is_reconnection,
            'previousSessionData': {
                'questionsAsked': getattr(candidate, 'interview_total_questions', 0),
                'questionsAnswered': getattr(candidate, 'interview_answered_questions', 0),
                'duration': getattr(candidate, 'interview_duration', 0)
            }
        }
        
        return create_interview_landing_page(interview_data, token), 200
        
    except Exception as e:
        logger.exception("Error in interview route")
        return create_error_page(token, str(e)), 500
    finally:
        session.close()

# Add this debug endpoint to your backend.py
@app.route('/api/debug/check-token/<token>', methods=['GET'])
def debug_check_token(token):
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if candidate:
            return jsonify({
                "found": True,
                "candidate_id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "interview_scheduled": candidate.interview_scheduled,
                "interview_token": candidate.interview_token
            }), 200
        else:
            # Try to find similar tokens
            similar = session.query(Candidate).filter(
                Candidate.interview_token.like(f'%{token[:8]}%')
            ).all()
            
            return jsonify({
                "found": False,
                "token_searched": token,
                "similar_tokens": [
                    {"id": c.id, "name": c.name, "token": c.interview_token}
                    for c in similar
                ]
            }), 404
    finally:
        session.close()

def create_error_page(token, error):
    """Create enhanced error page with debugging info"""
    
    # Try to find if there are any scheduled interviews
    session = SessionLocal()
    try:
        recent_interviews = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).order_by(Candidate.interview_date.desc()).limit(5).all()
        
        debug_info = {
            "recent_interviews": [
                {"id": c.id, "name": c.name, "token": c.interview_token[:8] + "..."}
                for c in recent_interviews
            ] if recent_interviews else []
        }
    finally:
        session.close()
    
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview System Error</title>
  <style>
    body {{ 
      font-family: Arial, sans-serif; 
      text-align: center; 
      margin-top: 100px; 
      background: #f8f9fa; 
    }}
    .container {{ 
      max-width: 600px; 
      margin: 0 auto; 
      padding: 2rem; 
      background: white; 
      border-radius: 10px; 
      box-shadow: 0 2px 10px rgba(0,0,0,0.1); 
    }}
    .debug-info {{
      background: #f0f0f0;
      padding: 1rem;
      border-radius: 5px;
      margin-top: 1rem;
      text-align: left;
      font-size: 0.9em;
    }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;">🚫 Interview System Error</h1>
    <p>There was an error loading the interview.</p>
    <p><strong>Error:</strong> {error}</p>
    <p><strong>Token:</strong> {token}</p>
    
    <div class="debug-info">
      <strong>Debugging Steps:</strong>
      <ol>
        <li>Check if token exists: <code>GET /api/debug/check-token/{token}</code></li>
        <li>Create test interview: <code>POST /api/create-test-interview</code></li>
        <li>View recent interviews: <code>GET /api/candidates?interview_scheduled=true</code></li>
      </ol>
    </div>
    
    <p>The system has been notified. Please try again or contact support.</p>
    <button onclick="window.location.reload()" style="padding:10px 20px; background:#007bff; color:#fff; border:none; border-radius:5px; cursor:pointer; margin-top:1rem;">
      Retry
    </button>
  </div>
</body>
</html>"""

# 2. Add a new endpoint to validate and refresh interview tokens
@app.route('/api/interview/validate-token/<token>', methods=['GET', 'POST'])
def validate_interview_token(token):
    """Validate interview token - PREVENT access after completion"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"valid": False, "error": "Invalid token"}), 404
        
        # CHECK 1: If interview is already completed, DENY access
        if getattr(candidate, 'interview_completed_at', None):
            return jsonify({
                "valid": False,
                "error": "Interview already completed",
                "completed_at": candidate.interview_completed_at.isoformat(),
                "message": "This interview has been completed and cannot be accessed again"
            }), 403
        
        # CHECK 2: Check expiration (with proper getattr)
        expires_at = getattr(candidate, 'interview_expires_at', None)
        if expires_at and expires_at < datetime.now():
            return jsonify({
                "valid": False,
                "error": "Interview link has expired",
                "expired_at": expires_at.isoformat(),
                "message": "This interview link has expired. Please contact HR for assistance."
            }), 403
        
        # CHECK 3: If already started but abandoned for too long (optional)
        started_at = getattr(candidate, 'interview_started_at', None)
        if started_at:
            time_since_start = datetime.now() - started_at
            if time_since_start.total_seconds() > 7200:  # 2 hours
                # Auto-complete abandoned interviews
                candidate.interview_completed_at = datetime.now()
                candidate.interview_ai_analysis_status = 'abandoned'
                session.commit()
                
                return jsonify({
                    "valid": False,
                    "error": "Interview session timeout",
                    "message": "This interview session has timed out after 2 hours."
                }), 403
        
        # Only allow access if NOT completed and NOT expired
        session_info = {
            "valid": True,
            "candidate_name": candidate.name,
            "position": candidate.job_title,
            "interview_started": bool(started_at),
            "interview_completed": False,  # Always false here since we checked above
            "knowledge_base_id": getattr(candidate, 'knowledge_base_id', None),
            "can_continue": True,
            "questions_asked": getattr(candidate, 'interview_total_questions', 0),
            "questions_answered": getattr(candidate, 'interview_answered_questions', 0)
        }
        
        # Update last accessed time
        if request.method == 'POST':
            if hasattr(candidate, 'last_accessed'):
                candidate.last_accessed = datetime.now()
                session.commit()
        
        return jsonify(session_info), 200
        
    except Exception as e:
        logger.error(f"Error validating token: {e}")
        return jsonify({"valid": False, "error": str(e)}), 500
    finally:
        session.close()


# 3. Update the create_interview_landing_page to show reconnection status
def create_interview_landing_page(interview_data, token):
    """Create the landing page HTML with reconnection support without illegal f-strings."""
    import json
    interview_json = json.dumps(interview_data)

    # Pre-build dynamic fragments (NO backslashes in f-string expressions):
    is_reconnection = bool(interview_data.get('isReconnection'))
    prev = interview_data.get('previousSessionData') or {}
    reconnect_block = ""
    if is_reconnection:
        reconnect_block = (
            "<div class=\"reconnect-info\">"
            "<h3> Welcome Back!</h3>"
            "<p>You're reconnecting to your interview session.</p>"
            f"<p><strong>Questions asked:</strong> {prev.get('questionsAsked', 0)}</p>"
            f"<p><strong>Questions answered:</strong> {prev.get('questionsAnswered', 0)}</p>"
            "</div>"
        )
    session_type_label = "Reconnection" if is_reconnection else "New Session"
    continue_or_start = "Continue" if is_reconnection else "Start"
    continue_lower = "continue" if is_reconnection else "start"
    progress_line = (
        f"<p><strong> Progress:</strong> {prev.get('questionsAnswered', 0)} questions completed</p>"
        if is_reconnection else ""
    )

    # Build the HTML
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    candidate_name = interview_data['candidateName']
    position = interview_data['position']
    company = interview_data['company']
    knowledge_base_id = interview_data.get('knowledgeBaseId', '')
    candidate_id = interview_data['candidateId']

    return f"""<!DOCTYPE html>
<html>
<head>
  <title>AI Interview - {position}</title>
  <meta charset="UTF-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1.0"/>
  <style>
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
      margin: 0; padding: 0;
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      min-height: 100vh; display: flex; align-items: center; justify-content: center;
    }}
    .container {{
      background: white; padding: 2rem; border-radius: 15px;
      box-shadow: 0 20px 40px rgba(0,0,0,0.1);
      text-align: center; max-width: 600px; width: 90%;
    }}
    .header {{ color: #333; margin-bottom: 1.5rem; }}
    .info-box {{
      background: #f8f9fa; padding: 1.5rem; border-radius: 10px;
      margin: 1rem 0; border-left: 5px solid #667eea;
    }}
    .reconnect-info {{
      background: #e3f2fd; padding: 1rem; border-radius: 8px;
      margin: 1rem 0; border-left: 5px solid #2196f3;
    }}
    .success-badge {{ color: #28a745; font-weight: bold; font-size: 1.1em; }}
    .start-btn {{
      background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
      color: white; padding: 15px 30px; border: none; border-radius: 50px;
      cursor: pointer; font-size: 18px; font-weight: bold; text-decoration: none;
      display: inline-block; margin: 20px 10px; transition: transform 0.3s, box-shadow 0.3s;
    }}
    .start-btn:hover {{ transform: translateY(-2px); box-shadow: 0 10px 25px rgba(102,126,234,0.3); }}
    .instructions {{ text-align: left; margin: 1.5rem 0; padding: 1.5rem; background: #e3f2fd; border-radius: 10px; }}
    .debug-info {{ background: #fff3cd; padding: 1rem; border-radius: 8px; font-size: 14px; margin-top: 15px; text-align: left; }}
  </style>
</head>
<body>
  <div class="container">
    <h1 class="header">ðŸ¤– AI Interview Portal</h1>

    {reconnect_block}

    <div class="info-box">
      <p><strong> Candidate:</strong> {candidate_name}</p>
      <p><strong> Position:</strong> {position}</p>
      <p><strong> Company:</strong> {company}</p>
      <p class="success-badge"> Interview Link Active & Ready</p>
      <p><strong> Session Type:</strong> {session_type_label}</p>
    </div>

    <div class="instructions">
      <h3> Before {continue_or_start} Your Interview:</h3>
      <ul>
        <li><strong>Internet:</strong> Ensure stable connection</li>
        <li><strong>Camera & Mic:</strong> Test and allow permissions</li>
        <li><strong>Environment:</strong> Find a quiet, well-lit space</li>
        <li><strong>Materials:</strong> Have your resume ready</li>
      </ul>
    </div>

    <div style="margin: 25px 0;">
      <p><strong>â± Duration:</strong> 30-45 minutes</p>
      <p><strong> Format:</strong> AI-powered video interview</p>
      {progress_line}
    </div>

    <button onclick="startInterview()" class="start-btn"> {continue_or_start} AI Interview</button>

    <div class="debug-info">
      <strong> System Status:</strong><br/>
      Token: {token}<br/>
      Knowledge Base: {knowledge_base_id}<br/>
      Candidate ID: {candidate_id}<br/>
      Status: Ready <br/>
      Session Type: {session_type_label}<br/>
      Time: {now_str}
    </div>

    <div style="margin-top: 30px; font-size: 12px; color: #666;">
      <p> Need help? Contact our support team</p>
      <p> Interview link valid for multiple sessions</p>
    </div>
  </div>

  <script>
    const interviewData = {interview_json};
    function startInterview() {{
      console.log(' Starting interview with data:', interviewData);
      sessionStorage.setItem('interviewData', JSON.stringify(interviewData));
      fetch('/api/avatar/interview/{token}', {{
        method: 'POST',
        headers: {{ 'Content-Type': 'application/json' }},
        body: JSON.stringify({{ action: '{continue_lower}' }})
      }}).catch(() => {{ /* best effort */ }});
      window.location.href = 'http://localhost:3001/interview/{token}';
    }}

    fetch('/api/interview/validate-token/{token}', {{ method: 'POST' }})
      .then(r => r.json())
      .then(d => {{
        if (!d.valid) {{
          alert('This interview link has expired. Please contact HR for assistance.');
          document.querySelector('.start-btn').disabled = true;
        }}
      }})
      .catch(() => {{ /* noop */ }});
  </script>
</body>
</html>"""


def create_expired_interview_page(token):
    """Create page for expired interviews"""
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview Expired</title>
  <style>
    body {{ font-family: Arial, sans-serif; text-align: center; margin-top: 100px; background: #f8f9fa; }}
    .container {{ max-width: 500px; margin: 0 auto; padding: 2rem; background: white; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;"> Interview Link Expired</h1>
    <p>This interview link has expired. Please contact HR for a new interview link.</p>
    <p><strong>Token:</strong> {token}</p>
  </div>
</body>
</html>""", 410


def create_error_page(token, error):
    """Create error page"""
    return f"""<!DOCTYPE html>
<html>
<head>
  <title>Interview Error</title>
  <style>
    body {{ font-family: Arial, sans-serif; text-align: center; margin-top: 100px; background: #f8f9fa; }}
    .container {{ max-width: 500px; margin: 0 auto; padding: 2rem; background: white; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
  </style>
</head>
<body>
  <div class="container">
    <h1 style="color: #e74c3c;"> Interview System Error</h1>
    <p>There was an error loading the interview.</p>
    <p><strong>Error:</strong> {error}</p>
    <p><strong>Token:</strong> {token}</p>
    <p>The system has been notified. Please try again or contact support.</p>
    <button onclick="window.location.reload()" style="padding:10px 20px; background:#007bff; color:#fff; border:none; border-radius:5px; cursor:pointer;">
      Retry
    </button>
  </div>
</body>
</html>"""

@app.route('/api/scrape_assessment_results', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=3, time_window=300)  # Max 3 scraping requests per 5 minutes
def api_scrape_assessment_results():
    """API endpoint to scrape assessment results for a specific assessment"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        data = request.json
        assessment_name = data.get('assessment_name')
        
        if not assessment_name:
            return jsonify({"success": False, "message": "assessment_name is required"}), 400
        
        logger.info(f"Starting results scraping for assessment: {assessment_name}")
        
        # Start scraping in a separate thread
        scraping_thread = threading.Thread(
            target=lambda: run_scraping_with_monitoring(assessment_name),
            daemon=True,
            name=f"scraping_{assessment_name.replace(' ', '_')}_{int(time.time())}"
        )
        scraping_thread.start()
        
        return jsonify({
            "success": True,
            "message": f"Started scraping results for '{assessment_name}'",
            "estimated_time": "2-5 minutes"
        }), 200
        
    except Exception as e:
        logger.error(f"Error in scrape_assessment_results: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/scrape_all_pending_results', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=1, time_window=600)  # Max 1 bulk scraping per 10 minutes
def api_scrape_all_pending_results():
    """API endpoint to scrape all pending assessment results"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        logger.info("Starting bulk results scraping for all pending assessments")
        
        # Start bulk scraping in a separate thread
        scraping_thread = threading.Thread(
            target=lambda: run_bulk_scraping_with_monitoring(),
            daemon=True,
            name=f"bulk_scraping_{int(time.time())}"
        )
        scraping_thread.start()
        
        return jsonify({
            "success": True,
            "message": "Started bulk scraping for all pending assessments",
            "estimated_time": "5-15 minutes"
        }), 200
        
    except Exception as e:
        logger.error(f"Error in scrape_all_pending_results: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500

@app.route('/api/scrape_assessment_results', methods=['GET','OPTIONS'])
def api_scraping_status():
    """Get status of running scraping operations"""
    try:
        # Get active scraping threads
        active_threads = []
        for thread in threading.enumerate():
            if thread.name.startswith(('scraping_', 'bulk_scraping_')):
                thread_info = {
                    "name": thread.name,
                    "is_alive": thread.is_alive(),
                    "daemon": thread.daemon
                }
                active_threads.append(thread_info)
        
        return jsonify({
            "success": True,
            "active_operations": len(active_threads),
            "operations": active_threads
        }), 200
        
    except Exception as e:
        logger.error(f"Error in scraping_status: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500
def run_scraping_with_monitoring(assessment_name: str):
    """Wrapper to run scraping with monitoring and error handling"""
    start_time = time.time()
    
    try:
        logger.info(f"Starting monitored scraping for assessment: {assessment_name}")
        
        # Import and run the scraping function
        try:
            from testlify_results_scraper import scrape_assessment_results_by_name
        except ImportError as e:
            logger.error(f"Failed to import scraper: {e}")
            notify_admin(
                "Scraper Import Error",
                f"Could not import results scraper: {str(e)}. Please ensure testlify_results_scraper.py is available."
            )
            return
        
        # Run the async scraping function
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            results = loop.run_until_complete(scrape_assessment_results_by_name(assessment_name))
        finally:
            loop.close()
        
        duration = time.time() - start_time
        logger.info(f"Scraping completed successfully in {duration:.2f} seconds. Found {len(results)} candidates.")
        
        # Send success notification
        notify_admin(
            "Assessment Results Scraping Completed",
            f"Assessment: {assessment_name}\nCandidates processed: {len(results)}\nDuration: {duration:.2f} seconds"
        )
        
    except Exception as e:
        duration = time.time() - start_time
        error_msg = f"Scraping failed for assessment '{assessment_name}' after {duration:.2f} seconds"
        logger.error(error_msg, exc_info=True)
        
        # Send failure notification
        notify_admin(
            "Assessment Results Scraping Failed",
            error_msg,
            error_details=traceback.format_exc()
        )

@app.route('/api/get-interview/<token>', methods=['GET', 'POST'])
def get_interview(token):
    from datetime import timezone  # ensure timezone available here
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        if not candidate:
            return jsonify({'error': 'Interview not found'}), 404

        if request.method == 'POST':
            body = request.get_json(silent=True) or {}
            action = body.get('action')

            if action == 'start':
                candidate.interview_started_at = datetime.now(timezone.utc)
                session.commit()
                return jsonify({"success": True}), 200

            elif action == 'complete':
                candidate.interview_completed_at = datetime.now(timezone.utc)
                transcript = body.get('transcript')
                if transcript:
                    candidate.interview_transcript = transcript
                session.commit()
                return jsonify({"success": True}), 200

        # unify KB id across both possible columns
        kb = getattr(candidate, 'knowledge_base_id', None) or getattr(candidate, 'interview_kb_id', None)
        position = getattr(candidate, "job_title", None) or getattr(candidate, "position", "") or "Interview"

        data = {
            "id": getattr(candidate, "id", None),
            "token": getattr(candidate, "interview_token", None),
            "candidateId": getattr(candidate, "id", None),
            "candidateName": getattr(candidate, "name", "") or "",
            "candidateEmail": getattr(candidate, "email", "") or "",
            "position": position,
            "company": getattr(candidate, "company_name", None)
                       or getattr(candidate, "company", None)
                       or os.getenv("COMPANY_NAME", "Our Company"),
            "knowledgeBaseId": kb,
            "status": "active" if getattr(candidate, "interview_scheduled", False) else "inactive",
            "jobDescription": getattr(candidate, "job_description", None)
                              or f"Interview for {position} position",
            "resumeLink": getattr(candidate, "resume_link", None) or getattr(candidate, "resume_path", None),
            "createdAt": getattr(candidate, "interview_created_at", None).isoformat()
                         if getattr(candidate, "interview_created_at", None) else None,
        }
        return jsonify(data), 200

    except Exception as e:
        logger.error(f"Error in get_interview: {e}", exc_info=True)
        return jsonify({'error': str(e)}), 500
    finally:
        session.close()


# --- TEMP DEBUG: check KB columns by interview token ---
@app.route('/api/debug/kb/<token>', methods=['GET'])
def debug_kb(token):
    s = SessionLocal()
    try:
        c = s.query(Candidate).filter_by(interview_token=token).first()
        if not c:
            return jsonify({"error": "not found", "token": token}), 404
        return jsonify({
            "id": c.id,
            "name": c.name,
            "email": c.email,
            "interview_token": c.interview_token,
            "knowledge_base_id": getattr(c, "knowledge_base_id", None),
            "interview_kb_id": getattr(c, "interview_kb_id", None),
        }), 200
    except Exception as e:
        logger.exception("debug_kb error")
        return jsonify({"error": str(e)}), 500
    finally:
        s.close()


@app.route('/api/verify-knowledge-base/<candidate_id>', methods=['GET'])
def verify_knowledge_base(candidate_id):
    """Verify knowledge base content for debugging"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Extract resume content
        resume_content = ""
        if candidate.resume_path and os.path.exists(candidate.resume_path):
            resume_content = extract_resume_content(candidate.resume_path)
        
        # Build preview WITHOUT backslashes in f-string expressions
        questions = generate_interview_questions(
            candidate_name=candidate.name,
            position=candidate.job_title,
            resume_content=resume_content,
            job_description=getattr(candidate, 'job_description', f"Position: {candidate.job_title}")
        )
        questions_preview = (questions[:1000] + "...") if len(questions) > 1000 else questions
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "position": candidate.job_title,
            "knowledge_base_id": getattr(candidate, "knowledge_base_id", None),
            "resume_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "resume_content_length": len(resume_content),
            "interview_scheduled": candidate.interview_scheduled,
            "generated_questions_preview": questions_preview,
            "total_questions_length": len(questions)
        }), 200
    finally:
        session.close()


def extract_skills_from_resume(resume_content):
    """Extract technical skills from resume"""
    skills = []
    resume_lower = resume_content.lower()
    tech_skills = [
        'python', 'javascript', 'java', 'c++', 'c#', 'react', 'angular', 'vue',
        'node.js', 'django', 'flask', 'spring', 'sql', 'nosql', 'mongodb',
        'postgresql', 'mysql', 'aws', 'azure', 'gcp', 'docker', 'kubernetes',
        'git', 'ci/cd', 'machine learning', 'data science', 'api', 'rest',
        'graphql', 'typescript', 'golang', 'rust', 'swift', 'kotlin'
    ]
    for skill in tech_skills:
        if skill in resume_lower:
            skills.append(skill.title())
    return skills[:10]


def extract_projects_from_resume(resume_content):
    """Extract project names from resume"""
    projects = []
    lines = resume_content.split('\n')
    for i, line in enumerate(lines):
        if 'project' in line.lower():
            if ':' in line:
                project_name = line.split(':', 1)[1].strip()[:50]
                if project_name:
                    projects.append(project_name)
    return projects[:5]


def extract_experience_years(resume_content):
    """Extract years of experience from resume"""
    import re
    pattern = r'(\d+)\+?\s*years?\s*(?:of\s*)?experience'
    match = re.search(pattern, resume_content.lower())
    if match:
        return f"{match.group(1)}+ years"
    year_pattern = r'20\d{2}'
    years = re.findall(year_pattern, resume_content)
    if len(years) >= 2:
        min_year = min(int(y) for y in years)
        max_year = max(int(y) for y in years)
        experience = max_year - min_year
        if experience > 0:
            return f"{experience}+ years"
    return "Not specified"


@app.route('/api/avatar/get-access-token', methods=['POST', 'OPTIONS'])
def get_avatar_access_token():
    """Generate HeyGen access token for avatar session"""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        heygen_key = os.getenv('HEYGEN_API_KEY', 'your_heygen_api_key_here')
        return heygen_key, 200, {'Content-Type': 'text/plain'}
    except Exception as e:
        logger.error(f"Error getting access token: {e}")
        return jsonify({"error": "Failed to get access token"}), 500


@app.route('/api/debug-schedule-interview', methods=['POST'])
def debug_schedule_interview():
    """Debug: try to create a KB and echo responses for troubleshooting"""
    try:
        data = request.json or {}
        api_key = os.getenv('HEYGEN_API_KEY')
        if not api_key:
            return jsonify({
                "error": "HEYGEN_API_KEY not found in environment variables",
                "fix": "Add HEYGEN_API_KEY to your .env file"
            }), 400
        if len(api_key) < 20:
            return jsonify({"error": "HEYGEN_API_KEY seems too short", "length": len(api_key)}), 400

        test_payload = {
            'name': f'Test_Interview_{int(time.time())}',
            'description': 'Test knowledge base',
            'content': 'Test interview questions: 1. Tell me about yourself. 2. Why this role?',
            'opening_line': 'Hello, this is a test interview.'
        }
        try:
            heygen_response = requests.post(
                'https://api.heygen.com/v1/streaming/knowledge_base',
                headers={
                    'X-Api-Key': api_key,
                    'Content-Type': 'application/json',
                    'Accept': 'application/json'
                },
                json=test_payload,
                timeout=30
            )
            if heygen_response.ok:
                kb_data = heygen_response.json()
                return jsonify({
                    "success": True,
                    "knowledge_base_id": (kb_data.get('data', {}) or {}).get('knowledge_base_id') or
                                         (kb_data.get('data', {}) or {}).get('id') or
                                         kb_data.get('knowledge_base_id') or kb_data.get('id'),
                    "full_response": kb_data
                }), 200
            else:
                return jsonify({
                    "success": False,
                    "error": "HeyGen API error",
                    "status_code": heygen_response.status_code,
                    "body": heygen_response.text[:500]
                }), 400
        except requests.exceptions.RequestException as e:
            return jsonify({"success": False, "error": "Request failed", "details": str(e)}), 500
    except Exception as e:
        return jsonify({
            "success": False,
            "error": "Unexpected error",
            "details": str(e),
            "traceback": traceback.format_exc()
        }), 500


# Add this enhanced function to your backend.py

def create_structured_interview_kb(candidate_name, position, company, resume_content, job_description):
    """Create a highly structured knowledge base for professional interviews"""
    
    # Extract key information
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    experience = extract_experience_years(resume_content) if resume_content else "Not specified"
    projects = extract_projects_from_resume(resume_content) if resume_content else []
    
    # Build the structured interview content
    structured_content = f"""
STRICT INTERVIEW PROTOCOL - FOLLOW EXACTLY:

YOU ARE: A professional AI interviewer conducting a structured technical interview.
YOUR BEHAVIOR: Professional, clear, structured. NO casual conversation.

CANDIDATE DETAILS:
- Name: {candidate_name}
- Position: {position}
- Company: {company}
- Experience: {experience}
- Key Skills: {', '.join(skills[:5]) if skills else 'General skills'}

INTERVIEW STRUCTURE - ASK THESE QUESTIONS IN EXACT ORDER:

=== QUESTION 1: SELF INTRODUCTION (ALWAYS ASK FIRST) ===
"Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background and what led you to apply for this role?"

=== QUESTION 2: TECHNICAL EXPERIENCE ===
{"I see from your resume that you have experience with " + skills[0] + ". Can you tell me about a specific project where you used " + skills[0] + " and what challenges you faced?" if skills else "Can you tell me about your most significant technical project and the technologies you used?"}

=== QUESTION 3: PROBLEM SOLVING ===
"That's interesting. Now, let me ask you about problem-solving. Can you describe a time when you encountered a complex technical problem and walk me through how you approached and solved it?"

=== QUESTION 4: TEAMWORK ===
"Great. Let's talk about teamwork. Tell me about a time when you had to collaborate with other team members on a challenging project. How did you handle any conflicts or disagreements?"

=== QUESTION 5: SPECIFIC SKILL DEEP DIVE ===
{"I noticed you also have experience with " + skills[1] + ". What's the most complex thing you've built using " + skills[1] + "? Please be specific about the technical details." if len(skills) > 1 else "What would you say is your strongest technical skill, and can you give me a detailed example of how you've applied it?"}

=== QUESTION 6: LEARNING ABILITY ===
"Technology evolves rapidly. Can you tell me about a time when you had to quickly learn a new technology or framework for a project? How did you approach the learning process?"

=== QUESTION 7: ROLE-SPECIFIC ===
"Let's talk specifically about this {position} role. Based on your understanding of the position, how do you see your skills and experience contributing to our team in the first 90 days?"

=== QUESTION 8: CHALLENGES ===
"What do you think would be the biggest challenge for you in this role, and how would you address it?"

=== QUESTION 9: CAREER GOALS ===
"Where do you see your career heading in the next 3-5 years, and how does this {position} role fit into those plans?"

=== QUESTION 10: CANDIDATE QUESTIONS ===
"Thank you for your answers. Now, do you have any questions for me about the role, the team, or {company}?"

CRITICAL RULES:
1. When you receive "INIT_INTERVIEW", IMMEDIATELY ask Question 1
2. Ask ONE question at a time
3. Wait for complete answer before next question
4. After each answer say: "Thank you for sharing that. [Next question]"
5. Stay professional - NO casual chat
6. If candidate says just "Hello", respond with Question 1
7. Track which question you're on to avoid repetition

FORBIDDEN BEHAVIORS:
- Do NOT say: "hey", "cool", "chat", "Oh", or use casual language
- Do NOT have conversations outside these 10 questions
- Do NOT ask random questions
- Do NOT give short one-word responses

RESUME CONTEXT:
{resume_content[:2000] if resume_content else 'No resume provided'}

JOB REQUIREMENTS:
{job_description[:1000] if job_description else f'Standard requirements for {position}'}

Remember: You are conducting a PROFESSIONAL STRUCTURED INTERVIEW. Stay on script!
"""
    
    return structured_content


@app.route('/api/create-knowledge-base', methods=['POST', 'OPTIONS'])
def create_knowledge_base_enhanced():
    """Create a structured interview knowledge base"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json or {}
        candidate_name = data.get('candidateName')
        position = data.get('position')
        company = data.get('company', 'Our Company')
        token = data.get('token')
        
        if not candidate_name or not position:
            return jsonify({"error": "candidateName and position are required"}), 400
        
        logger.info(f" Creating structured KB for: {candidate_name} - {position}")
        
        # Get resume content
        resume_content = ""
        candidate_id = None
        job_description = data.get('jobDescription', '')
        
        if token:
            session = SessionLocal()
            try:
                cand = session.query(Candidate).filter_by(interview_token=token).first()
                if cand:
                    candidate_id = cand.id
                    if cand.resume_path and os.path.exists(cand.resume_path):
                        resume_content = extract_resume_content(cand.resume_path)
                    job_description = job_description or getattr(cand, 'job_description', '')
            finally:
                session.close()
        
        # Create structured interview content
        kb_content = create_structured_interview_kb(
            candidate_name=candidate_name,
            position=position,
            company=company,
            resume_content=resume_content,
            job_description=job_description
        )
        
        # Create opening line that immediately asks first question
        opening_line = f"Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background and what led you to apply for this role?"
        
        heygen_key = os.getenv('HEYGEN_API_KEY')
        if not heygen_key:
            # Fallback
            fallback_kb_id = f"kb_structured_{candidate_name.replace(' ', '_')}_{int(time.time())}"
            logger.warning(f"No HeyGen API key, using fallback: {fallback_kb_id}")
            
            if candidate_id:
                session = SessionLocal()
                try:
                    cand = session.query(Candidate).filter_by(id=candidate_id).first()
                    if cand:
                        cand.knowledge_base_id = fallback_kb_id
                        cand.interview_kb_content = kb_content
                        session.commit()
                finally:
                    session.close()
            
            return jsonify({
                "success": True,
                "knowledgeBaseId": fallback_kb_id,
                "fallback": True
            }), 200
        
        # Create HeyGen knowledge base with strict instructions
        heygen_payload = {
            "name": f"Structured_Interview_{candidate_name.replace(' ', '_')}_{int(time.time())}",
            "description": f"Structured technical interview for {candidate_name} - {position}",
            "content": kb_content,
            "opening_line": opening_line,
            "custom_prompt": "You are a professional interviewer. Follow the structured questions EXACTLY as provided. No casual chat. When you receive INIT_INTERVIEW, immediately ask the first question about self-introduction.",
            "prompt": kb_content,  # Some endpoints use 'prompt' instead of 'content'
            "voice_settings": {
                "rate": 1.0,
                "emotion": "professional"
            }
        }
        
        # Try multiple endpoints
        endpoints = [
            "https://api.heygen.com/v1/streaming/knowledge_base/create",
            "https://api.heygen.com/v1/streaming/knowledge_base",
            "https://api.heygen.com/v1/streaming_avatar/knowledge_base"
        ]
        
        kb_id = None
        successful_endpoint = None
        
        for endpoint in endpoints:
            try:
                logger.info(f"Trying endpoint: {endpoint}")
                resp = requests.post(
                    endpoint,
                    headers={
                        "X-Api-Key": heygen_key,
                        "Content-Type": "application/json",
                        "Accept": "application/json"
                    },
                    json=heygen_payload,
                    timeout=30
                )
                
                if resp.ok:
                    data = resp.json()
                    kb_id = (
                        data.get('data', {}).get('knowledge_base_id') or
                        data.get('knowledge_base_id') or
                        data.get('id')
                    )
                    if kb_id:
                        successful_endpoint = endpoint
                        logger.info(f"KB created successfully: {kb_id}")
                        break
                else:
                    logger.error(f"Endpoint {endpoint} failed: {resp.status_code} - {resp.text[:200]}")
                    
            except Exception as e:
                logger.error(f"Error with endpoint {endpoint}: {e}")
        
        if not kb_id:
            kb_id = f"kb_structured_{candidate_name.replace(' ', '_')}_{int(time.time())}"
            logger.warning(f"All endpoints failed, using fallback: {kb_id}")
        
        # Save to database
        if candidate_id:
            session = SessionLocal()
            try:
                cand = session.query(Candidate).filter_by(id=candidate_id).first()
                if cand:
                    cand.knowledge_base_id = kb_id
                    cand.interview_kb_content = kb_content
                    session.commit()
            finally:
                session.close()
        
        return jsonify({
            "success": True,
            "knowledgeBaseId": kb_id,
            "endpoint_used": successful_endpoint,
            "structured": True,
            "question_count": 10
        }), 200
        
    except Exception as e:
        logger.error(f"KB creation error: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

def create_enhanced_kb_content(candidate_name, position, company, resume_content):
    """Create enhanced knowledge base content WITHOUT putting backslashes inside f-string expressions."""
    # Build any strings that contain backslashes/newlines first:
    resume_highlights = (
        f"Resume Highlights:\n{resume_content[:2000]}..."
        if resume_content else
        "No resume content available - focus on standard interview questions"
    )
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    experience_years = extract_experience_years(resume_content) if resume_content else "Not specified"

    skills_line = ", ".join(skills) if skills else "General software engineering skills"
    # Now the f-string only injects already-built variables (safe):
    return (
        "INTERVIEW SYSTEM CONFIGURATION\n"
        "==============================\n"
        f"MISSION: Conduct a professional, comprehensive interview for {candidate_name}\n"
        f"POSITION: {position}\n"
        f"COMPANY: {company}\n"
        "â±DURATION: 30-45 minutes\n"
        "MODE: AI-Powered Structured Interview\n\n"
        "CANDIDATE BACKGROUND\n"
        "====================\n"
        f"{resume_highlights}\n"
        f"Experience: {experience_years}\n"
        f"Key Skills: {skills_line}\n\n"
        "INTERVIEW FLOW\n"
        "==============\n"
        "1) Intro & Warm-up\n"
        "2) Skills Deep Dive\n"
        "3) Problem Solving\n"
        "4) Behavioral\n"
        "5) Wrap-up\n"
    )

# backend.py - Add this debug endpoint
@app.route('/api/verify-interview-system/<token>', methods=['GET'])
def verify_interview_system(token):
    """Verify complete interview system setup"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Check all components
        checks = {
            "candidate_found": True,
            "resume_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "knowledge_base_created": bool(candidate.knowledge_base_id),
            "heygen_api_configured": bool(os.getenv('HEYGEN_API_KEY')),
            "session_configured": bool(candidate.interview_session_id),
            "recording_ready": True,
            "qa_tracking_ready": True
        }
        
        # Extract sample questions
        sample_questions = []
        if candidate.resume_path:
            resume_content = extract_resume_content(candidate.resume_path)
            skills = extract_skills_from_resume(resume_content)
            sample_questions = [
                f"Tell me about your experience with {skills[0]}" if skills else "Tell me about yourself",
                "What interests you about this position?",
                "Describe a challenging project you've worked on"
            ]
        
        return jsonify({
            "status": "ready" if all(checks.values()) else "issues_found",
            "checks": checks,
            "candidate_info": {
                "name": candidate.name,
                "position": candidate.job_title,
                "knowledge_base_id": candidate.knowledge_base_id
            },
            "sample_questions": sample_questions,
            "recommendations": [
                "Ensure HEYGEN_API_KEY is set" if not checks["heygen_api_configured"] else None,
                "Upload candidate resume" if not checks["resume_exists"] else None,
                "Create knowledge base" if not checks["knowledge_base_created"] else None
            ]
        }), 200
        
    finally:
        session.close()
@app.route('/api/verify-knowledge-base', methods=['GET'])
def verify_kb_by_query():
    kb_id = request.args.get('id')
    if not kb_id:
        return jsonify({"error": "id is required"}), 400
    # If you want: actually check existence in your DB or HeyGen here
    return jsonify({"ok": True, "knowledge_base_id": kb_id}), 200


@app.route('/api/avatar/interviews', methods=['POST','OPTIONS'], endpoint='save_interview_v2')
def save_interview_v2():
    """Create/refresh an interview token for a candidate and persist expiry."""
    if request.method == 'OPTIONS':
        resp = jsonify({})
        resp.status_code = 200
        resp.headers['Access-Control-Allow-Origin'] = '*'
        resp.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        resp.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        return resp

    from sqlalchemy.exc import SQLAlchemyError
    try:
        data = request.get_json(silent=True) or {}
        candidate_email = data.get('candidateEmail')
        incoming_kb_id = data.get('knowledgeBaseId')

        if not candidate_email:
            return jsonify({"error": "candidateEmail is required"}), 400

        interview_token = str(uuid.uuid4())
        expires_at = datetime.now(timezone.utc) + timedelta(days=7)

        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(email=candidate_email).first()
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404

            candidate.interview_token = interview_token
            candidate.interview_expires_at = expires_at
            if incoming_kb_id:
                # only overwrite if caller explicitly sends
                candidate.knowledge_base_id = incoming_kb_id

            session.commit()

            return jsonify({
                "token": interview_token,
                "expiresAt": expires_at.isoformat(),
                "knowledgeBaseId": getattr(candidate, "knowledge_base_id", None)
            }), 200

        except SQLAlchemyError:
            session.rollback()
            logger.exception("DB error saving interview")
            return jsonify({"error": "Database error saving interview"}), 500
        finally:
            session.close()

    except Exception:
        logger.exception("Error saving interview")
        return jsonify({"error": "Failed to save interview"}), 500

@app.route('/api/debug/find-candidate', methods=['POST'])
def debug_find_candidate():
    from sqlalchemy.exc import SQLAlchemyError
    data = request.get_json(silent=True) or {}
    email = data.get('candidateEmail')
    if not email:
        return jsonify({"error": "candidateEmail required"}), 400
    s = SessionLocal()
    try:
        c = s.query(Candidate).filter_by(email=email).first()
        if not c:
            return jsonify({"found": False, "email": email}), 404
        return jsonify({
            "found": True,
            "id": c.id,
            "name": c.name,
            "email": c.email,
            "has_interview_token": hasattr(c, "interview_token"),
            "has_interview_expires_at": hasattr(c, "interview_expires_at"),
            "has_knowledge_base_id": hasattr(c, "knowledge_base_id"),
            "current_interview_token": getattr(c, "interview_token", None),
            "current_kb": getattr(c, "knowledge_base_id", None)
        }), 200
    except SQLAlchemyError as e:
        return jsonify({"error": "db", "details": str(e)}), 500
    finally:
        s.close()

@app.route('/api/avatar/interviews', methods=['POST','OPTIONS'])
def save_interview():
    if request.method == 'OPTIONS':
        return ('', 200)

    data = request.get_json(silent=True) or {}
    email = data.get('candidateEmail')
    if not email:
        return jsonify({"error":"candidateEmail is required"}), 400

    interview_token = str(uuid.uuid4())
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)

    s = SessionLocal()
    try:
        c = s.query(Candidate).filter_by(email=email).first()
        if not c:
            return jsonify({"error":"Candidate not found","email":email}), 404

        # Write fields (these must exist in your model/DB)
        c.interview_token = interview_token
        c.interview_expires_at = expires_at

        incoming_kb = data.get('knowledgeBaseId')
        if incoming_kb:
            c.knowledge_base_id = incoming_kb

        s.commit()
        return jsonify({
            "token": interview_token,
            "expiresAt": expires_at.isoformat(),
            "knowledgeBaseId": getattr(c,"knowledge_base_id",None),
            "candidateEmail": email
        }), 200

    except SQLAlchemyError as e:
        s.rollback()
        # TEMP: return the DB error so we can see it quickly
        return jsonify({"error":"db","details":str(e)}), 500
    except Exception as e:
        s.rollback()
        return jsonify({"error":"server","details":str(e)}), 500
    finally:
        s.close()


@app.route('/api/avatar/interview/<token>', methods=['POST'])
def api_avatar_interview(token):
    """Handle avatar interview updates"""
    try:
        data = request.json
        action = data.get('action')
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(interview_token=token).first()
            
            if not candidate:
                return jsonify({"error": "Interview not found"}), 404
            
            if action == 'start':
                candidate.interview_started_at = datetime.now()
                message = "Interview started"
            elif action == 'complete':
                candidate.interview_completed_at = datetime.now()
                if data.get('transcript'):
                    candidate.interview_transcript = json.dumps(data['transcript'])
                message = "Interview completed"
            else:
                message = "Interview updated"
            
            session.commit()
            
            return jsonify({
                "success": True,
                "message": message
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error in avatar interview {token}: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview-automation/status', methods=['GET', 'OPTIONS'])
def get_automation_status():
    """Get interview automation system status"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        from interview_automation import interview_automation
        
        status = {
            'is_running': interview_automation.is_running,
            'check_interval_minutes': interview_automation.check_interval / 60,
            'next_check': 'Running' if interview_automation.is_running else 'Stopped'
        }
        
        # Get statistics
        session = SessionLocal()
        try:
            stats = {
                'candidates_pending_interview': session.query(Candidate).filter(
                    and_(
                        Candidate.exam_completed == True,
                        Candidate.exam_percentage >= 70,
                        Candidate.interview_scheduled == False
                    )
                ).count(),
                'interviews_scheduled': session.query(Candidate).filter(
                    Candidate.interview_scheduled == True
                ).count(),
                'interviews_completed': session.query(Candidate).filter(
                    Candidate.interview_completed_at.isnot(None)
                ).count()
            }
            status['statistics'] = stats
        finally:
            session.close()
        
        return jsonify(status), 200
        
    except Exception as e:
        logger.error(f"Error getting automation status: {e}")
        return jsonify({"error": str(e)}), 500
    
@app.route('/api/v1/streaming_new', methods=['POST', 'OPTIONS'])
def streaming_new():
    """Proxy requests to HeyGen streaming API"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json or {}
        
        # Get HeyGen API key
        heygen_key = os.getenv('HEYGEN_API_KEY')
        if not heygen_key:
            return jsonify({"error": "HeyGen API key not configured"}), 500
        
        # Use correct HeyGen streaming endpoint
        heygen_streaming_url = "https://api.heygen.com/v1/streaming.new"
        
        headers = {
            'Content-Type': 'application/json',
            'X-Api-Key': heygen_key,
            'Accept': 'application/json'
        }
        
        print(f"Forwarding request to HeyGen: {heygen_streaming_url}")
        print(f"Request data: {data}")
        
        response = requests.post(
            heygen_streaming_url, 
            json=data, 
            headers=headers,
            timeout=30
        )
        
        print(f"HeyGen response status: {response.status_code}")
        
        if response.ok:
            response_data = response.json()
            return jsonify(response_data), response.status_code
        else:
            error_text = response.text
            print(f"HeyGen API error: {error_text}")
            return jsonify({
                "error": "HeyGen API error", 
                "details": error_text,
                "status": response.status_code
            }), response.status_code
            
    except requests.exceptions.Timeout:
        print("HeyGen API timeout")
        return jsonify({"error": "Request timeout"}), 504
    except requests.exceptions.ConnectionError:
        print("HeyGen API connection error")
        return jsonify({"error": "Connection error"}), 503
    except Exception as e:
        print(f"Streaming proxy error: {e}")
        return jsonify({"error": "Avatar service unavailable", "details": str(e)}), 500

@app.route('/api/get-access-token', methods=['POST', 'OPTIONS'])
def get_access_token():
    """Get HeyGen streaming access token"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        heygen_key = os.getenv('HEYGEN_API_KEY')
        if not heygen_key:
            return jsonify({"error": "HeyGen API key not configured"}), 500
        
        # Call HeyGen token creation API
        response = requests.post(
            'https://api.heygen.com/v1/streaming.create_token',
            headers={
                'Accept': 'application/json',
                'Content-Type': 'application/json',
                'X-Api-Key': heygen_key,
            },
            timeout=10
        )
        
        if not response.ok:
            error_text = response.text
            print(f"HeyGen token error: {error_text}")
            return jsonify({"error": "Failed to get token", "details": error_text}), response.status_code
        
        data = response.json()
        if not data.get('data', {}).get('token'):
            return jsonify({"error": "No token in response"}), 500
        
        token = data['data']['token']
        print(f"Token obtained successfully")
        
        # Return as plain text
        return Response(token, mimetype='text/plain', status=200)
        
    except Exception as e:
        print(f"Token error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/debug/avatar', methods=['GET'])
def debug_avatar():
    """Debug avatar configuration"""
    return jsonify({
        "heygen_key_configured": bool(os.getenv('HEYGEN_API_KEY')),
        "heygen_key_length": len(os.getenv('HEYGEN_API_KEY', '')),
        "cors_enabled": True,
        "endpoints": [
            "/api/get-access-token",
            "/api/v1/streaming_new",
            "/secure-interview/<token>"
        ]
    }), 200

@app.route('/api/interview/results', methods=['GET'])
def get_interview_results():
    """Get all interview results with filtering options"""
    session = SessionLocal()
    try:
        # Get query parameters
        position = request.args.get('position')
        status = request.args.get('status')
        date_from = request.args.get('date_from')
        date_to = request.args.get('date_to')
        
        # Build query
        query = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        )
        
        # Apply filters
        if position:
            query = query.filter(Candidate.job_title == position)
        
        if status:
            if status == 'completed':
                query = query.filter(Candidate.interview_completed_at.isnot(None))
            elif status == 'pending':
                query = query.filter(Candidate.interview_completed_at.is_(None))
        
        if date_from:
            query = query.filter(Candidate.interview_date >= date_from)
        
        if date_to:
            query = query.filter(Candidate.interview_date <= date_to)
        
        # Get results
        candidates = query.all()
        
        # Format results
        results = []
        for candidate in candidates:
            results.append({
                'id': candidate.id,
                'name': candidate.name,
                'email': candidate.email,
                'job_title': candidate.job_title,
                'interview_date': candidate.interview_date.isoformat() if candidate.interview_date else None,
                'interview_completed_at': candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                'interview_ai_score': candidate.interview_ai_score,
                'interview_ai_technical_score': candidate.interview_ai_technical_score,
                'interview_ai_communication_score': candidate.interview_ai_communication_score,
                'interview_ai_problem_solving_score': candidate.interview_ai_problem_solving_score,
                'interview_ai_cultural_fit_score': candidate.interview_ai_cultural_fit_score,
                'interview_ai_overall_feedback': candidate.interview_ai_overall_feedback,
                'interview_final_status': candidate.interview_final_status,
                'interview_recording_url': candidate.interview_recording_url,
                'interview_transcript': candidate.interview_transcript
            })
        
        return jsonify({
            'success': True,
            'results': results,
            'total': len(results)
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting interview results: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/stats', methods=['GET'])
def get_interview_stats():
    """Get interview statistics"""
    session = SessionLocal()
    try:
        # Get all interviewed candidates
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        # Calculate statistics
        total = len(candidates)
        completed = len([c for c in candidates if c.interview_completed_at])
        with_scores = len([c for c in candidates if c.interview_ai_score])
        passed = len([c for c in candidates if c.interview_ai_score and c.interview_ai_score >= 70])
        
        avg_score = 0
        if with_scores > 0:
            total_score = sum(c.interview_ai_score for c in candidates if c.interview_ai_score)
            avg_score = total_score / with_scores
        
        # Skills averages
        skills = {
            'technical': 0,
            'communication': 0,
            'problem_solving': 0,
            'cultural_fit': 0
        }
        
        for c in candidates:
            if c.interview_ai_technical_score:
                skills['technical'] += c.interview_ai_technical_score
            if c.interview_ai_communication_score:
                skills['communication'] += c.interview_ai_communication_score
            if c.interview_ai_problem_solving_score:
                skills['problem_solving'] += c.interview_ai_problem_solving_score
            if c.interview_ai_cultural_fit_score:
                skills['cultural_fit'] += c.interview_ai_cultural_fit_score
        
        # Calculate averages
        for skill in skills:
            if with_scores > 0:
                skills[skill] = skills[skill] / with_scores
        
        return jsonify({
            'success': True,
            'stats': {
                'total_interviews': total,
                'completed_interviews': completed,
                'average_score': round(avg_score, 1),
                'pass_rate': round((passed / with_scores * 100), 1) if with_scores > 0 else 0,
                'pending_analysis': completed - with_scores,
                'skills_average': skills
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting interview stats: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/conversation/update', methods=['POST', 'OPTIONS'])
@cross_origin()
def update_conversation_snapshot():
    """Update conversation snapshot with complete data"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        conversation_data = data.get('conversation_data', [])
        
        if not session_id:
            return jsonify({"error": "session_id required"}), 400
        
        session = SessionLocal()
        try:
            # Find candidate by session_id
            candidate = session.query(Candidate).filter_by(
                interview_session_id=session_id
            ).first()
            
            if not candidate:
                # Try alternate lookup
                parts = session_id.split('_')
                if len(parts) >= 2 and parts[1].isdigit():
                    candidate = session.query(Candidate).filter_by(id=int(parts[1])).first()
            
            if not candidate:
                return jsonify({"error": "Session not found"}), 404
            
            # Store structured conversation data
            candidate.interview_conversation_structured = json.dumps(conversation_data)
            
            # Create formatted transcript
            formatted_transcript = format_conversation_transcript(conversation_data, candidate.name)
            candidate.interview_transcript = formatted_transcript
            
            # Update statistics
            questions = [entry for entry in conversation_data if entry.get('type') == 'question']
            answers = [entry for entry in conversation_data if entry.get('type') == 'answer']
            
            candidate.interview_total_questions = len(questions)
            candidate.interview_answered_questions = len(answers)
            
            # Update progress
            if len(questions) > 0:
                progress = (len(answers) / len(questions)) * 100
                candidate.interview_progress_percentage = min(progress, 100)
            
            # Update last activity
            candidate.interview_last_activity = datetime.now()
            
            # Check for auto-completion
            if len(answers) >= 10 or (len(questions) > 0 and len(answers) >= len(questions)):
                if not candidate.interview_completed_at:
                    candidate.interview_completed_at = datetime.now()
                    candidate.interview_ai_analysis_status = 'pending'
                    
                    # Calculate duration
                    if candidate.interview_started_at:
                        duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                        candidate.interview_duration = int(duration)
                    
                    logger.info(f"Auto-completed interview for {candidate.name} with {len(answers)} answers")
                    
                    # Trigger scoring
                    executor.submit(trigger_auto_scoring, candidate.id)
            
            session.commit()
            
            return jsonify({
                "success": True,
                "questions": len(questions),
                "answers": len(answers),
                "progress": candidate.interview_progress_percentage,
                "auto_completed": candidate.interview_completed_at is not None
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Conversation update error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/api/interview/conversation/get/<session_id>', methods=['GET'])
def get_conversation_data(session_id):
    """Get complete conversation data for a session"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(
            interview_session_id=session_id
        ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Get structured conversation data
        conversation_data = json.loads(candidate.interview_conversation_structured or '[]')
        
        return jsonify({
            "success": True,
            "session_id": session_id,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "position": candidate.job_title
            },
            "conversation": conversation_data,
            "statistics": {
                "total_exchanges": len(conversation_data),
                "questions": len([e for e in conversation_data if e.get('type') == 'question']),
                "answers": len([e for e in conversation_data if e.get('type') == 'answer']),
                "progress": candidate.interview_progress_percentage or 0,
                "completed": candidate.interview_completed_at is not None
            },
            "formatted_transcript": candidate.interview_transcript
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting conversation: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


def format_conversation_transcript(conversation_data, candidate_name):
    """Format conversation data into readable transcript"""
    transcript = f"Interview Transcript - {candidate_name}\n"
    transcript += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    transcript += "=" * 70 + "\n\n"
    
    for entry in conversation_data:
        try:
            timestamp = datetime.fromisoformat(entry['timestamp'].replace('Z', '+00:00'))
            time_str = timestamp.strftime('%H:%M:%S')
            speaker = 'AI Interviewer' if entry['speaker'] == 'avatar' else 'Candidate'
            content = entry['content'].strip()
            
            transcript += f"[{time_str}] {speaker}:\n{content}\n\n"
            
        except (KeyError, ValueError) as e:
            logger.warning(f"Error formatting entry: {e}")
            continue
    
    return transcript


@app.route('/api/interview/qa/track-enhanced', methods=['POST', 'OPTIONS'])
@cross_origin()
def track_qa_enhanced():
    """Enhanced Q&A tracking with proper conversation structure"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        entry_type = data.get('type')  # 'question' or 'answer'
        content = data.get('content', '').strip()
        metadata = data.get('metadata', {})
        
        if not session_id or not content or not entry_type:
            return jsonify({"error": "Missing required fields"}), 400
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(
                interview_session_id=session_id
            ).first()
            
            if not candidate:
                return jsonify({"error": "Session not found"}), 404
            
            # Get existing conversation data
            conversation_data = json.loads(candidate.interview_conversation_structured or '[]')
            
            # Create new entry
            entry = {
                'id': metadata.get('entry_id', f"{entry_type}_{int(time.time())}"),
                'type': entry_type,
                'speaker': 'avatar' if entry_type == 'question' else 'candidate',
                'content': content,
                'timestamp': metadata.get('timestamp', datetime.now().isoformat()),
                'sequence': metadata.get('sequence', len(conversation_data) + 1),
                'linked_question_id': metadata.get('linked_question_id'),
                'word_count': metadata.get('word_count', len(content.split())),
                'confidence': metadata.get('confidence', 1.0),
                'is_complete': metadata.get('is_complete', True)
            }
            
            # Add to conversation data
            conversation_data.append(entry)
            
            # Update candidate record
            candidate.interview_conversation_structured = json.dumps(conversation_data)
            candidate.interview_last_activity = datetime.now()
            
            # Update counters
            questions = [e for e in conversation_data if e.get('type') == 'question']
            answers = [e for e in conversation_data if e.get('type') == 'answer']
            
            candidate.interview_total_questions = len(questions)
            candidate.interview_answered_questions = len(answers)
            
            # Update progress
            if len(questions) > 0:
                progress = (len(answers) / len(questions)) * 100
                candidate.interview_progress_percentage = min(progress, 100)
            
            # Update formatted transcript
            candidate.interview_transcript = format_conversation_transcript(
                conversation_data, candidate.name
            )
            
            session.commit()
            
            logger.info(f"Enhanced Q&A tracking: {entry_type} for {candidate.name}")
            
            return jsonify({
                "success": True,
                "entry_id": entry['id'],
                "total_questions": len(questions),
                "answered_questions": len(answers),
                "progress": candidate.interview_progress_percentage
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Enhanced Q&A tracking error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/api/interview/conversation/export/<session_id>', methods=['GET'])
def export_conversation_enhanced(session_id):
    """Export conversation with multiple format options"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(
            interview_session_id=session_id
        ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Get format from query parameter
        format_type = request.args.get('format', 'json')
        
        conversation_data = json.loads(candidate.interview_conversation_structured or '[]')
        
        if format_type == 'json':
            export_data = {
                "session_id": session_id,
                "candidate": {
                    "id": candidate.id,
                    "name": candidate.name,
                    "email": candidate.email,
                    "position": candidate.job_title
                },
                "conversation": conversation_data,
                "statistics": {
                    "total_exchanges": len(conversation_data),
                    "questions": len([e for e in conversation_data if e.get('type') == 'question']),
                    "answers": len([e for e in conversation_data if e.get('type') == 'answer']),
                    "duration": candidate.interview_duration,
                    "progress": candidate.interview_progress_percentage
                },
                "exported_at": datetime.now().isoformat()
            }
            
            return jsonify(export_data), 200
        
        elif format_type == 'text':
            transcript = format_conversation_transcript(conversation_data, candidate.name)
            
            response = Response(
                transcript,
                mimetype='text/plain',
                headers={
                    'Content-Disposition': f'attachment; filename=interview_{candidate.name}_{session_id}.txt'
                }
            )
            return response
        
        else:
            return jsonify({"error": "Invalid format"}), 400
            
    except Exception as e:
        logger.error(f"Export error: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


@app.route('/api/interview/conversation/validate/<session_id>', methods=['GET'])
def validate_conversation_data(session_id):
    """Validate conversation data integrity"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(
            interview_session_id=session_id
        ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        conversation_data = json.loads(candidate.interview_conversation_structured or '[]')
        
        # Validate data integrity
        issues = []
        questions = []
        answers = []
        
        for entry in conversation_data:
            if not entry.get('content') or len(entry['content'].strip()) < 5:
                issues.append(f"Entry {entry.get('id', 'unknown')} has invalid content")
            
            if entry.get('type') == 'question':
                questions.append(entry)
            elif entry.get('type') == 'answer':
                answers.append(entry)
        
        # Check for orphaned answers
        orphaned_answers = [a for a in answers if not a.get('linked_question_id')]
        if orphaned_answers:
            issues.append(f"{len(orphaned_answers)} answers without linked questions")
        
        # Check sequence integrity
        sequences = [entry.get('sequence', 0) for entry in conversation_data]
        if len(set(sequences)) != len(sequences):
            issues.append("Duplicate sequence numbers detected")
        
        validation_result = {
            "is_valid": len(issues) == 0,
            "issues": issues,
            "statistics": {
                "total_entries": len(conversation_data),
                "questions": len(questions),
                "answers": len(answers),
                "orphaned_answers": len(orphaned_answers),
                "expected_questions": candidate.interview_total_questions or 0,
                "expected_answers": candidate.interview_answered_questions or 0
            },
            "data_integrity": {
                "has_structured_data": len(conversation_data) > 0,
                "has_transcript": bool(candidate.interview_transcript),
                "counters_match": (
                    len(questions) == (candidate.interview_total_questions or 0) and
                    len(answers) == (candidate.interview_answered_questions or 0)
                )
            }
        }
        
        return jsonify(validation_result), 200
        
    except Exception as e:
        logger.error(f"Validation error: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


# Database migration function (run once to add new columns)
def migrate_conversation_storage():
    """Migrate existing interview data to new conversation storage format"""
    from sqlalchemy import text
    
    session = SessionLocal()
    try:
        # Add new column if it doesn't exist
        try:
            session.execute(text("""
                ALTER TABLE candidates 
                ADD COLUMN interview_conversation_structured TEXT DEFAULT NULL
            """))
            session.commit()
            print("✅ Added interview_conversation_structured column")
        except Exception:
            print("ℹ️  Column interview_conversation_structured already exists")
        
        # Migrate existing data
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True,
            Candidate.interview_conversation_structured.is_(None)
        ).all()
        
        migrated = 0
        for candidate in candidates:
            try:
                # Try to reconstruct conversation from existing data
                qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
                conversation_data = []
                
                for i, qa in enumerate(qa_pairs):
                    if qa.get('question'):
                        conversation_data.append({
                            'id': f'q_{i}_{int(time.time())}',
                            'type': 'question',
                            'speaker': 'avatar',
                            'content': qa['question'],
                            'timestamp': qa.get('timestamp', datetime.now().isoformat()),
                            'sequence': len(conversation_data) + 1
                        })
                    
                    if qa.get('answer'):
                        conversation_data.append({
                            'id': f'a_{i}_{int(time.time())}',
                            'type': 'answer',
                            'speaker': 'candidate',
                            'content': qa['answer'],
                            'timestamp': qa.get('answered_at', datetime.now().isoformat()),
                            'sequence': len(conversation_data) + 1,
                            'linked_question_id': f'q_{i}_{int(time.time())}'
                        })
                
                if conversation_data:
                    candidate.interview_conversation_structured = json.dumps(conversation_data)
                    migrated += 1
            
            except Exception as e:
                print(f"❌ Failed to migrate candidate {candidate.id}: {e}")
        
        session.commit()
        print(f"✅ Migrated {migrated} candidates to new conversation storage")
        
    except Exception as e:
        print(f"❌ Migration failed: {e}")
        session.rollback()
    finally:
        session.close()

@app.route('/api/interview-automation/toggle', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=5, time_window=60)
def toggle_automation():
    """Start or stop the interview automation system"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        action = data.get('action', 'toggle')
        
        from interview_automation import interview_automation
        
        if action == 'start':
            start_interview_automation()
            message = "Interview automation started"
        elif action == 'stop':
            stop_interview_automation()
            message = "Interview automation stopped"
        else:
            # Toggle
            if interview_automation.is_running:
                stop_interview_automation()
                message = "Interview automation stopped"
            else:
                start_interview_automation()
                message = "Interview automation started"
        
        return jsonify({
            'success': True,
            'message': message,
            'is_running': interview_automation.is_running
        }), 200
        
    except Exception as e:
        logger.error(f"Error toggling automation: {e}")
        return jsonify({"error": str(e)}), 500

# 3. Fix the api_schedule_interview function to properly handle company_name
@app.route('/api/schedule-interview', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=10, time_window=60)
def api_schedule_interview():
    """Schedule interview with enhanced knowledge base creation for proactive questioning"""
    if request.method == 'OPTIONS':
        return '', 200
        
    try:
        data = request.json
        candidate_id = data.get('candidate_id')
        email = data.get('email')
        interview_date = data.get('date')
        time_slot = data.get('time_slot')
        job_description_override = data.get('job_description')
        
        logger.info(f"Schedule interview request: candidate_id={candidate_id}")
        
        if not candidate_id and not email:
            return jsonify({"success": False, "message": "candidate_id or email is required"}), 400
        
        session = SessionLocal()
        try:
            # Find candidate
            if candidate_id:
                candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            else:
                candidate = session.query(Candidate).filter_by(email=email).first()
            
            if not candidate:
                return jsonify({"success": False, "message": "Candidate not found"}), 404
            
            # Check if already scheduled
            if candidate.interview_scheduled and candidate.interview_token:
                existing_link = f"{request.host_url.rstrip('/')}/secure-interview/{candidate.interview_token}"
                return jsonify({
                    "success": True,
                    "message": "Interview already scheduled",
                    "interview_link": existing_link,
                    "knowledge_base_id": getattr(candidate, 'knowledge_base_id', None),
                    "already_scheduled": True
                }), 200
            
            # Extract resume content
            resume_content = ""
            resume_extracted = False
            
            if candidate.resume_path and os.path.exists(candidate.resume_path):
                logger.info(f"Extracting resume from: {candidate.resume_path}")
                resume_content = extract_resume_content(candidate.resume_path)
                if resume_content:
                    resume_extracted = True
                    logger.info(f"Resume extracted: {len(resume_content)} characters")
                else:
                    logger.error("Resume extraction returned empty content")
            
            # Fallback to candidate profile if no resume
            if not resume_content:
                logger.warning("Using candidate profile as fallback")
                resume_content = f"""
CANDIDATE: {candidate.name}
EMAIL: {candidate.email}
POSITION: {candidate.job_title}
ATS SCORE: {candidate.ats_score}
STATUS: {candidate.status}
{f"SCORING: {candidate.score_reasoning}" if candidate.score_reasoning else ""}
"""
            
            # Get company name
            company_name = os.getenv('COMPANY_NAME', 'Our Company')
            
            # Get job description
            job_description = job_description_override or getattr(candidate, 'job_description', f"Position: {candidate.job_title}")
            
            # CREATE HEYGEN KNOWLEDGE BASE WITH INTERVIEW QUESTIONS
            knowledge_base_id = None
            kb_creation_method = "none"
            
            if os.getenv('HEYGEN_API_KEY') and resume_content:
                try:
                    logger.info("Creating HeyGen knowledge base with interview questions...")
                    
                    # Generate structured interview questions
                    interview_questions = generate_interview_questions(
                        candidate_name=candidate.name,
                        position=candidate.job_title,
                        resume_content=resume_content,
                        job_description=job_description
                    )
                    
                    kb_name = f"Interview_{candidate.name.replace(' ', '_')}_{candidate.id}"
                    
                    # Create comprehensive knowledge base content
                    kb_content = f"""
INTERVIEW CONFIGURATION:
- Mode: Structured Technical Interview
- Candidate: {candidate.name}
- Position: {candidate.job_title}
- Company: {company_name}
- Interview Type: Technical and Behavioral
- Duration: 30-45 minutes

SPECIAL COMMANDS:
- When you receive "INIT_INTERVIEW": Start with the warm greeting and first question
- When you receive "NEXT_QUESTION": Move to the next question in the list
- If user is silent for 15+ seconds: Gently prompt or ask if they need more time


CANDIDATE BACKGROUND:
{resume_content[:8000]}

JOB REQUIREMENTS:
{job_description[:2000]}

{interview_questions}

INTERVIEW BEHAVIOR INSTRUCTIONS:
1. When stream starts, wait for "INIT_INTERVIEW" command
2. Upon receiving "INIT_INTERVIEW", immediately greet the candidate and ask the first question
3. Listen to complete answers before proceeding
4. Ask follow-up questions when appropriate
5. Keep track of which questions you've asked
6. Be encouraging if candidate seems nervous
7. End professionally after covering all questions

CONVERSATION STARTERS:
- If you receive any greeting like "Hello", "Hi", respond with: "Hello {candidate.name}! Welcome to your interview for {candidate.job_title} at {company_name}. I'm excited to learn about your experience. Let's start with you telling me about yourself and your journey to applying for this role."
- If candidate asks "Can you hear me?", respond: "Yes, I can hear you clearly! Let's begin with our interview. Please tell me about yourself."
- If candidate seems confused, say: "No worries! This is an AI-powered interview. I'll be asking you questions about your experience and the {candidate.job_title} role. Shall we start?"

IMPORTANT RULES:
- Start immediately when you receive "INIT_INTERVIEW"
- Always maintain a professional yet friendly tone
- Give candidates time to think (10-15 seconds)
- If no response after 20 seconds, ask: "Take your time, or would you like me to rephrase the question?"
- Track answered questions to avoid repetition
"""                    
                    # Prepare HeyGen payload with proper configuration
                    heygen_payload = {
                        'name': kb_name,
                        'description': f'Structured interview for {candidate.name} - {candidate.job_title}',
                        'content': kb_content,
                        'opening_line': f"Hello {candidate.name}, welcome to your interview for the {candidate.job_title} position at {company_name}. I'm your AI interviewer today. I've reviewed your resume and I'm excited to learn more about your experiences. Let's start with you telling me a bit about yourself and your journey to applying for this role.",
                        'custom_prompt': f"""You are conducting a professional technical interview for {candidate.name}. 
                        
Your personality: Professional, friendly, encouraging, and engaged.

Key behaviors:
1. Ask questions from the provided list ONE AT A TIME
2. Wait for complete answers before proceeding
3. Show active listening with phrases like "That's interesting", "I see", "Tell me more"
4. If they struggle, offer encouragement: "Take your time", "No worries"
5. Ask follow-up questions based on their responses
6. Keep track of which questions you've asked to avoid repetition

Interview style:
- Conversational, not robotic
- Professional but warm
- Encouraging when candidate seems nervous
- Patient with responses

Remember: This is a conversation, not an interrogation. Make {candidate.name} feel comfortable while thoroughly assessing their qualifications for the {candidate.job_title} role."""
                    }
                    
                    # Make API call to HeyGen
                    heygen_response = requests.post(
                        'https://api.heygen.com/v1/streaming/knowledge_base',
                        headers={
                            'X-Api-Key': os.getenv('HEYGEN_API_KEY'),
                            'Content-Type': 'application/json',
                            'Accept': 'application/json'
                        },
                        json=heygen_payload,
                        timeout=30
                    )
                    
                    if heygen_response.ok:
                        kb_data = heygen_response.json()
                        knowledge_base_id = kb_data.get('data', {}).get('knowledge_base_id')
                        kb_creation_method = "heygen_api"
                        logger.info(f"HeyGen KB created successfully: {knowledge_base_id}")
                    else:
                        error_text = heygen_response.text
                        logger.error(f"HeyGen API error: {heygen_response.status_code} - {error_text}")
                        
                except Exception as e:
                    logger.error(f"HeyGen KB creation failed: {e}", exc_info=True)
            
            # Fallback KB ID if HeyGen fails
            if not knowledge_base_id:
                knowledge_base_id = f"kb_{candidate.id}_{int(time.time())}"
                kb_creation_method = "fallback"
                logger.warning(f"Using fallback KB: {knowledge_base_id}")
            
            # Create interview session
            interview_token = str(uuid.uuid4())
            interview_session_id = f"session_{candidate.id}_{int(time.time())}"
            
            # Parse interview date
            if isinstance(interview_date, str):
                interview_datetime = datetime.fromisoformat(interview_date.replace('Z', '+00:00'))
            else:
                interview_datetime = datetime.now() + timedelta(days=3)
            
            # Update candidate record
            candidate.interview_scheduled = True
            candidate.interview_date = interview_datetime
            candidate.interview_token = interview_token
            candidate.interview_link = f"{request.host_url.rstrip('/')}/secure-interview/{interview_token}"
            candidate.final_status = 'Interview Scheduled'
            
            # Safe attribute setting for optional fields
            safe_attrs = {
                'interview_session_id': interview_session_id,
                'knowledge_base_id': knowledge_base_id,
                'interview_created_at': datetime.now(),
                'interview_expires_at': datetime.now() + timedelta(days=7),
                'company_name': company_name,
                'interview_time_slot': time_slot,
                'interview_questions_asked': '[]',
                'interview_answers_given': '[]',
                'interview_total_questions': 0,
                'interview_answered_questions': 0,
                'job_description': job_description if job_description_override else None
            }
            
            for attr, value in safe_attrs.items():
                if hasattr(candidate, attr):
                    setattr(candidate, attr, value)
            
            # Commit changes
            session.commit()
            
            # Send email
            email_sent = False
            try:
                send_interview_link_email(
                    candidate_email=candidate.email,
                    candidate_name=candidate.name,
                    interview_link=candidate.interview_link,
                    interview_date=interview_datetime,
                    time_slot=time_slot,
                    position=candidate.job_title
                )
                email_sent = True
                logger.info(f"Interview email sent to {candidate.email}")
            except Exception as e:
                logger.error(f"Email failed: {e}")
            
            # Clear caches
            cache.delete_memoized(get_cached_candidates)
            
            return jsonify({
                "success": True,
                "message": f"Interview scheduled for {candidate.name}",
                "interview_link": candidate.interview_link,
                "interview_date": interview_datetime.isoformat(),
                "knowledge_base_id": knowledge_base_id,
                "kb_creation_method": kb_creation_method,
                "resume_extracted": resume_extracted,
                "resume_content_length": len(resume_content),
                "email_sent": email_sent,
                "session_id": interview_session_id
            }), 200
            
        except Exception as e:
            session.rollback()
            logger.error(f"Error in schedule_interview: {e}", exc_info=True)
            return jsonify({"success": False, "message": str(e)}), 500
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Critical error: {e}", exc_info=True)
        return jsonify({"success": False, "message": str(e)}), 500


def generate_interview_questions(candidate_name, position, resume_content, job_description):
    """Generate structured interview questions based on resume and job"""
    
    # Extract key skills from resume
    skills = extract_skills_from_resume(resume_content)
    
    questions = f"""
INTERVIEW QUESTIONS:

1. INTRODUCTION (Ask first):
   - "Tell me about yourself and your journey to applying for this {position} role."
   - "What attracted you to our company and this position?"

2. TECHNICAL QUESTIONS (Based on resume):"""
    
    # Add technical questions based on skills found
    if 'python' in resume_content.lower():
        questions += """
   - "I see you have Python experience. Can you tell me about a complex Python project you've worked on?"
   - "How do you handle error handling and debugging in Python?"""
   
    if 'javascript' in resume_content.lower() or 'react' in resume_content.lower():
        questions += """
   - "Tell me about your experience with JavaScript/React. What was the most challenging frontend problem you've solved?"
   - "How do you manage state in React applications?"""
    
    if 'database' in resume_content.lower() or 'sql' in resume_content.lower():
        questions += """
   - "Describe your experience with databases. How do you optimize slow queries?"
   - "Tell me about a time you designed a database schema."""
    
    questions += f"""

3. BEHAVIORAL QUESTIONS:
   - "Describe a time when you had to work under pressure. How did you handle it?"
   - "Tell me about a project where you had to collaborate with a difficult team member."
   - "Give me an example of when you had to learn a new technology quickly."

4. ROLE-SPECIFIC QUESTIONS:
   - "How do you see yourself contributing to our team in the first 90 days?"
   - "What aspects of this {position} role excite you the most?"

5. CLOSING QUESTIONS:
   - "What questions do you have for me about the role or the company?"
   - "Is there anything else you'd like me to know about your qualifications?"

REMEMBER: Ask these questions one at a time, wait for complete responses, and ask relevant follow-up questions based on their answers.
"""
    
    return questions


# Add this helper endpoint to check interview status
@app.route('/api/interview-status/<int:candidate_id>', methods=['GET'])
def get_interview_v1_status(candidate_id):
    """Check if interview is already scheduled for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "exam_completed": candidate.exam_completed,
            "exam_percentage": candidate.exam_percentage,
            "interview_scheduled": candidate.interview_scheduled,
            "interview_token": candidate.interview_token,
            "interview_link": candidate.interview_link,
            "final_status": candidate.final_status
        }), 200
    finally:
        session.close()

interview_executor = ThreadPoolExecutor(max_workers=2)

@app.route('/api/interview/track-link-click/<token>', methods=['POST', 'OPTIONS'])
@cross_origin()
def track_interview_link_click(token):
    """Track when a candidate clicks on their interview link"""
    
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        logger.info(f"Tracking link click for token: {token}")
        
        # Import your models and db session
        from models import Candidate, InterviewSession, db
        
        # Find the candidate
        candidate = Candidate.query.filter_by(interview_token=token).first()
        
        if not candidate:
            logger.error(f"No candidate found with token: {token}")
            return jsonify({'error': 'Invalid interview token'}), 404
        
        logger.info(f"Found candidate: {candidate.id} - {candidate.name}")
        
        # Update the interview status from "Interview Scheduled" to "in_progress"
        if candidate.interview_status == "Interview Scheduled":
            candidate.interview_status = "in_progress"
        
        # Set the link clicked timestamp
        if not hasattr(candidate, 'interview_link_clicked_at') or not candidate.interview_link_clicked_at:
            candidate.interview_link_clicked_at = datetime.utcnow()
        
        # Also update the interview session if it exists
        session = InterviewSession.query.filter_by(
            candidate_id=candidate.id
        ).order_by(InterviewSession.created_at.desc()).first()
        
        if session:
            session.status = 'in_progress'
            session.last_activity = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'candidate_id': candidate.id,
            'message': 'Link click tracked successfully',
            'interview_status': candidate.interview_status
        }), 200
        
    except Exception as e:
        logger.error(f"Error tracking link click: {str(e)}")
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# Update your force_complete_interview function with more logging
@app.route('/api/interview/force-complete/<token>', methods=['POST'])
def force_complete_interview(token):
    """Force complete an interview regardless of Q&A count"""
    session = SessionLocal()
    try:
        logger.info(f"Force completing interview for token: {token}")
        
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        if not candidate:
            logger.error(f"No candidate found with token: {token}")
            return jsonify({"error": "Interview not found"}), 404
        
        logger.info(f"Found candidate: {candidate.id} - {candidate.name}")
        logger.info(f"Current completed_at: {candidate.interview_completed_at}")
        
        # Force completion
        if not candidate.interview_completed_at:
            candidate.interview_completed_at = datetime.now()
            candidate.interview_status = 'completed'
            candidate.interview_progress_percentage = 100
            
            # Calculate duration
            if candidate.interview_started_at:
                duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                candidate.interview_duration = int(duration)
            
            # Mark for AI analysis
            candidate.interview_ai_analysis_status = 'pending'
            candidate.final_status = 'Interview Completed'
            
            # Log the changes before commit
            logger.info(f"Setting completed_at to: {candidate.interview_completed_at}")
            logger.info(f"Setting final_status to: {candidate.final_status}")
            
            # Commit the changes
            session.commit()
            logger.info("Changes committed to database")
            
            # Verify the commit worked
            session.refresh(candidate)
            logger.info(f"After commit - completed_at: {candidate.interview_completed_at}")
            
            # Trigger scoring
            try:
                trigger_auto_scoring(candidate.id)
            except Exception as e:
                logger.error(f"Failed to trigger auto-scoring: {e}")
            
        return jsonify({
            "success": True,
            "message": "Interview completed",
            "candidate_id": candidate.id,
            "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error force completing interview: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


@app.route('/api/interview/session/progress', methods=['POST', 'OPTIONS'])
def update_session_progress():
    if request.method == 'OPTIONS':
        return '', 200
    
    data = request.json or {}
    session_id = data.get('session_id')
    progress = data.get('progress', 0)
    status = data.get('status', 'in_progress')
    
    if not session_id:
        return jsonify({"error": "session_id is required"}), 400
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_session_id=session_id).first()
        
        if not candidate:
            # Try splitting session_id and matching by candidate ID
            parts = session_id.split('_')
            if len(parts) >= 2 and parts[1].isdigit():
                candidate_id = int(parts[1])
                candidate = session.query(Candidate).filter_by(id=candidate_id).first()
                if candidate:
                    # Update session_id if found
                    candidate.interview_session_id = session_id
        
        if not candidate:
            logger.warning(f"No candidate found for session_id: {session_id}")
            return jsonify({"error": "Session not found"}), 404
        
        # Update progress fields
        if hasattr(candidate, 'interview_progress_percentage'):
            candidate.interview_progress_percentage = float(progress)
        
        if hasattr(candidate, 'interview_last_activity'):
            candidate.interview_last_activity = datetime.now()
        
        # Handle completion status
        if status == 'completed' or progress >= 100:
            if hasattr(candidate, 'interview_completed_at'):
                if not candidate.interview_completed_at:
                    candidate.interview_completed_at = datetime.now()
                    if hasattr(candidate, 'interview_status'):
                        candidate.interview_status = 'completed'
                    candidate.final_status = 'Interview Completed'
                    
                    if hasattr(candidate, 'interview_started_at') and candidate.interview_started_at:
                        duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                        if hasattr(candidate, 'interview_duration'):
                            candidate.interview_duration = int(duration)
                    
                    if hasattr(candidate, 'interview_ai_analysis_status'):
                        candidate.interview_ai_analysis_status = 'pending'
                    
                    logger.info(f"Interview marked as completed for candidate {candidate.id}")
                    # session.commit()

                    try:
                        if 'trigger_auto_scoring' in globals():
                            trigger_auto_scoring(candidate.id)
                        else:
                            logger.warning("Auto-scoring function not found.")
                    except Exception as e:
                        logger.error(f"Error triggering auto-scoring: {e}")
        
        elif status == 'in_progress' and hasattr(candidate, 'interview_started_at'):
            if not candidate.interview_started_at:
                candidate.interview_started_at = datetime.now()
                if hasattr(candidate, 'interview_status'):
                    candidate.interview_status = 'in_progress'
        
        # if session.dirty:
        session.commit()

        cache.delete_memoized(get_cached_candidates)

        return jsonify({
            "success": True,
            "progress": progress,
            "status": status,
            "candidate_id": candidate.id,
            "candidate_name": candidate.name
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error updating progress: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


@app.route('/api/interview/session/complete', methods=['POST', 'OPTIONS'])
def complete_interview_session():
    """Mark interview session as complete"""
    if request.method == 'OPTIONS':
        return '', 200
    
    data = request.json or {}
    session_id = data.get('session_id')
    interview_token = data.get('interview_token')
    
    session = SessionLocal()
    try:
        # Find candidate by session_id or token
        candidate = None
        if session_id:
            candidate = session.query(Candidate).filter_by(
                interview_session_id=session_id
            ).first()
        elif interview_token:
            candidate = session.query(Candidate).filter_by(
                interview_token=interview_token
            ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Mark as completed
        candidate.interview_completed_at = datetime.now()
        candidate.interview_progress_percentage = 100
        candidate.final_status = 'Interview Completed - Pending Review'
        
        # Update interview status if it exists
        if hasattr(candidate, 'interview_status'):
            candidate.interview_status = 'completed'
        
        # Calculate interview duration if interview started
        if candidate.interview_started_at:
            duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
            if hasattr(candidate, 'interview_duration'):
                candidate.interview_duration = int(duration)
        
        # Update Q&A stats if interview_total_questions exists
        if hasattr(candidate, 'interview_qa_completion_rate'):
            if candidate.interview_total_questions and candidate.interview_total_questions > 0:
                completion_rate = (candidate.interview_answered_questions / candidate.interview_total_questions) * 100
                candidate.interview_qa_completion_rate = completion_rate
        
        # Set AI analysis status to pending if the attribute exists
        if hasattr(candidate, 'interview_ai_analysis_status'):
            candidate.interview_ai_analysis_status = 'pending'
        
        session.commit()
        
        logger.info(f"Interview completed for candidate {candidate.id} - {candidate.name}")
        
        # Trigger analysis in background (auto-scoring)
        try:
            if 'trigger_auto_scoring' in globals():
                from concurrent.futures import ThreadPoolExecutor
                executor = ThreadPoolExecutor(max_workers=2)
                executor.submit(trigger_auto_scoring, candidate.id)
        except Exception as e:
            logger.error(f"Failed to trigger analysis: {e}")
        
        return jsonify({
            "success": True,
            "message": "Interview completed successfully",
            "candidate_id": candidate.id,
            "duration_seconds": candidate.interview_duration if hasattr(candidate, 'interview_duration') else None,
            "next_step": "AI analysis in progress"
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error completing interview session: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

# Add to backend.py
def periodic_interview_completion_check():
    """Periodically check for interviews that should be completed"""
    while True:
        try:
            time.sleep(300)  # Check every 5 minutes
            
            session = SessionLocal()
            try:
                # Find uncompleted interviews
                uncompleted = session.query(Candidate).filter(
                    Candidate.interview_started_at.isnot(None),
                    Candidate.interview_completed_at.is_(None)
                ).all()
                
                for candidate in uncompleted:
                    check_and_complete_interview(candidate.id)
                    
            finally:
                session.close()
                
        except Exception as e:
            logger.error(f"Error in periodic completion check: {e}")

# Start this when the app starts
completion_check_thread = threading.Thread(
    target=periodic_interview_completion_check,
    daemon=True
)
completion_check_thread.start()


def check_and_update_expired_interviews():
    """Automatically check and mark expired interviews"""
    session = SessionLocal()
    try:
        now = datetime.now()
        
        # Find expired interviews
        expired_candidates = session.query(Candidate).filter(
            Candidate.interview_expires_at < now,
            Candidate.interview_completed_at.is_(None),
            Candidate.interview_status != 'expired'
        ).all()
        
        for candidate in expired_candidates:
            candidate.interview_status = 'expired'
            candidate.final_status = 'Interview Link Expired'
            logger.info(f"Marked interview as expired for {candidate.name}")
        
        # Find abandoned interviews (no activity for 2 hours)
        two_hours_ago = now - timedelta(hours=2)
        abandoned_candidates = session.query(Candidate).filter(
            Candidate.interview_started_at.isnot(None),
            Candidate.interview_completed_at.is_(None),
            Candidate.interview_last_activity < two_hours_ago,
            Candidate.interview_status != 'abandoned'
        ).all()
        
        for candidate in abandoned_candidates:
            candidate.interview_status = 'abandoned'
            candidate.final_status = 'Interview Abandoned'
            logger.info(f"Marked interview as abandoned for {candidate.name}")
        
        session.commit()
        
    except Exception as e:
        logger.error(f"Error checking expired interviews: {e}")
        session.rollback()
    finally:
        session.close()

# # 1. LIGHTWEIGHT STATUS ENDPOINT (for frequent checks)
@app.route('/api/interview/status/<token>', methods=['GET'])
def get_interview_status(token):
    """Get current interview status - LIGHTWEIGHT"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Determine current status
        status = 'not_started'
        if candidate.interview_completed_at:
            status = 'completed'
        elif hasattr(candidate, 'interview_status') and candidate.interview_status:
            status = candidate.interview_status
        elif candidate.interview_started_at:
            status = 'in_progress'
        elif candidate.link_clicked or (hasattr(candidate, 'interview_link_clicked') and candidate.interview_link_clicked):
            status = 'link_clicked'
        
        # Return only essential data
        return jsonify({
            "status": status,
            "progress": candidate.interview_progress_percentage if hasattr(candidate, 'interview_progress_percentage') else 0,
            "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
            "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
            "expires_at": candidate.interview_expires_at.isoformat() if candidate.interview_expires_at else None
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/dashboard', methods=['GET'])
def interview_tracking_dashboard():
    """Get dashboard data for all interviews"""
    session = SessionLocal()
    try:
        # Get all scheduled interviews
        all_interviews = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        stats = {
            "total_scheduled": len(all_interviews),
            "not_started": 0,
            "link_clicked": 0,
            "in_progress": 0,
            "completed": 0,
            "expired": 0,
            "abandoned": 0
        }
        
        interviews = []
        
        for candidate in all_interviews:
            # Determine status
            if candidate.interview_completed_at:
                status = 'completed'
                stats['completed'] += 1
            elif candidate.interview_status == 'expired':
                status = 'expired'
                stats['expired'] += 1
            elif candidate.interview_status == 'abandoned':
                status = 'abandoned'
                stats['abandoned'] += 1
            elif candidate.interview_started_at:
                status = 'in_progress'
                stats['in_progress'] += 1
            elif candidate.interview_link_clicked:
                status = 'link_clicked'
                stats['link_clicked'] += 1
            else:
                status = 'not_started'
                stats['not_started'] += 1
            
            interviews.append({
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title,
                "status": status,
                "scheduled_date": candidate.interview_date.isoformat() if candidate.interview_date else None,
                "link_clicked": bool(candidate.interview_link_clicked),
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "progress": candidate.interview_progress_percentage or 0,
                "ai_score": candidate.interview_ai_score,
                "final_status": candidate.final_status
            })
        
        return jsonify({
            "success": True,
            "stats": stats,
            "interviews": interviews,
            "last_updated": datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting dashboard data: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

def analyze_answer_quality(answer_text):
    """Analyze individual answer quality"""
    score = 50  # Base score
    
    # Length bonus
    word_count = len(answer_text.split())
    if word_count > 100:
        score += 20
    elif word_count > 50:
        score += 10
    elif word_count < 10:
        score -= 20
    
    # Technical keywords bonus
    tech_keywords = ['implemented', 'developed', 'designed', 'architecture', 
                    'framework', 'database', 'algorithm', 'optimization']
    for keyword in tech_keywords:
        if keyword.lower() in answer_text.lower():
            score += 5
    
    # STAR method indicators
    star_keywords = ['situation', 'task', 'action', 'result', 'challenge', 'solution']
    for keyword in star_keywords:
        if keyword.lower() in answer_text.lower():
            score += 3
    
    return min(100, max(0, score))

# Add this improved scoring function to backend.py
def analyze_interview_with_ai_enhanced(qa_pairs, candidate):
    """Enhanced AI analysis with actual content evaluation"""
    try:
        # Check if OpenAI is available
        openai_key = os.getenv('OPENAI_API_KEY')
        if openai_key:
            import openai
            openai.api_key = openai_key
            
            # Prepare Q&A for analysis
            qa_text = "\n".join([
                f"Q: {qa.get('question', '')}\nA: {qa.get('answer', '')}" 
                for qa in qa_pairs if qa.get('answer')
            ])
            
            # Get AI analysis
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{
                    "role": "system",
                    "content": f"You are evaluating interview responses for {candidate.job_title} position. Score each category 0-100."
                }, {
                    "role": "user", 
                    "content": f"""Analyze these interview Q&As and provide scores:
                    
{qa_text}

Return JSON with:
- technical_score: (0-100)
- communication_score: (0-100)  
- problem_solving_score: (0-100)
- cultural_fit_score: (0-100)
- overall_score: (0-100)
- feedback: (detailed text feedback)
- strengths: (list of strengths)
- weaknesses: (list of areas to improve)"""
                }],
                temperature=0.3
            )
            
            result = json.loads(response.choices[0].message.content)
            return result
            
        else:
            # Fallback to enhanced rule-based scoring
            return analyze_with_enhanced_rules(qa_pairs, candidate)
            
    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        return analyze_with_enhanced_rules(qa_pairs, candidate)

def analyze_with_enhanced_rules(qa_pairs, candidate):
    """Enhanced rule-based scoring with better metrics"""
    if not qa_pairs:
        return {
            'technical_score': 0,
            'communication_score': 0,
            'problem_solving_score': 0,
            'cultural_fit_score': 0,
            'overall_score': 0,
            'confidence': 0,
            'feedback': 'No interview data available'
        }
    
    # Calculate metrics
    answered = sum(1 for qa in qa_pairs if qa.get('answer'))
    total = len(qa_pairs)
    completion_rate = (answered / total * 100) if total > 0 else 0
    
    # Calculate answer quality metrics
    answer_lengths = [len(qa.get('answer', '')) for qa in qa_pairs if qa.get('answer')]
    avg_length = sum(answer_lengths) / len(answer_lengths) if answer_lengths else 0
    
    # Response times
    response_times = [qa.get('response_time', 0) for qa in qa_pairs if qa.get('response_time')]
    avg_response_time = sum(response_times) / len(response_times) if response_times else 0
    
    # Calculate scores with better logic
    technical_score = min(100, max(0, 
        50 + # Base score
        (completion_rate * 0.2) + # 20% from completion
        (min(avg_length / 10, 30)) # Up to 30 points for detailed answers
    ))
    
    communication_score = min(100, max(0,
        40 + # Base score
        (completion_rate * 0.3) + # 30% from completion
        (30 if avg_response_time < 30 else 15) # Quick responses
    ))
    
    problem_solving_score = min(100, max(0,
        45 + # Base score
        (completion_rate * 0.25) + # 25% from completion
        (min(avg_length / 8, 30)) # Detail in answers
    ))
    
    cultural_fit_score = min(100, max(0,
        50 + # Base score
        (completion_rate * 0.3) + # 30% from completion
        (20 if answered >= total * 0.8 else 0) # Bonus for high completion
    ))
    
    overall_score = (
        technical_score * 0.35 +
        communication_score * 0.25 +
        problem_solving_score * 0.25 +
        cultural_fit_score * 0.15
    )

     # Generate strengths and weaknesses
    strengths = []
    weaknesses = []
    
    if completion_rate >= 80:
        strengths.append("High completion rate")
    else:
        weaknesses.append("Low completion rate")
    
    if avg_length > 150:
        strengths.append("Detailed responses")
    elif avg_length < 50:
        weaknesses.append("Brief responses")
    
    if avg_response_time < 15:
        strengths.append("Quick response time")
    elif avg_response_time > 30:
        weaknesses.append("Slow response time")
    
    return {
        'technical_score': round(technical_score),
        'communication_score': round(communication_score),
        'problem_solving_score': round(problem_solving_score),
        'cultural_fit_score': round(cultural_fit_score),
        'overall_score': round(overall_score),
        'confidence': 0.75,
        'feedback': f"""
Interview Analysis (Rule-Based):
- Completion Rate: {completion_rate:.1f}%
- Questions Answered: {answered}/{total}
- Average Response Detail: {'Detailed' if avg_length > 150 else 'Brief' if avg_length > 50 else 'Very Brief'}
- Response Speed: {'Quick' if avg_response_time < 15 else 'Moderate' if avg_response_time < 30 else 'Thoughtful'}

Overall: {'Strong candidate' if overall_score >= 70 else 'Potential candidate' if overall_score >= 50 else 'Needs improvement'}
"""
    }

# Update the trigger_auto_scoring function in backend.py:
def trigger_auto_scoring(candidate_id):
    """Automatically trigger AI scoring when interview completes"""
    def run_scoring():
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                logger.error(f"Candidate {candidate_id} not found for scoring")
                return
            
            # Update status to processing
            candidate.interview_ai_analysis_status = 'processing'
            session.commit()
            
            # Parse Q&A data
            qa_pairs = []
            try:
                # Try to parse Q&A pairs
                if candidate.interview_qa_pairs:
                    qa_pairs = json.loads(candidate.interview_qa_pairs)
                elif candidate.interview_questions_asked and candidate.interview_answers_given:
                    # Build Q&A pairs from separate fields
                    questions = json.loads(candidate.interview_questions_asked or '[]')
                    answers = json.loads(candidate.interview_answers_given or '[]')
                    
                    for i, question in enumerate(questions):
                        qa_pair = {
                            'question': question.get('text', ''),
                            'answer': answers[i].get('text', '') if i < len(answers) else '',
                            'timestamp': question.get('timestamp', None)
                        }
                        qa_pairs.append(qa_pair)
            except Exception as e:
                logger.error(f"Failed to parse Q&A data: {e}")
                qa_pairs = []
            
            # Perform analysis
            if len(qa_pairs) > 0:
                # Use your enhanced analysis function
                scores = analyze_interview_with_ai_enhanced(qa_pairs, candidate)
            else:
                # Default scores if no Q&A data
                scores = {
                    'technical_score': 0,
                    'communication_score': 0,
                    'problem_solving_score': 0,
                    'cultural_fit_score': 0,
                    'overall_score': 0,
                    'feedback': 'No interview data available for analysis',
                    'confidence': 0
                }
            
            # Update candidate with scores
            candidate.interview_ai_score = scores.get('overall_score', 0)
            candidate.interview_ai_technical_score = scores.get('technical_score', 0)
            candidate.interview_ai_communication_score = scores.get('communication_score', 0)
            candidate.interview_ai_problem_solving_score = scores.get('problem_solving_score', 0)
            candidate.interview_ai_cultural_fit_score = scores.get('cultural_fit_score', 0)
            candidate.interview_ai_overall_feedback = scores.get('feedback', '')
            
            # Extract strengths and weaknesses
            if 'strengths' in scores:
                candidate.interview_ai_strengths = json.dumps(scores['strengths'])
            if 'weaknesses' in scores:
                candidate.interview_ai_weaknesses = json.dumps(scores['weaknesses'])
            
            # Set final status based on score
            if candidate.interview_ai_score >= 70:
                candidate.interview_final_status = 'Passed'
                candidate.final_status = 'Interview Passed'
            else:
                candidate.interview_final_status = 'Failed'
                candidate.final_status = 'Interview Failed'
            
            # Mark analysis as complete
            candidate.interview_ai_analysis_status = 'completed'
            
            # Store completion timestamp
            if hasattr(candidate, 'interview_ai_analysis_completed_at'):
                candidate.interview_ai_analysis_completed_at = datetime.now()
            
            session.commit()
            
            # Clear cache to update frontend
            cache.delete_memoized(get_cached_candidates)
            
            logger.info(f"Auto-scoring completed for candidate {candidate_id}: {candidate.interview_ai_score}%")
            
            # Send notification if configured
            if hasattr(globals(), 'notify_scoring_complete'):
                notify_scoring_complete(candidate)
            
        except Exception as e:
            logger.error(f"Auto-scoring failed for candidate {candidate_id}: {e}", exc_info=True)
            if candidate:
                candidate.interview_ai_analysis_status = 'failed'
                session.commit()
        finally:
            session.close()
    
    # Run in background thread
    if 'executor' in globals():
        executor.submit(run_scoring)
    else:
        # Run directly if no executor
        run_scoring()


def trigger_ai_analysis(candidate_id):
    """Trigger REAL AI analysis for completed interview - NO RANDOM SCORES"""
    def run_analysis():
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                logger.error(f"Candidate {candidate_id} not found")
                return
            
            # Update status
            candidate.interview_ai_analysis_status = 'processing'
            session.commit()
            
            # Parse Q&A data
            questions = json.loads(candidate.interview_questions_asked or '[]')
            answers = json.loads(candidate.interview_answers_given or '[]')
            
            # Build Q&A pairs for analysis
            qa_pairs = []
            for i, question in enumerate(questions):
                qa_pair = {
                    'question': question.get('text', '') if isinstance(question, dict) else str(question),
                    'answer': ''
                }
                if i < len(answers):
                    answer = answers[i]
                    qa_pair['answer'] = answer.get('text', '') if isinstance(answer, dict) else str(answer)
                qa_pairs.append(qa_pair)
            
            # CHECK FOR INVALID/TEST RESPONSES
            has_invalid_responses = False
            invalid_patterns = ['INIT_INTERVIEW', 'TEST_RESPONSE', 'undefined', 'null']
            
            for qa in qa_pairs:
                answer_text = qa.get('answer', '').strip()
                if any(pattern in answer_text for pattern in invalid_patterns) or len(answer_text) < 5:
                    has_invalid_responses = True
                    logger.warning(f"Invalid response detected: {answer_text[:50]}")
                    break
            
            # DYNAMIC SCORING BASED ON ACTUAL CONTENT
            if has_invalid_responses or len(qa_pairs) == 0:
                # Failed interview - invalid or no responses
                candidate.interview_ai_score = 0
                candidate.interview_ai_technical_score = 0
                candidate.interview_ai_communication_score = 0
                candidate.interview_ai_problem_solving_score = 0
                candidate.interview_ai_cultural_fit_score = 0
                candidate.interview_ai_overall_feedback = """
Interview Analysis: FAILED
- No valid responses provided
- Interview appears to be incomplete or contains test data
- Candidate did not properly complete the interview

Recommendation: Schedule a new interview
"""
                candidate.interview_final_status = 'Failed - Invalid Response'
                
            else:
                # ACTUAL DYNAMIC ANALYSIS
                scores = analyze_interview_content(qa_pairs, candidate)
                
                candidate.interview_ai_score = scores['overall']
                candidate.interview_ai_technical_score = scores['technical']
                candidate.interview_ai_communication_score = scores['communication']
                candidate.interview_ai_problem_solving_score = scores['problem_solving']
                candidate.interview_ai_cultural_fit_score = scores['cultural_fit']
                candidate.interview_ai_overall_feedback = scores['feedback']
                
                # Set final status based on ACTUAL score
                if scores['overall'] >= 70:
                    candidate.interview_final_status = 'Recommended'
                elif scores['overall'] >= 50:
                    candidate.interview_final_status = 'Review Required'
                else:
                    candidate.interview_final_status = 'Not Recommended'
            
            candidate.interview_ai_analysis_status = 'completed'
            session.commit()
            
            logger.info(f"Dynamic analysis completed for {candidate.name}: {candidate.interview_ai_score}%")
            
            # Clear cache
            cache.delete_memoized(get_cached_candidates)
            
        except Exception as e:
            logger.error(f"AI analysis failed for candidate {candidate_id}: {e}", exc_info=True)
            if candidate:
                candidate.interview_ai_analysis_status = 'failed'
                candidate.interview_ai_overall_feedback = f"Analysis failed: {str(e)}"
                session.commit()
        finally:
            session.close()
    
    # Run in background thread
    executor.submit(run_analysis)


def analyze_interview_content(qa_pairs, candidate):
    """REAL content analysis - not random!"""
    
    # Initialize scores
    scores = {
        'technical': 0,
        'communication': 0,
        'problem_solving': 0,
        'cultural_fit': 0,
        'overall': 0
    }
    
    if not qa_pairs:
        return scores
    
    # Technical keywords to look for
    tech_keywords = [
        'implement', 'develop', 'design', 'architecture', 'algorithm',
        'database', 'api', 'framework', 'optimize', 'scale', 'debug',
        'testing', 'deployment', 'code', 'programming', 'software'
    ]
    
    soft_keywords = [
        'team', 'collaborate', 'communicate', 'lead', 'manage',
        'problem', 'solution', 'challenge', 'learn', 'adapt'
    ]
    
    total_score = 0
    answered_count = 0
    total_word_count = 0
    technical_mentions = 0
    soft_skill_mentions = 0
    
    for qa in qa_pairs:
        answer = qa.get('answer', '').lower()
        question = qa.get('question', '').lower()
        
        if not answer or len(answer) < 5:
            continue
            
        answered_count += 1
        words = answer.split()
        word_count = len(words)
        total_word_count += word_count
        
        # Score based on answer length (longer = more detailed)
        length_score = min(30, word_count / 3)  # Max 30 points for length
        
        # Check for technical keywords
        tech_found = sum(1 for kw in tech_keywords if kw in answer)
        technical_mentions += tech_found
        tech_score = min(40, tech_found * 10)  # Max 40 points for technical content
        
        # Check for soft skills
        soft_found = sum(1 for kw in soft_keywords if kw in answer)
        soft_skill_mentions += soft_found
        soft_score = min(30, soft_found * 10)  # Max 30 points for soft skills
        
        # Calculate answer score
        answer_score = length_score + tech_score + soft_score
        
        # Categorize score
        if 'technical' in question or 'code' in question or 'implement' in question:
            scores['technical'] += answer_score
        elif 'team' in question or 'collaborate' in question:
            scores['cultural_fit'] += answer_score
        elif 'problem' in question or 'challenge' in question:
            scores['problem_solving'] += answer_score
        else:
            scores['communication'] += answer_score
        
        total_score += answer_score
    
    # Calculate completion rate
    completion_rate = (answered_count / len(qa_pairs)) * 100 if qa_pairs else 0
    
    # Average word count per answer
    avg_word_count = total_word_count / answered_count if answered_count > 0 else 0
    
    # Normalize scores (0-100 scale)
    num_questions = len(qa_pairs)
    if num_questions > 0:
        for key in ['technical', 'communication', 'problem_solving', 'cultural_fit']:
            scores[key] = min(100, (scores[key] / num_questions) * 2)
    
    # Calculate overall score with penalties
    base_score = (scores['technical'] * 0.35 + 
                  scores['communication'] * 0.25 + 
                  scores['problem_solving'] * 0.25 + 
                  scores['cultural_fit'] * 0.15)
    
    # Apply penalties
    if completion_rate < 50:
        base_score *= 0.5  # 50% penalty for low completion
    elif completion_rate < 80:
        base_score *= 0.8  # 20% penalty
    
    if avg_word_count < 20:
        base_score *= 0.7  # 30% penalty for very brief answers
    
    scores['overall'] = min(100, max(0, base_score))
    
    # Generate detailed feedback
    scores['feedback'] = f"""
INTERVIEW ANALYSIS REPORT (Dynamic Scoring)
==========================================
Candidate: {candidate.name}
Position: {candidate.job_title}
Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}

SCORING METRICS:
- Completion Rate: {completion_rate:.1f}%
- Questions Answered: {answered_count}/{len(qa_pairs)}
- Average Response Length: {avg_word_count:.0f} words
- Technical Keywords Found: {technical_mentions}
- Soft Skill Keywords Found: {soft_skill_mentions}

DETAILED SCORES:
- Technical Skills: {scores['technical']:.1f}/100
- Communication: {scores['communication']:.1f}/100
- Problem Solving: {scores['problem_solving']:.1f}/100
- Cultural Fit: {scores['cultural_fit']:.1f}/100

OVERALL SCORE: {scores['overall']:.1f}/100

ASSESSMENT:
{get_assessment_text(scores['overall'], completion_rate, avg_word_count)}

RECOMMENDATION: {'Highly Recommended' if scores['overall'] >= 80 else 'Recommended' if scores['overall'] >= 70 else 'Consider for Further Review' if scores['overall'] >= 50 else 'Not Recommended'}
"""
    
    return scores


def get_assessment_text(overall_score, completion_rate, avg_word_count):
    """Generate assessment text based on metrics"""
    
    if overall_score >= 80:
        return """The candidate demonstrated excellent performance across all areas. 
Their responses were detailed, thoughtful, and showed strong technical knowledge 
combined with good communication skills. They are highly recommended for this position."""
    
    elif overall_score >= 70:
        return """The candidate showed good performance with solid responses to most questions. 
They demonstrated adequate technical knowledge and communication skills. 
Consider for the next round of interviews."""
    
    elif overall_score >= 50:
        return f"""The candidate showed moderate performance with some areas of concern.
Completion rate: {completion_rate:.0f}%
Average response detail: {avg_word_count:.0f} words
Additional evaluation may be needed to make a final decision."""
    
    else:
        return f"""The candidate's performance was below expectations.
Key concerns:
- Low completion rate: {completion_rate:.0f}%
- Brief responses: {avg_word_count:.0f} words average
- Limited demonstration of required skills
Not recommended for this position."""


# Also add this helper function if using OpenAI
def analyze_with_openai(qa_pairs, candidate):
    """Use OpenAI for more intelligent analysis"""
    try:
        import openai
        openai.api_key = os.getenv('OPENAI_API_KEY')
        
        if not openai.api_key:
            return None
        
        # Prepare Q&A text
        qa_text = "\n".join([
            f"Q: {qa['question']}\nA: {qa['answer']}" 
            for qa in qa_pairs if qa.get('answer')
        ])
        
        prompt = f"""
        Analyze this interview for {candidate.job_title} position.
        Score each aspect 0-100 based on ACTUAL content quality.
        
        Interview:
        {qa_text}
        
        Return JSON with: technical, communication, problem_solving, cultural_fit, overall scores and feedback.
        """
        
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert interview evaluator."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        
        # Parse response and return scores
        import json
        result = json.loads(response.choices[0].message.content)
        return result
        
    except Exception as e:
        logger.error(f"OpenAI analysis failed: {e}")
        return None

# Add to backend.py:
@app.route('/api/interview-results', methods=['GET'])
@cache.memoize(timeout=60)
def get_all_interview_results():
    """Get all interview results with proper filtering and pagination"""
    try:
        # Get query parameters
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 50))
        status_filter = request.args.get('status')
        date_from = request.args.get('date_from')
        date_to = request.args.get('date_to')
        
        session = SessionLocal()
        try:
            # Build query
            query = session.query(Candidate).filter(
                Candidate.interview_scheduled == True
            )
            
            # Apply filters
            if status_filter:
                if status_filter == 'completed':
                    query = query.filter(Candidate.interview_completed_at.isnot(None))
                elif status_filter == 'in_progress':
                    query = query.filter(
                        Candidate.interview_started_at.isnot(None),
                        Candidate.interview_completed_at.is_(None)
                    )
                elif status_filter == 'analyzed':
                    query = query.filter(
                        Candidate.interview_ai_analysis_status == AnalysisStatus.COMPLETED.value
                    )
            
            if date_from:
                query = query.filter(Candidate.interview_date >= date_from)
            if date_to:
                query = query.filter(Candidate.interview_date <= date_to)
            
            # Order by most recent
            query = query.order_by(Candidate.interview_date.desc())
            
            # Paginate
            total = query.count()
            candidates = query.offset((page - 1) * per_page).limit(per_page).all()
            
            # Format results
            results = []
            for candidate in candidates:
                # Safe JSON parsing
                def safe_json_parse(data, default):
                    try:
                        return json.loads(data) if data else default
                    except:
                        return default
                
                result = {
                    'id': candidate.id,
                    'name': candidate.name,
                    'email': candidate.email,
                    'phone': candidate.phone,
                    'job_title': candidate.job_title,
                    'resume_url': candidate.resume_path,
                    
                    # Interview details
                    'interview_date': candidate.interview_date.isoformat() if candidate.interview_date else None,
                    'interview_scheduled': candidate.interview_scheduled,
                    'interview_started_at': candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                    'interview_completed_at': candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                    'interview_duration': candidate.interview_duration or 0,
                    'interview_token': candidate.interview_token,
                    
                    # Progress
                    'interview_progress': candidate.interview_progress_percentage or 0,
                    'interview_questions_answered': candidate.interview_answered_questions or 0,
                    'interview_total_questions': candidate.interview_total_questions or 0,
                    
                    # Scores
                    'interview_ai_score': candidate.interview_ai_score,
                    'interview_ai_technical_score': candidate.interview_ai_technical_score,
                    'interview_ai_communication_score': candidate.interview_ai_communication_score,
                    'interview_ai_problem_solving_score': candidate.interview_ai_problem_solving_score,
                    'interview_ai_cultural_fit_score': candidate.interview_ai_cultural_fit_score,
                    
                    # Analysis
                    'interview_ai_analysis_status': candidate.interview_ai_analysis_status,
                    'interview_ai_overall_feedback': candidate.interview_ai_overall_feedback,
                    'interview_final_status': candidate.interview_final_status,
                    
                    # Insights
                    'strengths': safe_json_parse(candidate.interview_strengths, []),
                    'weaknesses': safe_json_parse(candidate.interview_weaknesses, []),
                    'recommendations': safe_json_parse(candidate.interview_recommendations, []),
                    
                    # Metadata
                    'interview_confidence_score': candidate.interview_confidence_score,
                    'interview_scoring_method': candidate.interview_scoring_method,
                    'interview_recording_url': candidate.interview_recording_url
                }
                
                results.append(result)
            
            # Return paginated response
            return jsonify({
                'results': results,
                'pagination': {
                    'page': page,
                    'per_page': per_page,
                    'total': total,
                    'pages': (total + per_page - 1) // per_page
                }
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error getting interview results: {e}", exc_info=True)
        return jsonify({"error": "Internal server error"}), 500

@app.route('/api/interview/trigger-analysis/<int:candidate_id>', methods=['POST'])
@rate_limit(max_calls=5, time_window=60)
def trigger_analysis_manually(candidate_id):
    """Manually trigger analysis for a candidate"""
    try:
        # Validate candidate exists
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404
            
            if not candidate.interview_completed_at:
                return jsonify({"error": "Interview not completed"}), 400
            
            # Reset flags to trigger analysis
            candidate.interview_auto_score_triggered = False
            candidate.interview_ai_analysis_status = None
            session.commit()
        finally:
            session.close()
        
        # Queue for analysis
        success = interview_analysis_service.analyze_single_interview(candidate_id)
        
        if success:
            return jsonify({
                "success": True,
                "message": "Analysis triggered successfully",
                "candidate_id": candidate_id
            }), 200
        else:
            return jsonify({
                "success": False,
                "message": "Failed to trigger analysis"
            }), 500
            
    except Exception as e:
        logger.error(f"Error triggering analysis: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview/validate-analysis/<int:candidate_id>', methods=['GET'])
def validate_analysis(candidate_id):
    """Validate that analysis is dynamic, not random"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Get Q&A data
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        
        # Check for invalid responses
        has_invalid = False
        invalid_answers = []
        
        for qa in qa_pairs:
            answer = qa.get('answer', '').strip()
            if answer in ['INIT_INTERVIEW', 'TEST', ''] or len(answer) < 5:
                has_invalid = True
                invalid_answers.append(answer)
        
        return jsonify({
            "candidate_id": candidate_id,
            "name": candidate.name,
            "has_valid_qa_data": len(qa_pairs) > 0 and not has_invalid,
            "total_questions": len(qa_pairs),
            "invalid_answers": invalid_answers,
            "current_score": candidate.interview_ai_score,
            "scoring_method": candidate.interview_scoring_method,
            "analysis_status": candidate.interview_ai_analysis_status,
            "should_fail": has_invalid,
            "expected_score_range": "0-30%" if has_invalid else "Based on actual content"
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/reanalyze/<int:candidate_id>', methods=['POST'])
def reanalyze_interview(candidate_id):
    """Force re-analysis with dynamic scoring"""
    try:
        from interview_analysis_service_production_fixed import dynamic_analyzer
        
        # Perform dynamic analysis
        result = dynamic_analyzer.analyze_interview(candidate_id)
        
        return jsonify({
            "success": True,
            "message": "Re-analysis completed",
            "new_score": result['overall_score'],
            "method": result['method'],
            "recommendation": result.get('recommendation')
        }), 200
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/api/interview/service-status', methods=['GET'])
def get_analysis_service_status():
    """Get analysis service health status"""
    try:
        stats = interview_analysis_service.get_service_stats()
        
        # Add database stats
        session = SessionLocal()
        try:
            db_stats = {
                'pending_analyses': session.query(Candidate).filter(
                    Candidate.interview_completed_at.isnot(None),
                    Candidate.interview_ai_analysis_status.is_(None)
                ).count(),
                'processing_analyses': session.query(Candidate).filter(
                    Candidate.interview_ai_analysis_status == AnalysisStatus.PROCESSING.value
                ).count(),
                'completed_today': session.query(Candidate).filter(
                    Candidate.interview_analysis_completed_at >= datetime.now().replace(hour=0, minute=0, second=0)
                ).count()
            }
            stats['database'] = db_stats
        finally:
            session.close()
        
        return jsonify({
            "status": "healthy" if stats['is_running'] else "stopped",
            "details": stats
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting service status: {e}")
        return jsonify({
            "status": "error",
            "error": str(e)
        }), 500

# Error handler for the service
@atexit.register
def cleanup_analysis_service():
    """Cleanup on shutdown"""
    try:
        logger.info("Stopping Interview Analysis Service...")
        interview_analysis_service.stop()
    except Exception as e:
        logger.error(f"Error stopping analysis service: {e}")

def analyze_interview_with_ai(qa_pairs, candidate):
    """Analyze interview Q&A with AI (or rule-based fallback)"""
    try:
        # Try AI analysis first (you can integrate OpenAI, Claude, etc.)
        if os.getenv('OPENAI_API_KEY'):
            return analyze_with_openai(qa_pairs, candidate)
        else:
            # Fallback to rule-based scoring
            return analyze_with_rules(qa_pairs, candidate)
    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        return analyze_with_rules(qa_pairs, candidate)

def analyze_with_rules(qa_pairs, candidate):
    """Rule-based interview scoring"""
    scores = {
        'technical': 0,
        'communication': 0,
        'problem_solving': 0,
        'cultural_fit': 0,
        'overall': 0,
        'confidence': 0.85,
        'feedback': ''
    }
    
    if not qa_pairs:
        return scores
    
    # Calculate completion rate
    answered = sum(1 for qa in qa_pairs if qa.get('answer'))
    total = len(qa_pairs)
    completion_rate = (answered / total * 100) if total > 0 else 0
    
    # Calculate average response time
    response_times = [qa.get('response_time', 0) for qa in qa_pairs if qa.get('response_time')]
    avg_response_time = sum(response_times) / len(response_times) if response_times else 0
    
    # Score based on completion
    if completion_rate >= 90:
        scores['communication'] = 85
    elif completion_rate >= 70:
        scores['communication'] = 70
    else:
        scores['communication'] = 50
    
    # Score based on response times
    if avg_response_time < 10:  # Quick responses
        scores['problem_solving'] = 80
    elif avg_response_time < 30:
        scores['problem_solving'] = 70
    else:
        scores['problem_solving'] = 60
    
    # Analyze answer lengths
    answer_lengths = [len(qa.get('answer', '')) for qa in qa_pairs if qa.get('answer')]
    avg_length = sum(answer_lengths) / len(answer_lengths) if answer_lengths else 0
    
    if avg_length > 100:  # Detailed answers
        scores['technical'] = 75
        scores['cultural_fit'] = 70
    else:
        scores['technical'] = 60
        scores['cultural_fit'] = 60
    
    # Calculate overall score
    scores['overall'] = (
        scores['technical'] * 0.35 +
        scores['communication'] * 0.25 +
        scores['problem_solving'] * 0.25 +
        scores['cultural_fit'] * 0.15
    )
    
    # Generate feedback
    scores['feedback'] = f"""
Interview Analysis Summary:
- Completion Rate: {completion_rate:.1f}%
- Questions Answered: {answered}/{total}
- Average Response Time: {avg_response_time:.1f} seconds
- Average Answer Length: {avg_length:.0f} characters

Strengths:
{'- Good completion rate' if completion_rate >= 80 else ''}
{'- Quick response times' if avg_response_time < 15 else ''}
{'- Detailed answers provided' if avg_length > 100 else ''}

Areas for Improvement:
{'- Complete all questions' if completion_rate < 100 else ''}
{'- Provide more detailed responses' if avg_length < 50 else ''}

Overall Assessment: {'Recommended' if scores['overall'] >= 70 else 'Needs Further Evaluation'}
"""
    
    return scores

@app.route('/api/interview/live-status/<int:candidate_id>', methods=['GET'])
def get_live_interview_status(candidate_id):
    """Get real-time interview status for frontend"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Parse Q&A data
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        
        # Calculate live stats
        status = {
            'candidate_id': candidate.id,
            'name': candidate.name,
            'interview_status': get_interview_status(candidate),
            'progress': candidate.interview_progress_percentage or 0,
            'link_clicked': candidate.interview_link_clicked,
            'started': bool(candidate.interview_started_at),
            'completed': bool(candidate.interview_completed_at),
            'total_questions': candidate.interview_total_questions or 0,
            'answered_questions': candidate.interview_answered_questions or 0,
            'unanswered': candidate.interview_total_questions - candidate.interview_answered_questions if candidate.interview_total_questions else 0,
            'duration': None,
            'ai_score': candidate.interview_ai_score,
            'analysis_status': candidate.interview_ai_analysis_status,
            'last_activity': candidate.interview_last_activity.isoformat() if candidate.interview_last_activity else None,
            'connection_quality': candidate.interview_connection_quality or 'unknown',
            'current_question': None,
            'is_active': False
        }
        
        # Calculate duration
        if candidate.interview_started_at:
            if candidate.interview_completed_at:
                duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
            else:
                duration = (datetime.now() - candidate.interview_started_at).total_seconds()
            status['duration'] = int(duration)
        
        # Check if currently active
        if candidate.interview_last_activity:
            time_since_activity = (datetime.now() - candidate.interview_last_activity).total_seconds()
            status['is_active'] = time_since_activity < 60  # Active if activity within last minute
        
        # Get current question being answered
        for qa in qa_pairs:
            if qa.get('question') and not qa.get('answer'):
                status['current_question'] = qa.get('question')
                break
        
        return jsonify(status), 200
        
    except Exception as e:
        logger.error(f"Error getting live status: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

def get_interview_status(candidate):
    """Determine current interview status"""
    if candidate.interview_ai_analysis_status == 'completed':
        if candidate.interview_ai_score >= 70:
            return 'passed'
        else:
            return 'failed'
    elif candidate.interview_ai_analysis_status == 'processing':
        return 'analyzing'
    elif candidate.interview_completed_at:
        return 'completed'
    elif candidate.interview_started_at:
        if candidate.interview_last_activity:
            time_since = (datetime.now() - candidate.interview_last_activity).total_seconds()
            if time_since < 300:  # Active within last 5 minutes
                return 'in_progress'
            else:
                return 'inactive'
        return 'in_progress'
    elif candidate.interview_link_clicked:
        return 'link_clicked'
    elif candidate.interview_scheduled:
        return 'scheduled'
    else:
        return 'not_started'

# In backend.py, update the complete_interview_auto function:

def complete_interview_auto(candidate_id):
    """Auto-complete interview and trigger scoring when all questions are answered"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate or candidate.interview_completed_at:
            return
        
        # Mark interview as completed
        candidate.interview_completed_at = datetime.now()
        candidate.interview_progress_percentage = 100
        
        # Calculate duration
        if candidate.interview_started_at:
            duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
            candidate.interview_duration = int(duration)
        
        # Set analysis status to pending
        candidate.interview_ai_analysis_status = 'pending'
        
        session.commit()
        
        # Trigger auto-scoring immediately
        trigger_auto_scoring(candidate_id)
        
        logger.info(f"Interview auto-completed and scoring triggered for candidate {candidate_id}")
        
    except Exception as e:
        logger.error(f"Error auto-completing interview: {e}")
        session.rollback()
    finally:
        session.close()

def schedule_auto_scoring(candidate_id, delay_minutes=45):
    """Schedule automatic scoring if interview doesn't complete normally"""
    def check_and_score():
        time.sleep(delay_minutes * 60)
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            
            # Check if interview is still incomplete
            if candidate and candidate.interview_started_at and not candidate.interview_completed_at:
                # Check for recent activity
                if candidate.interview_last_activity:
                    time_since = (datetime.now() - candidate.interview_last_activity).total_seconds()
                    if time_since > 600:  # No activity for 10 minutes
                        # Force complete and score
                        complete_interview_auto(candidate_id)
        except Exception as e:
            logger.error(f"Error in scheduled scoring: {e}")
        finally:
            session.close()
    
    # Schedule in background
    threading.Thread(target=check_and_score, daemon=True).start()

def send_realtime_update(candidate_id, data):
    """Send real-time updates to frontend via WebSocket or polling"""
    # Store update for polling
    update_key = f"interview_update_{candidate_id}"
    cache.set(update_key, json.dumps(data), timeout=60)

@app.route('/api/interview/poll-updates/<int:candidate_id>', methods=['GET'])
def poll_interview_updates(candidate_id):
    """Polling endpoint for real-time updates"""
    update_key = f"interview_update_{candidate_id}"
    update_data = cache.get(update_key)
    
    if update_data:
        return jsonify(json.loads(update_data)), 200
    else:
        return jsonify({"no_updates": True}), 204

def notify_scoring_complete(candidate):
    """Send notification when scoring is complete"""
    try:
        # Send email notification
        if candidate.email:
            subject = f"Interview Analysis Complete - {candidate.name}"
            body = f"""
            Interview analysis has been completed for {candidate.name}.
            
            Position: {candidate.job_title}
            Overall Score: {candidate.interview_ai_score:.1f}%
            Recommendation: {candidate.interview_final_status}
            
            View full results in the dashboard.
            """
            
            # Send to HR/Admin
            admin_email = os.getenv('ADMIN_EMAIL')
            if admin_email:
                send_email(admin_email, subject, body)
        
        logger.info(f"Scoring notification sent for candidate {candidate.id}")
        
    except Exception as e:
        logger.error(f"Error sending scoring notification: {e}")

def log_interview_activity(candidate_id, activity_type, data):
    """Log all interview activities for audit trail"""
    try:
        log_entry = {
            'candidate_id': candidate_id,
            'activity': activity_type,
            'timestamp': datetime.now().isoformat(),
            'data': data
        }
        
        # Store in logs directory
        log_dir = os.path.join('logs', 'interviews', str(candidate_id))
        os.makedirs(log_dir, exist_ok=True)
        
        log_file = os.path.join(log_dir, 'activity.jsonl')
        with open(log_file, 'a') as f:
            f.write(json.dumps(log_entry) + '\n')
            
    except Exception as e:
        logger.error(f"Error logging activity: {e}")


@app.route('/api/verify-interview-process/<int:candidate_id>', methods=['GET'])
def verify_interview_process(candidate_id):
    """Verify the interview link generation process for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Check assessment completion
        assessment_status = {
            "exam_completed": candidate.exam_completed,
            "exam_percentage": candidate.exam_percentage,
            "exam_score": candidate.exam_score,
            "exam_feedback": candidate.exam_feedback
        }
        
        # Check interview status
        interview_status = {
            "interview_scheduled": candidate.interview_scheduled,
            "interview_token": candidate.interview_token,
            "interview_link": candidate.interview_link,
            "knowledge_base_id": candidate.knowledge_base_id,
            "final_status": candidate.final_status
        }
        
        # Determine if interview should be scheduled
        should_schedule = (
            candidate.exam_completed and 
            candidate.exam_percentage >= 70 and 
            not candidate.interview_scheduled
        )
        
        # Generate test interview link if needed
        test_link = None
        if should_schedule and not candidate.interview_token:
            import uuid
            test_token = str(uuid.uuid4())
            test_link = f"{request.host_url.rstrip('/')}/secure-interview/{test_token}"
        
        return jsonify({
            "candidate_info": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "job_title": candidate.job_title
            },
            "assessment_status": assessment_status,
            "interview_status": interview_status,
            "should_schedule_interview": should_schedule,
            "test_interview_link": test_link,
            "process_ready": candidate.exam_completed and candidate.exam_percentage is not None
        }), 200
    finally:
        session.close()

@app.route('/api/interview/recording/start', methods=['POST', 'OPTIONS'])
def start_recording():
    if request.method == 'OPTIONS':
        return _ok_preflight()

    body = request.get_json(silent=True) or {}
    session_id = body.get('session_id')
    if not session_id:
        return jsonify({"success": False, "error": "session_id required"}), 400

    base = _ensure_dir(os.path.join('logs', 'interviews', session_id))
    _append_jsonl(os.path.join(base, 'session.jsonl'), {
        "event": "recording_started",
        "ts": datetime.now(timezone.utc).isoformat(),
        "meta": {
            "recording_format": body.get('recording_format'),
            "client_meta": body.get('client_meta'),
        }
    })

    # optional: flag on candidate
    try:
        db = SessionLocal()
        try:
            cand = db.query(Candidate).filter_by(interview_session_id=session_id).first()
            if cand:
                cand.recording_started_at = datetime.now(timezone.utc)
                db.commit()
        finally:
            db.close()
    except Exception:
        logger.exception("start_recording: candidate update failed")

    return jsonify({"success": True}), 200

@app.route("/api/interview/recording/upload", methods=["POST", "OPTIONS"])
def upload_recording():
    # CORS preflight
    if request.method == "OPTIONS":
        return _ok_preflight()

    # multipart/form-data expected
    f = request.files.get("recording")
    session_id = request.form.get("session_id") or request.args.get("session_id")

    if not f or not session_id:
        return jsonify({"success": False, "error": "missing recording file or session_id"}), 400

    # sanitize session_id for filesystem usage
    safe_session = secure_filename(str(session_id)) or f"session_{int(time.time())}"

    # Decide filename + ext
    ext = _ext_from_filename(getattr(f, "filename", None), default_ext="webm")
    base = _ensure_dir(os.path.join("logs", "interviews", safe_session))
    fname = f"interview_{safe_session}_{int(time.time())}.{ext}"
    path = os.path.join(base, fname)

    try:
        # Save file
        f.save(path)

        # Log event
        _append_jsonl(os.path.join(base, "session.jsonl"), {
            "event": "recording_uploaded",
            "ts": datetime.now(timezone.utc).isoformat(),
            "file": path,
            "size": os.path.getsize(path),
            "session_id": safe_session,
        })

        # Optional: update candidate record
        try:
            db = SessionLocal()
            try:
                cand = db.query(Candidate).filter_by(interview_session_id=session_id).first()
                if cand:
                    cand.recording_path = path
                    cand.interview_recording_format = ext
                    cand.interview_recording_status = "completed"
                    # only set completed_at if not already present
                    if not getattr(cand, "interview_completed_at", None):
                        cand.interview_completed_at = datetime.now(timezone.utc)
                    db.commit()
            except Exception as e:
                db.rollback()
                logger.exception(f"Failed to update candidate recording info: {e}")
            finally:
                db.close()
        except Exception as e:
            logger.exception(f"upload_recording: candidate update failed: {e}")

        return jsonify({"success": True, "path": path}), 200

    except Exception as e:
        logger.exception(f"upload_recording: saving failed: {e}")
        return jsonify({"success": False, "error": "failed to save recording"}), 500
 
@app.route('/api/interview/qa/track', methods=['POST', 'OPTIONS'])
def track_interview_qa_unified():
    """Unified Q&A tracking with real-time conversation capture and fallback mechanisms"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Get data from the incoming request
        data = request.json
        session_id = data.get('session_id')
        content_type = data.get('type')  # 'question' or 'answer'
        content = data.get('content', '').strip()
        metadata = data.get('metadata', {})
        
        # Validation: session_id and content are required
        if not session_id or not content:
            return jsonify({"error": "Missing required fields"}), 400
        
        session = SessionLocal()
        try:
            # Try to find the candidate using the session_id
            candidate = session.query(Candidate).filter_by(interview_session_id=session_id).first()

            if not candidate:
                # Fallback 1: Try parsing session_id (split by '_')
                if '_' in session_id:
                    parts = session_id.split('_')
                    if len(parts) >= 2 and parts[1].isdigit():
                        candidate = session.query(Candidate).filter_by(id=int(parts[1])).first()
                        if candidate:
                            # Update the session_id for the candidate
                            candidate.interview_session_id = session_id

                # Fallback 2: Try by interview token if session_id contains it
                if not candidate and 'token' in session_id:
                    token = session_id.split('token_')[-1]  # Assuming token is part of the session_id
                    candidate = session.query(Candidate).filter_by(interview_token=token).first()
            
            # If no candidate is found, return error
            if not candidate:
                logger.error(f"No candidate found for session_id={session_id}")
                return jsonify({"error": "Session not found"}), 404
            
            # Proceed to track the Q&A data for the candidate
            # Initialize or retrieve current interview data
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            conversation = json.loads(candidate.interview_conversation or '[]')
            transcript = candidate.interview_transcript or ""
            timestamp = datetime.now()

            if content_type == 'question':
                # This is a question from the knowledge base (KB)
                question_id = f"q_{len(qa_pairs) + 1}_{int(time.time())}"

                # Store the question data
                qa_entry = {
                    'id': question_id,
                    'question': content,
                    'answer': None,
                    'timestamp': timestamp.isoformat(),
                    'source': 'knowledge_base',
                    'is_transcribed': False,
                    'metadata': metadata
                }
                qa_pairs.append(qa_entry)

                # Add the question to the conversation
                conversation.append({
                    'type': 'question',
                    'speaker': 'Avatar',
                    'content': content,
                    'timestamp': timestamp.isoformat(),
                    'source': 'knowledge_base'
                })

                # Update transcript
                transcript += f"\n[{timestamp.strftime('%H:%M:%S')}] Avatar (KB): {content}\n"

                # Update the candidate's total question count
                candidate.interview_total_questions = len(qa_pairs)
                logger.info(f"Stored KB question #{len(qa_pairs)}: {content[:50]}...")

            elif content_type == 'answer':
                # This is a transcribed answer from the candidate
                # Find the most recent unanswered question
                for qa in reversed(qa_pairs):
                    if qa.get('question') and not qa.get('answer'):
                        qa['answer'] = content
                        qa['answer_timestamp'] = timestamp.isoformat()
                        qa['answer_source'] = 'voice_transcription'
                        qa['is_answer_transcribed'] = True
                        qa['answer_metadata'] = metadata
                        break

                # Add the answer to the conversation
                conversation.append({
                    'type': 'answer',
                    'speaker': 'Candidate',
                    'content': content,
                    'timestamp': timestamp.isoformat(),
                    'source': 'voice_transcription'
                })

                # Update transcript
                transcript += f"[{timestamp.strftime('%H:%M:%S')}] Candidate (Voice): {content}\n"

                # Update the answered question count
                answered = sum(1 for qa in qa_pairs if qa.get('answer'))
                candidate.interview_answered_questions = answered
                logger.info(f"Stored transcribed answer: {content[:50]}...")

            # Calculate interview progress
            if candidate.interview_total_questions > 0:
                progress = (candidate.interview_answered_questions / candidate.interview_total_questions) * 100
                candidate.interview_progress_percentage = progress

            # Save all updated data to the candidate record
            candidate.interview_qa_pairs = json.dumps(qa_pairs)
            candidate.interview_conversation = json.dumps(conversation)
            candidate.interview_transcript = transcript
            candidate.interview_last_activity = timestamp

            # Check for automatic interview completion if all questions are answered
            if (candidate.interview_total_questions >= 10 and 
                candidate.interview_answered_questions >= candidate.interview_total_questions):
                if not candidate.interview_completed_at:
                    candidate.interview_completed_at = datetime.now()
                    candidate.interview_ai_analysis_status = 'pending'
                    logger.info(f"Auto-completed interview for {candidate.name}")

            # Commit the changes to the database
            session.commit()

            return jsonify({
                "success": True,
                "stats": {
                    "total_questions": candidate.interview_total_questions,
                    "answered_questions": candidate.interview_answered_questions,
                    "progress": candidate.interview_progress_percentage
                }
            }), 200
            
        finally:
            session.close()

    except Exception as e:
        logger.error(f"Q&A tracking error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route('/api/interview/full-analysis/<token>', methods=['GET'])
def get_full_interview_analysis(token):
    """Get complete interview analysis data"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Parse Q&A data
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        questions_asked = json.loads(candidate.interview_questions_asked or '[]')
        answers_given = json.loads(candidate.interview_answers_given or '[]')
        
        return jsonify({
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "interview_status": {
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "duration_seconds": candidate.interview_duration,
                "duration_minutes": round(candidate.interview_duration / 60, 1) if candidate.interview_duration else None,
                "status": candidate.interview_status,
                "progress": candidate.interview_progress_percentage
            },
            "qa_statistics": {
                "total_qa_pairs": len(qa_pairs),
                "questions_asked": len(questions_asked),
                "answers_given": len(answers_given),
                "completion_rate": f"{(len(answers_given) / len(questions_asked) * 100) if questions_asked else 0:.1f}%"
            },
            "ai_analysis": {
                "analysis_status": candidate.interview_ai_analysis_status,
                "overall_score": candidate.interview_ai_score,
                "technical_score": candidate.interview_ai_technical_score,
                "communication_score": candidate.interview_ai_communication_score,
                "problem_solving_score": candidate.interview_ai_problem_solving_score,
                "cultural_fit_score": candidate.interview_ai_cultural_fit_score,
                "final_status": candidate.interview_final_status,
                "overall_feedback": candidate.interview_ai_overall_feedback
            },
            "recommendation": {
                "final_status": candidate.final_status,
                "interview_passed": candidate.interview_ai_score >= 70 if candidate.interview_ai_score else False
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting full analysis: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/qa/debug/<session_id>', methods=['GET'])
def debug_qa_tracking(session_id):
    """Debug endpoint to check Q&A tracking status"""
    
    session = SessionLocal()
    try:
        # Find candidate
        candidate = session.query(Candidate).filter_by(interview_session_id=session_id).first()
        
        if not candidate:
            # Try alternate lookups
            if '_' in session_id:
                parts = session_id.split('_')
                if len(parts) >= 2 and parts[1].isdigit():
                    candidate = session.query(Candidate).filter_by(id=int(parts[1])).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Parse Q&A data
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        qa_sequence = json.loads(candidate.interview_qa_sequence or '[]')
        conversation = json.loads(candidate.interview_conversation or '[]')
        
        # Analyze the data
        questions = [q for q in qa_pairs if q.get('question')]
        answered = [q for q in qa_pairs if q.get('answer')]
        unanswered = [q for q in qa_pairs if q.get('question') and not q.get('answer')]
        
        return jsonify({
            "candidate_id": candidate.id,
            "session_id": session_id,
            "status": candidate.interview_status,
            "qa_analysis": {
                "total_qa_pairs": len(qa_pairs),
                "questions_asked": len(questions),
                "questions_answered": len(answered),
                "unanswered_questions": len(unanswered),
                "conversation_length": len(conversation),
                "qa_sequence_length": len(qa_sequence)
            },
            "unanswered_questions": [
                {
                    "id": q.get('id'),
                    "question": q.get('question'),
                    "order": q.get('order')
                } for q in unanswered
            ],
            "completion_status": {
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "should_be_complete": len(questions) >= 10 and len(questions) == len(answered)
            }
        })
        
    finally:
        session.close()

@app.route('/api/interview/fix-status/<token>', methods=['POST'])
@cross_origin()
def fix_interview_status(token):
    """Fix interview status and completion tracking"""
    
    try:
        from models import Candidate, InterviewSession, InterviewQA, db
        
        candidate = Candidate.query.filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({'error': 'Invalid token'}), 404
        
        # Get current session
        session = InterviewSession.query.filter_by(
            candidate_id=candidate.id
        ).order_by(InterviewSession.created_at.desc()).first()
        
        # Count actual Q&A pairs with answers
        qa_count = InterviewQA.query.filter_by(
            candidate_id=candidate.id
        ).filter(InterviewQA.answer.isnot(None)).count()
        
        # Determine if interview should be complete
        # Based on your completion criteria
        should_complete = qa_count >= 10  # At least 10 answered questions
        
        if should_complete:
            # Mark as complete
            candidate.interview_completed_at = datetime.utcnow()
            candidate.interview_status = "Completed"
            
            if candidate.interview_started_at:
                duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                candidate.interview_duration = int(duration)
            
            if session:
                session.status = 'completed'
                session.completed_at = datetime.utcnow()
                session.progress = 100.0
            
            message = "Interview marked as completed"
        else:
            # Update to in_progress
            candidate.interview_status = "in_progress"
            
            if session:
                session.status = 'in_progress'
                # Calculate actual progress
                session.progress = (qa_count / 30) * 100  # Assuming 30 total questions
            
            message = f"Interview status updated to in_progress ({qa_count} questions answered)"
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': message,
            'candidate_id': candidate.id,
            'status': candidate.interview_status,
            'qa_count': qa_count,
            'progress': session.progress if session else 0
        }), 200
        
    except Exception as e:
        logger.error(f"Error fixing status: {str(e)}")
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/interview/qa/get-conversation/<session_id>', methods=['GET'])
def get_qa_conversation_unified(session_id):
    """Get the complete Q&A conversation from all tracking systems"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_session_id=session_id).first()
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Get data from all systems
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        qa_sequence = json.loads(getattr(candidate, 'interview_qa_sequence', '[]'))
        conversation = json.loads(getattr(candidate, 'interview_conversation', '[]'))
        
        # Format conversation from real-time conversation tracking (primary)
        formatted_conversation = []
        for entry in conversation:
            formatted_conversation.append({
                'index': entry.get('sequence', len(formatted_conversation) + 1),
                'speaker': entry['speaker'],
                'type': entry['type'],
                'content': entry['content'],
                'timestamp': entry['timestamp'],
                'metadata': entry.get('metadata', {})
            })
        
        # If conversation is empty, fallback to qa_sequence
        if not formatted_conversation and qa_sequence:
            for entry in qa_sequence:
                if entry.get('type') == 'question':
                    formatted_conversation.append({
                        'index': entry.get('sequence_number', len(formatted_conversation) + 1),
                        'speaker': 'Avatar',
                        'type': 'question',
                        'content': entry['content'],
                        'timestamp': entry.get('timestamp')
                    })
                    if entry.get('answered') and entry.get('answer'):
                        formatted_conversation.append({
                            'index': len(formatted_conversation) + 1,
                            'speaker': 'Candidate',
                            'type': 'answer',
                            'content': entry['answer'],
                            'timestamp': entry.get('answer_timestamp')
                        })
        
        # Create formatted text exactly as requested
        formatted_text = f"Interview Conversation - {candidate.name}\n"
        formatted_text += f"Position: {candidate.job_title}\n"
        formatted_text += f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        formatted_text += "="*70 + "\n\n"
        
        for entry in formatted_conversation:
            formatted_text += f"{entry['speaker']}: {entry['content']}\n\n"
        
        return jsonify({
            "success": True,
            "session_id": session_id,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "formatted_transcript": formatted_text,
            "structured_conversation": formatted_conversation,
            "raw_transcript": candidate.interview_transcript,
            "conversation": formatted_conversation,
            "stats": {
                "total_exchanges": len(formatted_conversation),
                "questions_asked": len([e for e in formatted_conversation if e['type'] == 'question']),
                "answers_given": len([e for e in formatted_conversation if e['type'] == 'answer']),
                "progress": candidate.interview_progress_percentage,
                "total_questions": candidate.interview_total_questions,
                "answered_questions": candidate.interview_answered_questions,
                "unanswered_questions": candidate.interview_total_questions - candidate.interview_answered_questions,
                "currently_waiting_for_answer": getattr(candidate, 'interview_waiting_for_answer', False)
            },
            "tracking_systems": {
                "qa_pairs_count": len(qa_pairs),
                "qa_sequence_count": len(qa_sequence),
                "conversation_count": len(conversation)
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting conversation: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

# Add this endpoint to your backend.py file

@app.route('/api/interview/debug-db/<token>', methods=['GET'])
def debug_interview_db(token):
    """Debug database state for an interview"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Get all interview-related fields
        interview_fields = {}
        for attr in dir(candidate):
            if attr.startswith('interview_') and not attr.startswith('_'):
                try:
                    value = getattr(candidate, attr)
                    if isinstance(value, datetime):
                        value = value.isoformat()
                    elif callable(value):
                        continue
                    interview_fields[attr] = value
                except:
                    interview_fields[attr] = "Error reading"
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "token": token,
            "interview_fields": interview_fields,
            "has_completed_at": candidate.interview_completed_at is not None,
            "final_status": candidate.final_status,
            "database_values": {
                "interview_completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "interview_started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "interview_status": getattr(candidate, 'interview_status', None),
                "interview_progress_percentage": getattr(candidate, 'interview_progress_percentage', None),
                "interview_duration": getattr(candidate, 'interview_duration', None),
                "final_status": candidate.final_status
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error in debug endpoint: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

# Add this simpler check endpoint
@app.route('/api/interview/check/<token>', methods=['GET'])
def check_interview_simple(token):
    """Simple check of interview status"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Force a fresh read from database
        session.refresh(candidate)
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "interview_completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
            "interview_started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
            "final_status": candidate.final_status,
            "fields_exist": {
                "has_completed_at_field": hasattr(candidate, 'interview_completed_at'),
                "has_status_field": hasattr(candidate, 'interview_status'),
                "has_duration_field": hasattr(candidate, 'interview_duration'),
            }
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/export-conversation/<session_id>', methods=['GET'])
def export_conversation_unified(session_id):
    """Export conversation in various formats"""
    format_type = request.args.get('format', 'text')
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_session_id=session_id).first()
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        conversation = json.loads(getattr(candidate, 'interview_conversation', '[]'))
        
        # If no conversation data, build from qa_pairs
        if not conversation:
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            for qa in qa_pairs:
                if qa.get('question'):
                    conversation.append({
                        'type': 'question',
                        'speaker': 'Avatar',
                        'content': qa['question'],
                        'timestamp': qa.get('timestamp')
                    })
                    if qa.get('answer'):
                        conversation.append({
                            'type': 'answer',
                            'speaker': 'Candidate',
                            'content': qa['answer'],
                            'timestamp': qa.get('answered_at')
                        })
        
        if format_type == 'text':
            # Plain text format
            output = f"Interview Transcript\n"
            output += f"Candidate: {candidate.name}\n"
            output += f"Position: {candidate.job_title}\n"
            output += f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            output += "="*50 + "\n\n"
            
            for entry in conversation:
                output += f"{entry.get('speaker', 'Unknown')}: {entry['content']}\n\n"
            
            response = Response(output, mimetype='text/plain')
            response.headers['Content-Disposition'] = f'attachment; filename=interview_{candidate.name}_{session_id}.txt'
            return response
            
        elif format_type == 'json':
            # JSON format with all metadata
            output = {
                "interview": {
                    "session_id": session_id,
                    "candidate": {
                        "id": candidate.id,
                        "name": candidate.name,
                        "email": candidate.email,
                        "position": candidate.job_title
                    },
                    "date": datetime.now().isoformat(),
                    "conversation": conversation,
                    "statistics": {
                        "total_exchanges": len(conversation),
                        "questions": len([e for e in conversation if e['type'] == 'question']),
                        "answers": len([e for e in conversation if e['type'] == 'answer']),
                        "completion": candidate.interview_progress_percentage
                    }
                }
            }
            
            response = Response(json.dumps(output, indent=2), mimetype='application/json')
            response.headers['Content-Disposition'] = f'attachment; filename=interview_{candidate.name}_{session_id}.json'
            return response
            
        elif format_type == 'html':
            # HTML format for viewing
            html = f"""
            <html>
            <head>
                <title>Interview Transcript - {candidate.name}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    .header {{ background: #f0f0f0; padding: 20px; margin-bottom: 30px; }}
                    .conversation {{ max-width: 800px; }}
                    .avatar {{ color: #2563eb; font-weight: bold; margin-top: 20px; }}
                    .candidate {{ color: #059669; font-weight: bold; margin-top: 20px; }}
                    .content {{ margin-left: 20px; margin-top: 5px; }}
                    .timestamp {{ color: #999; font-size: 0.8em; }}
                    .metadata {{ color: #666; font-size: 0.8em; font-style: italic; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>Interview Transcript</h1>
                    <p><strong>Candidate:</strong> {candidate.name}</p>
                    <p><strong>Position:</strong> {candidate.job_title}</p>
                    <p><strong>Date:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                    <p><strong>Progress:</strong> {candidate.interview_progress_percentage:.1f}%</p>
                </div>
                <div class="conversation">
            """
            
            for entry in conversation:
                speaker_class = 'avatar' if entry.get('speaker') == 'Avatar' else 'candidate'
                html += f'<div class="{speaker_class}">{entry.get("speaker", "Unknown")}:</div>'
                html += f'<div class="content">{entry["content"]}</div>'
                if entry.get('timestamp'):
                    html += f'<div class="timestamp">{entry["timestamp"]}</div>'
                if entry.get('metadata') and any(entry['metadata'].values()):
                    metadata_str = ', '.join([f"{k}: {v}" for k, v in entry['metadata'].items() if v])
                    html += f'<div class="metadata">{metadata_str}</div>'
            
            html += """
                </div>
            </body>
            </html>
            """
            
            response = Response(html, mimetype='text/html')
            response.headers['Content-Disposition'] = f'attachment; filename=interview_{candidate.name}_{session_id}.html'
            return response
            
    except Exception as e:
        logger.error(f"Error exporting conversation: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/fix-all-pending', methods=['POST'])
def fix_all_pending_interviews():
    """Fix all completed interviews that haven't been analyzed"""
    session = SessionLocal()
    fixed = 0
    
    try:
        # Find all completed interviews without scores
        pending = session.query(Candidate).filter(
            Candidate.interview_completed_at.isnot(None),
            Candidate.interview_ai_score.is_(None)
        ).all()
        
        for candidate in pending:
            # Trigger analysis
            candidate.interview_ai_analysis_status = 'pending'
            session.commit()
            
            # Trigger auto-scoring
            trigger_auto_scoring(candidate.id)
            fixed += 1
            
            logger.info(f"Triggered analysis for candidate {candidate.id} - {candidate.name}")
        
        # Clear cache
        cache.delete_memoized(get_cached_candidates)
        
        return jsonify({
            "success": True,
            "fixed": fixed,
            "message": f"Triggered analysis for {fixed} interviews"
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error fixing interviews: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/poll-updates', methods=['GET'])
def poll_for_updates():
    """Poll for interview analysis updates"""
    session = SessionLocal()
    try:
        # Find recently completed analyses
        cutoff = datetime.now() - timedelta(minutes=5)
        
        updated = session.query(Candidate).filter(
            Candidate.interview_analysis_completed_at >= cutoff,
            Candidate.interview_ai_score.isnot(None)
        ).all()
        
        updates = []
        for candidate in updated:
            updates.append({
                'candidate_id': candidate.id,
                'scores': {
                    'overall': candidate.interview_ai_score,
                    'technical': candidate.interview_ai_technical_score,
                    'communication': candidate.interview_ai_communication_score,
                    'problem_solving': candidate.interview_ai_problem_solving_score,
                    'cultural_fit': candidate.interview_ai_cultural_fit_score
                },
                'final_status': candidate.interview_final_status,
                'timestamp': candidate.interview_analysis_completed_at.isoformat()
            })
        
        return jsonify({
            'success': True,
            'updates': updates
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/qa/test/<session_id>', methods=['GET'])
def test_qa_tracking_unified(session_id):
    """Test endpoint to verify Q&A tracking is working correctly"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_session_id=session_id).first()
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Get raw data from all systems
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        qa_sequence = json.loads(getattr(candidate, 'interview_qa_sequence', '[]'))
        conversation = json.loads(getattr(candidate, 'interview_conversation', '[]'))
        
        # Analyze tracking health
        tracking_health = {
            "qa_pairs_healthy": all(qa.get('question') for qa in qa_pairs),
            "qa_sequence_healthy": all(entry.get('content') for entry in qa_sequence if entry.get('type') == 'question'),
            "conversation_healthy": all(entry.get('content') and entry.get('speaker') for entry in conversation),
            "answers_matched": sum(1 for qa in qa_pairs if qa.get('answer') is not None),
            "orphaned_answers": sum(1 for entry in qa_sequence if entry.get('type') == 'orphaned_answer'),
            "sync_status": {
                "qa_pairs_questions": len([q for q in qa_pairs if q.get('question')]),
                "qa_sequence_questions": len([e for e in qa_sequence if e.get('type') == 'question']),
                "conversation_questions": len([e for e in conversation if e.get('type') == 'question']),
                "all_synced": len(qa_pairs) == len([e for e in qa_sequence if e.get('type') == 'question']) == len([e for e in conversation if e.get('type') == 'question'])
            }
        }
        
        # Get last few entries for preview
        last_conversation = conversation[-5:] if conversation else []
        
        return jsonify({
            "session_id": session_id,
            "candidate": {
                "name": candidate.name,
                "position": candidate.job_title
            },
            "tracking_health": tracking_health,
            "last_conversation": last_conversation,
            "raw_data": {
                "qa_pairs": qa_pairs,
                "qa_sequence": qa_sequence,
                "conversation": conversation
            },
            "transcript_preview": candidate.interview_transcript[-500:] if candidate.interview_transcript else "No transcript",
            "stats": {
                "total_questions": candidate.interview_total_questions,
                "answered_questions": candidate.interview_answered_questions,
                "progress": candidate.interview_progress_percentage
            }
        }), 200
        
    finally:
        session.close()

@app.route('/api/resume-text/<int:candidate_id>', methods=['GET'])
def api_resume_text(candidate_id):
    session = SessionLocal()
    try:
        cand = session.query(Candidate).filter_by(id=candidate_id).first()
        if not cand:
            return jsonify({"error": "Candidate not found"}), 404

        text = ""
        if cand.resume_path and os.path.exists(cand.resume_path):
            # uses your existing extractor
            text = extract_resume_content(cand.resume_path)

        return jsonify({"resume_text": text, "length": len(text)}), 200
    finally:
        session.close()


@app.route('/api/interview/qa/verify/<session_id>', methods=['GET'])
def verify_qa_tracking_enhanced(session_id):
    """Enhanced verification of Q&A tracking with detailed analysis"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(
            interview_session_id=session_id
        ).first()
        
        if not candidate:
            # Try to find by token if session_id looks like it contains a token
            if 'token_' in session_id:
                token = session_id.split('token_')[-1]
                candidate = session.query(Candidate).filter_by(
                    interview_token=token
                ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Parse Q&A data
        try:
            questions = json.loads(candidate.interview_questions_asked or '[]')
        except:
            questions = []
            
        try:
            answers = json.loads(candidate.interview_answers_given or '[]')
        except:
            answers = []
        
        # Create detailed Q&A pairs analysis
        qa_pairs = []
        for i, question in enumerate(questions):
            # Find matching answer
            matching_answer = None
            for answer in answers:
                if answer.get('question_order') == i + 1:
                    matching_answer = answer
                    break
            
            # If no exact match, try to match by timing
            if not matching_answer and i < len(answers):
                matching_answer = answers[i]
            
            qa_pairs.append({
                'question_number': i + 1,
                'question': {
                    'text': question.get('text', ''),
                    'timestamp': question.get('timestamp', ''),
                    'id': question.get('id', '')
                },
                'answer': {
                    'text': matching_answer.get('text', '') if matching_answer else None,
                    'timestamp': matching_answer.get('timestamp', '') if matching_answer else None,
                    'id': matching_answer.get('id', '') if matching_answer else None
                } if matching_answer else None,
                'has_answer': matching_answer is not None,
                'time_to_answer': calculate_time_difference(
                    question.get('timestamp'),
                    matching_answer.get('timestamp')
                ) if matching_answer else None
            })
        
        # Calculate statistics
        total_questions = len(questions)
        total_answers = len(answers)
        answer_rate = (total_answers / total_questions * 100) if total_questions > 0 else 0
        
        # Get transcript preview
        transcript = candidate.interview_transcript or ''
        transcript_lines = transcript.strip().split('\n')[-10:]  # Last 10 lines
        
        return jsonify({
            "session_id": session_id,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "statistics": {
                "total_questions": total_questions,
                "total_answers": total_answers,
                "unanswered_questions": total_questions - total_answers,
                "answer_rate": f"{answer_rate:.1f}%",
                "transcript_length": len(transcript),
                "transcript_lines": len(transcript.strip().split('\n')) if transcript else 0
            },
            "qa_pairs": qa_pairs,
            "recent_transcript": transcript_lines,
            "session_info": {
                "interview_token": candidate.interview_token,
                "knowledge_base_id": candidate.knowledge_base_id,
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "last_accessed": candidate.last_accessed.isoformat() if hasattr(candidate, 'last_accessed') and candidate.last_accessed else None
            },
            "recording_status": getattr(candidate, 'interview_recording_status', 'unknown')
        }), 200
        
    except Exception as e:
        logger.error(f"Error in verify Q&A: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


def calculate_time_difference(timestamp1, timestamp2):
    """Calculate time difference between two timestamps in seconds"""
    if not timestamp1 or not timestamp2:
        return None
    try:
        t1 = datetime.fromisoformat(timestamp1.replace('Z', '+00:00'))
        t2 = datetime.fromisoformat(timestamp2.replace('Z', '+00:00'))
        return abs((t2 - t1).total_seconds())
    except:
        return None
    
@app.route('/api/interview/conversation/<int:candidate_id>', methods=['GET'])
def get_interview_conversation(candidate_id):
    """Get the formatted conversation for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Get formatted conversation
        conversation = candidate.interview_conversation or "No conversation recorded"
        
        return jsonify({
            "success": True,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "position": candidate.job_title
            },
            "conversation": conversation,
            "timestamp": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting conversation: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/export-conversation/<int:candidate_id>', methods=['GET'])
def export_conversation(candidate_id):
    """Export conversation in various formats"""
    format_type = request.args.get('format', 'text')  # text, json, pdf
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        if format_type == 'text':
            # Return plain text conversation
            conversation = candidate.interview_conversation or "No conversation recorded"
            
            response = Response(
                conversation,
                mimetype='text/plain',
                headers={
                    'Content-Disposition': f'attachment; filename=interview_{candidate.name}_{candidate.id}.txt'
                }
            )
            return response
            
        elif format_type == 'json':
            # Return structured JSON
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            
            export_data = {
                "candidate": {
                    "id": candidate.id,
                    "name": candidate.name,
                    "email": candidate.email,
                    "position": candidate.job_title
                },
                "interview_date": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "qa_pairs": qa_pairs,
                "formatted_conversation": candidate.interview_conversation
            }
            
            return jsonify(export_data), 200
            
        else:
            return jsonify({"error": "Invalid format"}), 400
            
    finally:
        session.close()

# Add this endpoint to check all interview data for a candidate
@app.route('/api/interview/candidate/<int:candidate_id>/full-data', methods=['GET'])
def get_candidate_interview_data(candidate_id):
    """Get complete interview data for a candidate including Q&A"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Parse all interview data
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        
        return jsonify({
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "interview": {
                "scheduled": candidate.interview_scheduled,
                "token": candidate.interview_token,
                "session_id": candidate.interview_session_id,
                "knowledge_base_id": candidate.knowledge_base_id,
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None
            },
            "qa_data": {
                "questions": questions,
                "answers": answers,
                "total_questions": len(questions),
                "total_answers": len(answers),
                "completion_rate": f"{(len(answers) / len(questions) * 100) if questions else 0:.1f}%"
            },
            "transcript": {
                "content": candidate.interview_transcript,
                "length": len(candidate.interview_transcript or ''),
                "lines": len((candidate.interview_transcript or '').strip().split('\n'))
            },
            "recording": {
                "status": candidate.interview_recording_status,
                "file": candidate.interview_recording_file,
                "duration": candidate.interview_recording_duration
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting candidate interview data: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

def upload_to_cloud_storage(local_path, filename):
    """Upload recording to cloud storage (implement based on your provider)"""
    try:
        # Example for S3
        if os.getenv('AWS_ACCESS_KEY_ID'):
            import boto3
            s3 = boto3.client('s3')
            bucket = os.getenv('S3_BUCKET_NAME', 'interview-recordings')
            key = f"interviews/{filename}"
            
            s3.upload_file(local_path, bucket, key)
            return f"https://{bucket}.s3.amazonaws.com/{key}"
       
        return None
        
    except Exception as e:
        logger.error(f"Cloud upload failed: {e}")
        return None


# Add this route to check recording and Q&A data
@app.route('/api/interview/session/data/<session_id>', methods=['GET'])
def get_interview_session_data(session_id):
    """Get complete interview session data including recording and Q&A"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(
            interview_session_id=session_id
        ).first()
        
        if not candidate:
            return jsonify({"error": "Session not found"}), 404
        
        # Parse Q&A data
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        
        # Create Q&A pairs
        qa_pairs = []
        for i, question in enumerate(questions):
            answer = next((a for a in answers if a.get('question_order') == i + 1), None)
            qa_pairs.append({
                'question': question,
                'answer': answer
            })
        
        return jsonify({
            "session_id": session_id,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "recording": {
                "status": candidate.interview_recording_status,
                "file": candidate.interview_recording_file,
                "url": candidate.interview_recording_url,
                "duration": candidate.interview_recording_duration,
                "size": candidate.interview_recording_size,
                "format": candidate.interview_recording_format
            },
            "qa_data": {
                "total_questions": candidate.interview_total_questions,
                "total_answers": candidate.interview_answered_questions,
                "qa_pairs": qa_pairs,
                "raw_questions": questions,
                "raw_answers": answers
            },
            "ai_analysis": {
                "status": candidate.interview_ai_analysis_status,
                "overall_score": candidate.interview_ai_score,
                "technical_score": candidate.interview_ai_technical_score,
                "communication_score": candidate.interview_ai_communication_score,
                "feedback": candidate.interview_ai_overall_feedback
            },
            "timestamps": {
                "started": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None
            }
        }), 200
        
    finally:
        session.close()

def _ok_preflight():
    resp = jsonify({})
    resp.status_code = 200
    resp.headers['Access-Control-Allow-Origin'] = '*'
    resp.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
    resp.headers['Access-Control-Allow-Methods'] = 'GET, POST, OPTIONS'
    return resp

def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path

def _append_jsonl(path, obj):
    with open(path, 'a', encoding='utf-8') as f:
        f.write(json.dumps(obj, ensure_ascii=False) + '\n')

# POST /api/interview/session/start
@app.route('/api/interview/session/start', methods=['POST', 'OPTIONS'])
def start_interview_session():
    if request.method == 'OPTIONS':
        return _ok_preflight()

    body = request.get_json(silent=True) or {}
    interview_token   = body.get('interview_token')
    candidate_id      = body.get('candidate_id')
    recording_config  = body.get('recording_config') or {}

    # make a session id if none provided
    session_id = body.get('session_id') or f"sess_{uuid.uuid4().hex[:8]}_{int(time.time())}"

    # persist basic linkage if we can find the candidate
    try:
        session = SessionLocal()
        try:
            candidate = None
            if candidate_id:
                candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            elif interview_token:
                candidate = session.query(Candidate).filter_by(interview_token=interview_token).first()

            if candidate:
                candidate.interview_session_id = session_id
                candidate.interview_started_at = datetime.now(timezone.utc)
                candidate.last_accessed = datetime.now(timezone.utc)
                # initialize counters if fields exist
                if not getattr(candidate, 'interview_total_questions', None):
                    candidate.interview_total_questions = 0
                if not getattr(candidate, 'interview_answered_questions', None):
                    candidate.interview_answered_questions = 0
                session.commit()
        finally:
            session.close()
    except Exception as e:
        logger.exception("start_interview_session: DB linkage failed")

    # create a place to log things to disk (optional but handy)
    base = _ensure_dir(os.path.join('logs', 'interviews', session_id))
    _append_jsonl(os.path.join(base, 'session.jsonl'), {
        "event": "session_started",
        "ts": datetime.now(timezone.utc).isoformat(),
        "session_id": session_id,
        "interview_token": interview_token,
        "candidate_id": candidate_id,
        "recording_config": recording_config,
    })

    return jsonify({
        "success": True,
        "session_id": session_id,
        "message": "Interview session initialized"
    }), 200


# 1. Add route to get candidate data by token
@app.route('/api/get-candidate-by-token/<token>', methods=['GET', 'OPTIONS'])
def get_candidate_by_token(token):
    """Get candidate data by interview token for KB creation"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        candidate_data = {
            "id": candidate.id,
            "name": candidate.name,
            "email": candidate.email,
            "job_title": candidate.job_title,
            "company": getattr(candidate, 'company_name', os.getenv('COMPANY_NAME', 'Our Company')),
            "resume_path": candidate.resume_path,
            "job_description": getattr(candidate, 'job_description', None),
            "ats_score": candidate.ats_score,
            "phone": getattr(candidate, 'phone', None)
        }
        
        return jsonify(candidate_data), 200
        
    except Exception as e:
        logger.error(f"Error getting candidate by token: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()


# 2. Add route to extract resume content
@app.route('/api/extract-resume/<int:candidate_id>', methods=['GET', 'OPTIONS'])
def extract_resume_content_api(candidate_id):
    """Extract resume content for a candidate"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        if not candidate.resume_path or not os.path.exists(candidate.resume_path):
            return jsonify({
                "content": "",
                "message": "No resume file available",
                "file_path": candidate.resume_path
            }), 200
        
        # Extract resume content
        resume_content = extract_resume_content(candidate.resume_path)
        
        # Extract additional metadata
        skills = extract_skills_from_resume(resume_content)
        experience = extract_experience_years(resume_content)
        projects = extract_projects_from_resume(resume_content)
        
        return jsonify({
            "content": resume_content,
            "skills": skills,
            "experience": experience,
            "projects": projects,
            "file_path": candidate.resume_path,
            "content_length": len(resume_content),
            "extraction_successful": len(resume_content) > 0
        }), 200
        
    except Exception as e:
        logger.error(f"Error extracting resume: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

# 3. Enhanced knowledge base storage
@app.route('/api/store-knowledge-base', methods=['POST', 'OPTIONS'])
def store_knowledge_base_enhanced():
    """Store knowledge base data with enhanced tracking"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        candidate_id = data.get('candidate_id')
        kb_id = data.get('knowledge_base_id')
        content = data.get('content', '')
        
        if not candidate_id or not kb_id:
            return jsonify({"error": "candidate_id and knowledge_base_id are required"}), 400
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404
            
            # Update candidate with knowledge base information
            candidate.knowledge_base_id = kb_id
            candidate.interview_kb_id = kb_id  # Also store in interview-specific field
            
            # Store knowledge base content if we have the field
            if hasattr(candidate, 'interview_kb_content'):
                candidate.interview_kb_content = content
            
            # Store metadata
            if hasattr(candidate, 'interview_kb_metadata'):
                metadata = {
                    'created_at': data.get('created_at', datetime.now().isoformat()),
                    'content_length': len(content),
                    'kb_type': 'heygen' if not kb_id.startswith('kb_') else 'fallback',
                    'version': '1.0'
                }
                candidate.interview_kb_metadata = json.dumps(metadata)
            
            # Update timestamps
            candidate.interview_created_at = datetime.now()
            if not candidate.interview_expires_at:
                candidate.interview_expires_at = datetime.now() + timedelta(days=7)
            
            session.commit()
            
            logger.info(f"Knowledge base {kb_id} stored for candidate {candidate.name}")
            
            return jsonify({
                "success": True,
                "message": "Knowledge base stored successfully",
                "candidate_id": candidate_id,
                "knowledge_base_id": kb_id
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error storing knowledge base: {e}")
        return jsonify({"error": str(e)}), 500


# 4. Enhanced resume extraction function
def extract_resume_content(resume_path):
    """Enhanced resume extraction with fixed f-string syntax"""
    try:
        if not os.path.exists(resume_path):
            logger.error(f"Resume file not found: {resume_path}")
            return ""
        
        file_ext = os.path.splitext(resume_path)[1].lower()
        logger.info(f"Extracting resume: {resume_path} (type: {file_ext})")
        
        resume_text = ""
        
        if file_ext == '.pdf':
            # Try multiple PDF extraction methods
            resume_text = extract_pdf_content(resume_path)
        elif file_ext in ['.docx', '.doc']:
            resume_text = extract_docx_content(resume_path)
        elif file_ext == '.txt':
            resume_text = extract_txt_content(resume_path)
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
        
        if resume_text and len(resume_text.strip()) > 50:
            logger.info(f"Resume extracted successfully: {len(resume_text)} characters")
            return resume_text.strip()
        else:
            logger.error("Resume extraction failed or produced minimal content")
            return ""
            
    except Exception as e:
        logger.error(f"Resume extraction error: {str(e)}", exc_info=True)
        return ""


def extract_pdf_content(pdf_path):
    """Extract content from PDF using multiple methods"""
    try:
        # Method 1: Try PyPDF2
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                if text.strip():
                    return text.strip()
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        # Method 2: Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # Method 3: Try pymupdf (fitz)
        try:
            import fitz  # pymupdf
            doc = fitz.open(pdf_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning(f"pymupdf extraction failed: {str(e)}")
            
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
    
    return ""


def extract_docx_content(docx_path):
    """Extract content from DOCX files"""
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Extract text from paragraphs
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text.strip()
        
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        return ""


def extract_txt_content(txt_path):
    """Extract content from text files"""
    try:
        # Try UTF-8 first
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Fallback to latin-1
            with open(txt_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
    except Exception as e:
        logger.error(f"TXT extraction failed: {str(e)}")
        return ""

# 5. Add validation endpoint for knowledge base creation
@app.route('/api/validate-kb-creation/<int:candidate_id>', methods=['GET', 'OPTIONS'])
def validate_kb_creation(candidate_id):
    """Validate that all components are ready for KB creation"""
    if request.method == 'OPTIONS':
        return '', 200
    
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Check all requirements
        checks = {
            "candidate_exists": True,
            "interview_token_exists": bool(candidate.interview_token),
            "resume_file_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "heygen_api_configured": bool(os.getenv('HEYGEN_API_KEY')),
            "job_title_available": bool(candidate.job_title),
            "company_name_available": bool(getattr(candidate, 'company_name', None) or os.getenv('COMPANY_NAME'))
        }
        
        # Try to extract resume content
        resume_content = ""
        if checks["resume_file_exists"]:
            resume_content = extract_resume_content_enhanced(candidate.resume_path)
            checks["resume_extractable"] = len(resume_content) > 0
        else:
            checks["resume_extractable"] = False
        
        # Check if KB already exists
        checks["kb_already_exists"] = bool(candidate.knowledge_base_id)
        
        # Overall readiness
        critical_checks = ["candidate_exists", "interview_token_exists", "heygen_api_configured", "job_title_available"]
        all_critical_passed = all(checks[check] for check in critical_checks)
        
        return jsonify({
            "candidate_id": candidate_id,
            "candidate_name": candidate.name,
            "ready_for_kb_creation": all_critical_passed,
            "checks": checks,
            "resume_content_length": len(resume_content),
            "existing_kb_id": candidate.knowledge_base_id,
            "recommendations": generate_kb_recommendations(checks)
        }), 200
        
    finally:
        session.close()


def generate_kb_recommendations(checks):
    """Generate recommendations based on validation checks"""
    recommendations = []
    
    if not checks.get("heygen_api_configured"):
        recommendations.append("Set HEYGEN_API_KEY environment variable")
    
    if not checks.get("resume_file_exists"):
        recommendations.append("Upload candidate resume file")
    elif not checks.get("resume_extractable"):
        recommendations.append("Resume file exists but content extraction failed - check file format")
    
    if not checks.get("interview_token_exists"):
        recommendations.append("Generate interview token for candidate")
    
    if not checks.get("company_name_available"):
        recommendations.append("Set COMPANY_NAME environment variable or add company to candidate record")
    
    if checks.get("kb_already_exists"):
        recommendations.append("Knowledge base already exists - consider updating instead of recreating")
    
    return recommendations


# 6. Add test endpoint for KB creation
@app.route('/api/test-kb-creation/<int:candidate_id>', methods=['POST', 'OPTIONS'])
def test_kb_creation(candidate_id):
    """Test knowledge base creation for debugging"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Step 1: Validate candidate
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404
            
            test_results = {
                "candidate_info": {
                    "id": candidate.id,
                    "name": candidate.name,
                    "email": candidate.email,
                    "job_title": candidate.job_title
                }
            }
            
            # Step 2: Test resume extraction
            resume_content = ""
            if candidate.resume_path:
                resume_content = extract_resume_content_enhanced(candidate.resume_path)
                test_results["resume_extraction"] = {
                    "success": len(resume_content) > 0,
                    "content_length": len(resume_content),
                    "file_path": candidate.resume_path,
                    "preview": resume_content[:200] + "..." if resume_content else "No content"
                }
            
            # Step 3: Test HeyGen API connection
            heygen_key = os.getenv('HEYGEN_API_KEY')
            if heygen_key:
                try:
                    # Test with a simple knowledge base
                    test_kb_content = f"Test knowledge base for {candidate.name}"
                    
                    test_response = requests.post(
                        'https://api.heygen.com/v1/streaming/knowledge_base',
                        headers={
                            'X-Api-Key': heygen_key,
                            'Content-Type': 'application/json'
                        },
                        json={
                            'name': f'Test_KB_{candidate.id}_{int(time.time())}',
                            'description': 'Test knowledge base',
                            'content': test_kb_content
                        },
                        timeout=30
                    )
                    
                    test_results["heygen_api_test"] = {
                        "success": test_response.ok,
                        "status_code": test_response.status_code,
                        "response_preview": test_response.text[:500]
                    }
                    
                    if test_response.ok:
                        kb_data = test_response.json()
                        test_kb_id = kb_data.get('data', {}).get('knowledge_base_id')
                        if test_kb_id:
                            test_results["heygen_api_test"]["kb_id_created"] = test_kb_id
                
                except Exception as e:
                    test_results["heygen_api_test"] = {
                        "success": False,
                        "error": str(e)
                    }
            else:
                test_results["heygen_api_test"] = {
                    "success": False,
                    "error": "HEYGEN_API_KEY not configured"
                }
            
            return jsonify({
                "success": True,
                "test_results": test_results,
                "ready_for_production": (
                    test_results.get("heygen_api_test", {}).get("success", False) and
                    test_results.get("resume_extraction", {}).get("success", False)
                )
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"KB creation test failed: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/interview/question/add', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=50, time_window=60)
def add_interview_question():
    """Add a question asked by the avatar during interview"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        question_data = data.get('question_data', {})
        
        if not session_id or not question_data.get('text'):
            return jsonify({"success": False, "message": "session_id and question_data.text are required"}), 400
        
        # Import session manager
        from interview_session_manager import interview_session_manager
        
        # Add question to session
        question_id = interview_session_manager.add_interview_question(session_id, question_data)
        
        if question_id:
            return jsonify({
                "success": True,
                "question_id": question_id,
                "message": "Question added successfully"
            }), 200
        else:
            return jsonify({"success": False, "message": "Failed to add question"}), 500
            
    except Exception as e:
        logger.error(f"Error adding interview question: {e}")
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/interview/answer/add', methods=['POST', 'OPTIONS'])
@rate_limit(max_calls=50, time_window=60)
def add_interview_answer():
    """Add a candidate's answer during interview"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        question_id = data.get('question_id')
        answer_data = data.get('answer_data', {})
        
        if not session_id or not question_id or not answer_data.get('text'):
            return jsonify({"success": False, "message": "session_id, question_id, and answer_data.text are required"}), 400
        
        # Import session manager
        from interview_session_manager import interview_session_manager
        
        # Add answer to session
        answer_id = interview_session_manager.add_interview_answer(session_id, question_id, answer_data)
        
        if answer_id:
            return jsonify({
                "success": True,
                "answer_id": answer_id,
                "message": "Answer added successfully"
            }), 200
        else:
            return jsonify({"success": False, "message": "Failed to add answer"}), 500
            
    except Exception as e:
        logger.error(f"Error adding interview answer: {e}")
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/api/interview/session/end', methods=['POST', 'OPTIONS'])
def end_interview_session():
    """End the interview session"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        
        if not session_id:
            return jsonify({"error": "session_id is required"}), 400
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(
                interview_session_id=session_id
            ).first()
            
            if not candidate:
                return jsonify({"error": "Session not found"}), 404
            
            # Update session status
            candidate.interview_completed_at = datetime.now()
            candidate.interview_status = 'completed'
            
            # Calculate duration
            if candidate.interview_started_at:
                duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                candidate.interview_duration = int(duration)
            
            session.commit()
            
            # Clear caches
            cache.delete_memoized(get_cached_candidates)
            
            logger.info(f"Interview session ended: {session_id}")
            
            return jsonify({
                "success": True,
                "message": "Interview session ended successfully",
                "session_id": session_id,
                "duration": getattr(candidate, 'interview_duration', 0)
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error ending interview session: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview/analysis/<int:candidate_id>', methods=['GET'])
@cache.memoize(timeout=30)
def get_interview_analysis_production(candidate_id):
    """Get analysis results with caching"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Parse stored data
        strengths = []
        weaknesses = []
        recommendations = []
        
        try:
            if candidate.interview_strengths:
                strengths = json.loads(candidate.interview_strengths)
            if candidate.interview_weaknesses:
                weaknesses = json.loads(candidate.interview_weaknesses)
            if candidate.interview_recommendations:
                recommendations = json.loads(candidate.interview_recommendations)
        except json.JSONDecodeError:
            logger.warning(f"Failed to parse JSON data for candidate {candidate_id}")
        
        analysis_data = {
            "candidate_id": candidate_id,
            "status": candidate.interview_ai_analysis_status,
            "scores": {
                "overall": candidate.interview_ai_score,
                "technical": candidate.interview_ai_technical_score,
                "communication": candidate.interview_ai_communication_score,
                "problem_solving": candidate.interview_ai_problem_solving_score,
                "cultural_fit": candidate.interview_ai_cultural_fit_score
            },
            "insights": {
                "strengths": strengths,
                "weaknesses": weaknesses,
                "recommendations": recommendations
            },
            "feedback": candidate.interview_ai_overall_feedback,
            "confidence": candidate.interview_confidence_score,
            "method": candidate.interview_scoring_method,
            "final_status": candidate.interview_final_status,
            "analyzed_at": candidate.interview_analysis_completed_at.isoformat() if candidate.interview_analysis_completed_at else None
        }
        
        return jsonify(analysis_data), 200
        
    except Exception as e:
        logger.error(f"Error getting analysis: {e}")
        return jsonify({"error": "Internal server error"}), 500
    finally:
        session.close()

# Add this to your backend.py to properly create knowledge base from resume/job

@app.route('/api/create-interview-knowledge-base', methods=['POST', 'OPTIONS'])
def create_interview_knowledge_base():
    """Create HeyGen knowledge base from candidate's resume and job description"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        candidate_id = data.get('candidate_id')
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return jsonify({"error": "Candidate not found"}), 404
            
            # Extract resume content
            resume_content = ""
            if candidate.resume_path and os.path.exists(candidate.resume_path):
                resume_content = extract_resume_content(candidate.resume_path)
            
            # Get job description
            job_description = candidate.job_description or f"Position: {candidate.job_title}"
            
            # Create knowledge base content
            kb_content = f"""
            CANDIDATE INFORMATION:
            Name: {candidate.name}
            Email: {candidate.email}
            Position Applied: {candidate.job_title}
            
            RESUME CONTENT:
            {resume_content}
            
            JOB DESCRIPTION:
            {job_description}
            
            INTERVIEW INSTRUCTIONS:
            - Ask questions based on the candidate's experience mentioned in resume
            - Focus on skills required for {candidate.job_title}
            - Assess technical competence based on job requirements
            - Ask behavioral questions related to their past experiences
            """
            
            # Call HeyGen API to create knowledge base
            heygen_key = os.getenv('HEYGEN_API_KEY')
            if heygen_key:
                response = requests.post(
                    'https://api.heygen.com/v1/streaming/knowledge_base',
                    headers={
                        'X-Api-Key': heygen_key,
                        'Content-Type': 'application/json'
                    },
                    json={
                        'name': f"Interview_{candidate.name}_{candidate.job_title}",
                        'content': kb_content,
                        'custom_prompt': generate_custom_interview_prompt(candidate, resume_content, job_description)
                    }
                )
                
                if response.ok:
                    kb_data = response.json()
                    kb_id = kb_data['data']['knowledge_base_id']
                    
                    # Update candidate record
                    candidate.knowledge_base_id = kb_id
                    candidate.interview_kb_id = kb_id
                    session.commit()
                    
                    return jsonify({
                        "success": True,
                        "knowledge_base_id": kb_id
                    }), 200
            
            # Fallback if HeyGen unavailable
            fallback_kb_id = f"kb_{candidate_id}_{int(time.time())}"
            candidate.knowledge_base_id = fallback_kb_id
            session.commit()
            
            return jsonify({
                "success": True,
                "knowledge_base_id": fallback_kb_id,
                "fallback": True
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Error creating knowledge base: {e}")
        return jsonify({"error": str(e)}), 500


def generate_custom_interview_prompt(candidate, resume_content, job_description):
    """Generate custom interview prompt based on resume and job"""
    return f"""
    You are interviewing {candidate.name} for {candidate.job_title} position.
    
    Based on their resume, ask questions about:
    1. Their experience with technologies mentioned in their resume
    2. Projects they've worked on
    3. Challenges they've faced
    4. Their approach to problem-solving
    
    Based on the job requirements, assess:
    1. Technical skills required for the role
    2. Soft skills and communication
    3. Cultural fit
    4. Career goals alignment
    
    Keep the interview conversational and professional.
    Ask follow-up questions based on their responses.
    """


# 2. Enhanced resume extraction function with better error handling
def extract_resume_content(resume_path):
    """Extract text content from resume with better error handling"""
    try:
        if not os.path.exists(resume_path):
            logger.error(f"Resume file not found: {resume_path}")
            return ""
        
        file_ext = os.path.splitext(resume_path)[1].lower()
        logger.info(f"Extracting resume: {resume_path} (type: {file_ext})")
        
        resume_text = ""
        
        if file_ext == '.pdf':
            # Try PyPDF2 first
            try:
                import PyPDF2
                with open(resume_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        resume_text += page.extract_text() + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with PyPDF2: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
            
            # Try pdfplumber as fallback
            try:
                import pdfplumber
                with pdfplumber.open(resume_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            resume_text += page_text + "\n"
                
                if resume_text.strip():
                    logger.info(f"PDF extracted with pdfplumber: {len(resume_text)} chars")
                    return resume_text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
                    
        elif file_ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(resume_path)
                resume_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                logger.info(f"DOCX extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                
        elif file_ext == '.txt':
            try:
                with open(resume_path, 'r', encoding='utf-8') as file:
                    resume_text = file.read()
                logger.info(f"TXT extracted: {len(resume_text)} chars")
                return resume_text.strip()
            except UnicodeDecodeError:
                with open(resume_path, 'r', encoding='latin-1') as file:
                    resume_text = file.read()
                return resume_text.strip()
        
        # If we couldn't extract anything, return empty string
        if not resume_text:
            logger.error(f"Failed to extract any text from {resume_path}")
            
    except Exception as e:
        logger.error(f"Resume extraction failed: {e}", exc_info=True)
    
    return ""

# Add this fixed version to your backend.py

@app.route('/api/verify-kb-creation/<int:candidate_id>', methods=['GET'])
def verify_kb_creation(candidate_id):
    """Verify knowledge base creation for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Check if we can extract resume
        resume_content = ""
        if candidate.resume_path and os.path.exists(candidate.resume_path):
            resume_content = extract_resume_content(candidate.resume_path)
        
        # Safely get knowledge_base_id using getattr
        knowledge_base_id = getattr(candidate, 'knowledge_base_id', None)
        interview_scheduled = getattr(candidate, 'interview_scheduled', False)
        job_description_available = bool(getattr(candidate, 'job_description', None))
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "resume_path": candidate.resume_path,
            "resume_exists": bool(candidate.resume_path and os.path.exists(candidate.resume_path)),
            "resume_content_length": len(resume_content),
            "resume_preview": resume_content[:500] + "..." if resume_content else "No content",
            "knowledge_base_id": knowledge_base_id,
            "interview_scheduled": interview_scheduled,
            "job_description_available": job_description_available,
            "database_ready": True
        }), 200
    except Exception as e:
        logger.error(f"Error in verify_kb_creation: {e}", exc_info=True)
        return jsonify({
            "error": str(e),
            "message": "Database schema may need updating. Run migration script."
        }), 500
    finally:
        session.close()

@app.route('/api/force-create-kb/<int:candidate_id>', methods=['POST'])
def force_create_kb(candidate_id):
    """Force create knowledge base for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Extract resume
        resume_content = ""
        if candidate.resume_path and os.path.exists(candidate.resume_path):
            resume_content = extract_resume_content(candidate.resume_path)
        
        # Try to create HeyGen KB
        kb_id = create_heygen_knowledge_base(
            candidate_name=candidate.name,
            position=candidate.job_title or "Software Engineer",
            resume_content=resume_content,
            company=os.getenv('COMPANY_NAME', 'Our Company')
        )
        
        # Check if HeyGen succeeded
        method = "heygen"
        if not kb_id or kb_id.startswith('kb_fallback'):
            method = "fallback"
            if not kb_id:
                kb_id = f"kb_fallback_{candidate.id}_{int(time.time())}"
        
        # Update candidate
        candidate.interview_kb_id = kb_id
        session.commit()
        
        # Clear cache
        cache.delete_memoized(get_cached_candidates)
        
        return jsonify({
            "success": True,
            "knowledge_base_id": kb_id,
            "resume_extracted": len(resume_content) > 0,
            "method": method
        }), 200
        
    finally:
        session.close()

# Add these endpoints to your backend.py file

@app.route('/api/fix-missing-knowledge-bases', methods=['POST'])
def fix_missing_knowledge_bases():
    """Fix all scheduled interviews that are missing knowledge bases"""
    session = SessionLocal()
    fixed_count = 0
    heygen_count = 0
    fallback_count = 0
    errors = []
    
    try:
        # Find all candidates with scheduled interviews but no KB
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True,
            Candidate.interview_kb_id.is_(None)
        ).all()
        
        logger.info(f"Found {len(candidates)} candidates with missing knowledge bases")
        
        for candidate in candidates:
            try:
                # Extract resume content
                resume_content = ""
                if candidate.resume_path and os.path.exists(candidate.resume_path):
                    resume_content = extract_resume_content(candidate.resume_path)
                    logger.info(f"Extracted {len(resume_content)} chars from resume for {candidate.name}")
                
                # Try to create HeyGen KB
                kb_id = create_heygen_knowledge_base(
                    candidate_name=candidate.name,
                    position=candidate.job_title,
                    resume_content=resume_content,
                    company=os.getenv('COMPANY_NAME', 'Our Company')
                )
                
                if kb_id and not kb_id.startswith('kb_fallback'):
                    heygen_count += 1
                    logger.info(f"Created HeyGen KB for {candidate.name}: {kb_id}")
                else:
                    # Use fallback if HeyGen failed
                    kb_id = f"kb_fallback_{candidate.id}_{int(time.time())}"
                    fallback_count += 1
                    logger.warning(f"Using fallback KB for {candidate.name}")
                
                # Update candidate
                candidate.interview_kb_id = kb_id
                fixed_count += 1
                
            except Exception as e:
                error_msg = f"Failed to fix {candidate.name} (ID: {candidate.id}): {str(e)}"
                logger.error(error_msg, exc_info=True)
                errors.append(error_msg)
        
        # Commit all changes
        session.commit()
        
        # Clear cache
        cache.delete_memoized(get_cached_candidates)
        
        return jsonify({
            "success": True,
            "fixed_count": fixed_count,
            "heygen_created": heygen_count,
            "fallback_used": fallback_count,
            "total_found": len(candidates),
            "errors": errors,
            "message": f"Fixed {fixed_count} candidates ({heygen_count} HeyGen, {fallback_count} fallback)"
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Batch KB fix failed: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/debug/heygen-test-fixed', methods=['GET'])
def debug_heygen_test_fixed():
    """Test HeyGen API with corrected field names"""
    heygen_key = os.getenv('HEYGEN_API_KEY')
    
    if not heygen_key:
        return jsonify({"error": "HEYGEN_API_KEY not found"}), 400
    
    # Test payload with CORRECT field names
    test_payload = {
        'name': f'Test_KB_{int(time.time())}',
        'opening': 'Hello, this is a test interview. Please tell me about yourself.',  # FIXED
        'prompt': 'Test interview questions: 1. Tell me about yourself. 2. Why this role? 3. What are your strengths?'
    }
    
    try:
        response = requests.post(
            "https://api.heygen.com/v1/streaming/knowledge_base/create",
            headers={
                'X-Api-Key': heygen_key,
                'Content-Type': 'application/json'
            },
            json=test_payload,
            timeout=30
        )
        
        if response.ok:
            data = response.json()
            kb_id = (
                data.get('data', {}).get('knowledge_base_id') or
                data.get('data', {}).get('id') or
                data.get('knowledge_base_id') or
                data.get('id')
            )
            
            return jsonify({
                "success": True,
                "status_code": response.status_code,
                "response": data,
                "knowledge_base_id": kb_id,
                "message": "HeyGen KB created successfully!"
            }), 200
        else:
            return jsonify({
                "success": False,
                "status_code": response.status_code,
                "error": response.text
            }), 400
            
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/api/check-interview-issues', methods=['GET'])
def check_interview_issues():
    """Check for all interview-related issues"""
    session = SessionLocal()
    try:
        issues = {
            "missing_kb": [],
            "missing_token": [],
            "missing_resume": [],
            "expired_interviews": []
        }
        
        # Get all scheduled interviews
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        for candidate in candidates:
            candidate_info = {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "job_title": candidate.job_title
            }
            
            # Check for missing KB using interview_kb_id
            if not candidate.interview_kb_id:
                issues["missing_kb"].append(candidate_info)
            
            # Check for missing token
            if not candidate.interview_token:
                issues["missing_token"].append(candidate_info)
            
            # Check for missing resume
            if not candidate.resume_path or not os.path.exists(candidate.resume_path):
                issues["missing_resume"].append(candidate_info)
            
            # Check for expired interviews
            if candidate.interview_expires_at and candidate.interview_expires_at < datetime.now():
                issues["expired_interviews"].append(candidate_info)
        
        summary = {
            "total_scheduled_interviews": len(candidates),
            "issues_found": {
                "missing_knowledge_bases": len(issues["missing_kb"]),
                "missing_tokens": len(issues["missing_token"]),
                "missing_resumes": len(issues["missing_resume"]),
                "expired_interviews": len(issues["expired_interviews"])
            },
            "details": issues
        }
        
        return jsonify(summary), 200
        
    except Exception as e:
        logger.error(f"Error checking interview issues: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/debug/heygen-test', methods=['GET'])
def debug_heygen_test():
    """Test HeyGen API connection and knowledge base creation"""
    heygen_key = os.getenv('HEYGEN_API_KEY')
    
    # Check if API key exists
    if not heygen_key:
        return jsonify({
            "error": "HEYGEN_API_KEY not found in environment variables",
            "fix": "Add HEYGEN_API_KEY to your .env file"
        }), 400
    
    # Test payload
    test_payload = {
        'name': f'Test_KB_{int(time.time())}',
        'description': 'Test knowledge base creation',
        'content': 'Test interview questions: 1. Tell me about yourself. 2. Why this role?',
        'opening_line': 'Hello, this is a test interview.'
    }
    
    # Try different endpoints
    endpoints = [
        "https://api.heygen.com/v1/streaming/knowledge_base/create",
        "https://api.heygen.com/v1/streaming/knowledge_base",
        "https://api.heygen.com/v1/streaming_avatar/knowledge_base",
        "https://api.heygen.com/v1/knowledge_base"
    ]
    
    results = []
    
    for endpoint in endpoints:
        try:
            response = requests.post(
                endpoint,
                headers={
                    'X-Api-Key': heygen_key,
                    'Content-Type': 'application/json',
                    'Accept': 'application/json'
                },
                json=test_payload,
                timeout=30
            )
            
            result = {
                "endpoint": endpoint,
                "status_code": response.status_code,
                "success": response.ok,
                "response": response.text[:500] if not response.ok else response.json()
            }
            results.append(result)
            
            if response.ok:
                # Try to extract KB ID
                data = response.json()
                kb_id = (
                    data.get('data', {}).get('knowledge_base_id') or
                    data.get('knowledge_base_id') or
                    data.get('id') or
                    data.get('data', {}).get('id')
                )
                if kb_id:
                    result["knowledge_base_id"] = kb_id
                    
        except Exception as e:
            results.append({
                "endpoint": endpoint,
                "error": str(e)
            })
    
    return jsonify({
        "api_key_length": len(heygen_key),
        "api_key_preview": f"{heygen_key[:10]}...{heygen_key[-5:]}",
        "test_results": results
    }), 200

def create_heygen_knowledge_base(candidate_name, position, resume_content, company):
    """Create HeyGen knowledge base with correct field names"""
    
    heygen_key = os.getenv('HEYGEN_API_KEY')
    if not heygen_key:
        logger.error("HEYGEN_API_KEY not set!")
        return None
    
    # Extract skills for better questions
    skills = extract_skills_from_resume(resume_content) if resume_content else []
    
    # Create a more HeyGen-friendly prompt format
    heygen_prompt = f"""You are an AI interviewer conducting a professional technical interview.

IMPORTANT: You must ask these exact questions in order when conducting the interview.

Candidate: {candidate_name}
Position: {position}
Company: {company}

When the interview starts or when you hear "INIT_INTERVIEW", immediately greet the candidate and ask Question 1.

INTERVIEW QUESTIONS TO ASK IN ORDER:

Question 1: "Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background?"

Question 2: {"I see from your resume that you have experience with " + skills[0] + ". Can you tell me about a specific project where you used this technology?" if skills else "Can you tell me about your most significant technical project and the technologies you used?"}

Question 3: "Can you describe a time when you encountered a complex technical problem and walk me through how you approached and solved it?"

Question 4: "Tell me about a time when you had to collaborate with team members on a challenging project. How did you handle any conflicts?"

Question 5: "What would you say is your strongest technical skill, and can you give me a detailed example of how you've applied it?"

Question 6: "Technology evolves rapidly. Can you tell me about a time when you had to quickly learn a new technology for a project?"

Question 7: "Based on your understanding of this {position} role, how do you see your skills contributing to our team?"

Question 8: "What do you think would be the biggest challenge for you in this role?"

Question 9: "Where do you see your career heading in the next 3-5 years?"

Question 10: "Do you have any questions for me about the role or {company}?"

INSTRUCTIONS:
- Ask ONE question at a time
- Wait for complete answers before proceeding
- Be professional and encouraging
- If the candidate seems stuck, offer to rephrase the question
- After the last question, thank them for their time"""
    
    # HeyGen payload with CORRECT field names
    payload = {
        "name": f"Interview_{candidate_name.replace(' ', '_')}_{int(time.time())}",
        "opening": f"Hello {candidate_name}, welcome to your interview for the {position} position at {company}. Let's begin. Could you please introduce yourself and tell me about your professional background?",
        "prompt": heygen_prompt
    }
    
    try:
        logger.info(f"Creating KB for {candidate_name} with payload keys: {payload.keys()}")
        
        response = requests.post(
            "https://api.heygen.com/v1/streaming/knowledge_base/create",
            headers={
                "X-Api-Key": heygen_key,
                "Content-Type": "application/json"
            },
            json=payload,
            timeout=45
        )
        
        logger.info(f"HeyGen response: {response.status_code}")
        
        if response.ok:
            data = response.json()
            logger.info(f"HeyGen success response: {json.dumps(data, indent=2)}")
            
            # Extract KB ID from response
            kb_id = None
            if isinstance(data, dict):
                kb_id = (
                    data.get('data', {}).get('knowledge_base_id') or
                    data.get('data', {}).get('id') or
                    data.get('knowledge_base_id') or
                    data.get('id')
                )
            
            if kb_id:
                logger.info(f"Successfully created HeyGen KB: {kb_id}")
                return kb_id
            else:
                logger.error(f"KB ID not found in response: {data}")
                
        else:
            error_text = response.text
            logger.error(f"HeyGen API error {response.status_code}: {error_text}")
            
    except Exception as e:
        logger.error(f"HeyGen API exception: {type(e).__name__}: {str(e)}")
    
    return None

# Add to backend.py

@app.route('/api/interview/realtime-analysis', methods=['GET'])
def get_realtime_analysis_status():
    """Get real-time analysis status for all pending interviews"""
    session = SessionLocal()
    try:
        pending = session.query(Candidate).filter(
            Candidate.interview_completed_at.isnot(None),
            Candidate.interview_ai_analysis_status.in_(['pending', 'processing'])
        ).all()
        
        results = []
        for candidate in pending:
            # Calculate progress
            progress = 0
            if candidate.interview_ai_analysis_status == 'processing':
                progress = 50
            
            results.append({
                'candidate_id': candidate.id,
                'name': candidate.name,
                'status': candidate.interview_ai_analysis_status,
                'progress': progress,
                'completed_at': candidate.interview_completed_at.isoformat()
            })
        
        return jsonify({
            'success': True,
            'pending_analyses': results
        }), 200
        
    finally:
        session.close()
        
# Add these endpoints to your backend.py file

@app.route('/api/interview/complete/<token>', methods=['POST', 'OPTIONS'])
@cross_origin()
def complete_interview_by_token(token):
    """Complete an interview with automatic database update"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        # Get request data
        data = request.json or {}
        trigger_source = data.get('completion_trigger', 'manual')
        
        # Use the completion handler
        result = completion_handler.complete_interview(token, trigger_source)
        
        if result["success"]:
            # Clear cache to update frontend
            cache.delete_memoized(get_cached_candidates)
            
            return jsonify({
                "success": True,
                "message": "Interview completed and saved to database",
                **result
            }), 200
        else:
            return jsonify({
                "success": False,
                "error": result.get("error", "Unknown error")
            }), 500
            
    except Exception as e:
        logger.error(f"Error in complete endpoint: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview/realtime-complete', methods=['POST'])
def realtime_complete():
    """Handle real-time completion from frontend"""
    try:
        data = request.json
        token = data.get('token')
        session_id = data.get('session_id')
        
        if not token:
            return jsonify({"error": "Token required"}), 400
        
        # Complete the interview
        result = completion_handler.complete_interview(token, "realtime_trigger")
        
        if result["success"]:
            # Also update session if provided
            if session_id:
                session = SessionLocal()
                try:
                    candidate = session.query(Candidate).filter_by(
                        interview_token=token
                    ).first()
                    if candidate:
                        candidate.interview_session_id = session_id
                        session.commit()
                finally:
                    session.close()
            
            return jsonify(result), 200
        else:
            return jsonify(result), 500
            
    except Exception as e:
        logger.error(f"Realtime completion error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview/verify-db-update/<token>', methods=['GET'])
def verify_db_update(token):
    """Verify that interview completion is saved in database"""
    session = SessionLocal()
    try:
        # Force fresh read from database
        session.expire_all()
        
        candidate = session.query(Candidate).filter_by(
            interview_token=token
        ).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Also check with direct SQL
        result = session.execute(
            text("""
                SELECT interview_completed_at, interview_status, final_status,
                       interview_progress_percentage, interview_ai_analysis_status
                FROM candidates WHERE interview_token = :token
            """),
            {"token": token}
        ).fetchone()
        
        return jsonify({
            "token": token,
            "candidate_name": candidate.name,
            "orm_data": {
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "status": candidate.interview_status,
                "final_status": candidate.final_status,
                "progress": candidate.interview_progress_percentage,
                "ai_status": candidate.interview_ai_analysis_status
            },
            "direct_sql_data": {
                "completed_at": str(result[0]) if result and result[0] else None,
                "status": result[1] if result else None,
                "final_status": result[2] if result else None,
                "progress": result[3] if result else None,
                "ai_status": result[4] if result else None
            },
            "is_completed": candidate.interview_completed_at is not None
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/verify-completion/<token>', methods=['GET'])
def verify_completion(token):
    """Verify if interview completion is persisted in database"""
    session = SessionLocal()
    try:
        # Try multiple ways to find the candidate
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            # Try finding by ID if token looks like an ID
            if token.isdigit():
                candidate = session.query(Candidate).filter_by(id=int(token)).first()
        
        if not candidate:
            return jsonify({"error": "Not found", "token": token}), 404
        
        # Force fresh read from database
        session.expire_all()
        session.commit()  # Commit any pending changes
        session.refresh(candidate)
        
        # Get raw SQL value
        from sqlalchemy import text
        raw_result = session.execute(
            text("SELECT interview_completed_at, interview_status, final_status FROM candidates WHERE id = :id"),
            {"id": candidate.id}
        ).fetchone()
        
        return jsonify({
            "candidate_id": candidate.id,
            "name": candidate.name,
            "orm_values": {
                "interview_completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "interview_status": candidate.interview_status,
                "final_status": candidate.final_status,
                "interview_progress_percentage": candidate.interview_progress_percentage
            },
            "raw_sql_values": {
                "interview_completed_at": str(raw_result[0]) if raw_result else None,
                "interview_status": str(raw_result[1]) if raw_result else None,
                "final_status": str(raw_result[2]) if raw_result else None
            },
            "is_completed": candidate.interview_completed_at is not None
        }), 200
        
    except Exception as e:
        logger.error(f"Verification error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/fix-null-completions', methods=['POST'])
def fix_null_completions():
    """Fix interviews that show as NULL in database but should be completed"""
    session = SessionLocal()
    try:
        from datetime import datetime, timezone
        from sqlalchemy import text
        
        # Find candidates with NULL completion but other signs of completion
        result = session.execute(
            text("""
                UPDATE candidates 
                SET interview_completed_at = CURRENT_TIMESTAMP,
                    interview_status = 'completed',
                    final_status = 'Interview Completed',
                    interview_progress_percentage = 100
                WHERE interview_completed_at IS NULL
                AND (
                    interview_progress_percentage >= 100
                    OR interview_status = 'completed'
                    OR final_status LIKE '%Completed%'
                    OR interview_total_questions > 0
                )
            """)
        )
        
        rows_updated = result.rowcount
        session.commit()
        
        # Clear all caches
        cache.clear()
        
        return jsonify({
            "success": True,
            "rows_fixed": rows_updated,
            "message": f"Fixed {rows_updated} interviews with NULL completion dates"
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Fix failed: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/check-incomplete', methods=['POST'])
def check_incomplete_interviews():
    """Check and fix incomplete interviews"""
    session = SessionLocal()
    fixed = 0
    
    try:
        # Find interviews that should be complete
        candidates = session.query(Candidate).filter(
            Candidate.interview_started_at.isnot(None),
            Candidate.interview_completed_at.is_(None)
        ).all()
        
        for candidate in candidates:
            # Get Q&A data
            questions = json.loads(candidate.interview_questions_asked or '[]')
            answers = json.loads(candidate.interview_answers_given or '[]')
            
            # Check if should be complete
            should_complete = False
            
            # Has enough Q&A
            if len(questions) >= 5 and len(answers) >= len(questions) * 0.8:
                should_complete = True
            
            # Or has been inactive for too long
            if candidate.interview_started_at:
                time_since = (datetime.now() - candidate.interview_started_at).total_seconds()
                if time_since > 7200:  # 2 hours
                    should_complete = True
            
            if should_complete:
                candidate.interview_completed_at = datetime.now()
                candidate.interview_status = 'completed'
                candidate.interview_progress_percentage = 100
                candidate.interview_ai_analysis_status = 'pending'
                fixed += 1
                logger.info(f"Fixed incomplete interview for {candidate.name}")
        
        session.commit()
        
        return jsonify({
            "success": True,
            "checked": len(candidates),
            "fixed": fixed
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error checking incomplete: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/check-completion/<token>', methods=['GET'])
def check_interview_completion(token):
    """Check if interview is properly marked as complete"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Get Q&A data
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        
        return jsonify({
            "token": token,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email
            },
            "completion_status": {
                "is_completed": candidate.interview_completed_at is not None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "duration": candidate.interview_duration,
                "progress": candidate.interview_progress_percentage,
                "status": candidate.interview_status,
                "final_status": candidate.final_status
            },
            "qa_stats": {
                "total_questions": len(questions),
                "total_answers": len(answers),
                "qa_pairs": len(qa_pairs),
                "stored_total": candidate.interview_total_questions,
                "stored_answered": candidate.interview_answered_questions
            },
            "ai_analysis": {
                "status": candidate.interview_ai_analysis_status,
                "score": candidate.interview_ai_score,
                "triggered": candidate.interview_auto_score_triggered
            }
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/fix-by-token/<token>', methods=['POST'])
def fix_interview_by_token(token):
    """Emergency fix for a specific interview"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Force complete if it has Q&A data but not marked complete
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        
        if qa_pairs and not candidate.interview_completed_at:
            candidate.interview_completed_at = datetime.now()
            candidate.interview_status = 'completed'
            candidate.interview_progress_percentage = 100
            candidate.final_status = 'Interview Completed'
            
            if candidate.interview_started_at:
                duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                candidate.interview_duration = int(duration)
            
            candidate.interview_ai_analysis_status = 'pending'
            session.commit()
            
            # Trigger scoring
            trigger_auto_scoring(candidate.id)
            
            return jsonify({
                "success": True,
                "message": "Interview fixed and marked as complete",
                "candidate": candidate.name
            }), 200
        else:
            return jsonify({
                "message": "No fix needed",
                "has_qa_data": len(qa_pairs) > 0,
                "already_completed": candidate.interview_completed_at is not None
            }), 200
            
    except Exception as e:
        session.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/tracking-status/<token>', methods=['GET'])
def get_interview_tracking_status(token):
    """Get comprehensive tracking status for an interview"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()

        if not candidate:
            return jsonify({"error": "Interview not found"}), 404

        # Calculate status
        status = 'not_started'
        if candidate.interview_completed_at:
            status = 'completed'
        elif candidate.interview_status == 'expired':
            status = 'expired'
        elif candidate.interview_status == 'abandoned':
            status = 'abandoned'
        elif candidate.interview_started_at:
            status = 'in_progress'
        elif candidate.interview_link_clicked:
            status = 'link_clicked'

        # Time calculations
        time_remaining = None
        if candidate.interview_expires_at and not candidate.interview_completed_at:
            time_remaining = (candidate.interview_expires_at - datetime.now()).total_seconds()
            if time_remaining < 0:
                time_remaining = 0

        tracking_data = {
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "status": status,
            "tracking": {
                "link_clicked": bool(candidate.interview_link_clicked),
                "link_clicked_at": candidate.interview_link_clicked_at.isoformat() if candidate.interview_link_clicked_at else None,
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "last_activity": candidate.interview_last_activity.isoformat() if candidate.interview_last_activity else None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "expires_at": candidate.interview_expires_at.isoformat() if candidate.interview_expires_at else None,
                "time_remaining_seconds": time_remaining
            },
            "progress": {
                "percentage": candidate.interview_progress_percentage or 0,
                "total_questions": candidate.interview_total_questions or 0,
                "answered_questions": candidate.interview_answered_questions or 0,
                "qa_completion_rate": candidate.interview_qa_completion_rate or 0
            },
            "analysis": {
                "status": candidate.interview_ai_analysis_status,
                "score": candidate.interview_ai_score,
                "recommendation": candidate.interview_final_status
            },
            "technical": {
                "session_id": candidate.interview_session_id,
                "browser_info": candidate.interview_browser_info,
                "duration_seconds": candidate.interview_duration
            }
        }

        return jsonify(tracking_data), 200

    except Exception as e:
        logger.error(f"Error getting tracking status: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/validate-completion/<token>', methods=['POST'])
def validate_interview_completion(token):
    """Validate and ensure interview is properly marked as complete"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(interview_token=token).first()
        
        if not candidate:
            return jsonify({"error": "Interview not found"}), 404
        
        # Force completion check
        qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        
        issues = []
        
        # Check for completion criteria
        if len(questions) < 5:
            issues.append("Less than 5 questions asked")
        
        if len(answers) < len(questions) * 0.8:  # 80% answer rate
            issues.append(f"Low answer rate: {len(answers)}/{len(questions)}")
        
        # Initialize inactive_time with a default value
        inactive_time = 0
        
        # Check for recent activity
        if hasattr(candidate, 'interview_last_activity') and candidate.interview_last_activity:
            inactive_time = (datetime.now() - candidate.interview_last_activity).total_seconds()
            if inactive_time > 3600:  # 1 hour
                issues.append("No activity for over 1 hour")
        
        # Force completion if criteria met
        should_complete = False
        if not candidate.interview_completed_at:
            if len(questions) >= 10 and len(answers) >= len(questions) - 1:
                should_complete = True
                logger.info(f"Completion criteria met: 10+ questions with most answered")
            elif inactive_time > 3600 and len(questions) >= 5:
                should_complete = True
                logger.info(f"Completion criteria met: Inactive for 1hr with 5+ questions")
            elif len(questions) >= 5 and len(answers) >= len(questions) * 0.8:
                should_complete = True
                logger.info(f"Completion criteria met: 5+ questions with 80% answered")
        
        if should_complete:
            candidate.interview_completed_at = datetime.now()
            candidate.interview_status = 'completed'
            candidate.interview_progress_percentage = 100
            
            if candidate.interview_started_at:
                duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                candidate.interview_duration = int(duration)
            
            candidate.interview_ai_analysis_status = 'pending'
            session.commit()
            
            logger.info(f"Interview force completed for {candidate.name} (ID: {candidate.id})")
            
            # Trigger scoring
            try:
                trigger_auto_scoring(candidate.id)
            except Exception as e:
                logger.error(f"Failed to trigger auto scoring: {e}")
            
            return jsonify({
                "success": True,
                "message": "Interview marked as complete",
                "forced_completion": True,
                "issues": issues
            }), 200
        
        return jsonify({
            "success": True,
            "message": "Interview status validated",
            "is_complete": candidate.interview_completed_at is not None,
            "issues": issues,
            "stats": {
                "questions": len(questions),
                "answers": len(answers),
                "inactive_minutes": int(inactive_time / 60) if inactive_time else 0
            }
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error validating completion: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/cleanup-incomplete', methods=['POST'])
def cleanup_incomplete_interviews():
    """Fix interviews that show as completed in UI but not in database"""
    session = SessionLocal()
    fixed_count = 0
    
    try:
        # Find candidates with started interviews but no completion
        incomplete = session.query(Candidate).filter(
            Candidate.interview_started_at.isnot(None),
            Candidate.interview_completed_at.is_(None),
            Candidate.interview_token.isnot(None)
        ).all()
        
        for candidate in incomplete:
            # Check if interview is actually complete based on other indicators
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            
            # If they have Q&A data or have been inactive for over 2 hours
            if qa_pairs or (candidate.interview_started_at and 
                          (datetime.now() - candidate.interview_started_at).total_seconds() > 7200):
                
                candidate.interview_completed_at = datetime.now()
                candidate.interview_status = 'completed'
                candidate.interview_progress_percentage = 100
                
                if candidate.interview_started_at:
                    duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                    candidate.interview_duration = int(duration)
                
                candidate.interview_ai_analysis_status = 'pending'
                fixed_count += 1
                
                # Trigger scoring
                trigger_auto_scoring(candidate.id)
                
                logger.info(f"Fixed incomplete interview for {candidate.name} (ID: {candidate.id})")
        
        session.commit()
        
        return jsonify({
            "success": True,
            "fixed_count": fixed_count,
            "message": f"Fixed {fixed_count} incomplete interviews"
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Cleanup error: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

def interview_auto_recovery_system():
    """Continuously monitor and fix incomplete interviews"""
    while True:
        time.sleep(120)  # Check every 2 minutes
        
        session = SessionLocal()
        try:
            from datetime import datetime, timedelta
            
            # Find potentially stuck interviews
            now = datetime.now()
            
            # Case 1: Started but not completed after 1 hour
            one_hour_ago = now - timedelta(hours=1)
            stuck_interviews = session.query(Candidate).filter(
                Candidate.interview_started_at.isnot(None),
                Candidate.interview_completed_at.is_(None),
                Candidate.interview_started_at < one_hour_ago
            ).all()
            
            for candidate in stuck_interviews:
                # Check if has Q&A data
                has_qa_data = False
                try:
                    qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
                    questions = json.loads(candidate.interview_questions_asked or '[]')
                    has_qa_data = len(qa_pairs) > 0 or len(questions) > 0
                except:
                    pass
                
                if has_qa_data or (now - candidate.interview_started_at).total_seconds() > 7200:
                    candidate.interview_completed_at = now
                    candidate.interview_status = 'completed'
                    candidate.interview_progress_percentage = 100
                    candidate.final_status = 'Interview Completed - Auto Recovery'
                    candidate.interview_ai_analysis_status = 'pending'
                    
                    if candidate.interview_started_at:
                        duration = (now - candidate.interview_started_at).total_seconds()
                        candidate.interview_duration = int(duration)
                    
                    logger.info(f"Auto-recovered interview for {candidate.name} (ID: {candidate.id})")
            
            # Case 2: Has 100% progress but no completion timestamp
            incomplete_100 = session.query(Candidate).filter(
                Candidate.interview_progress_percentage >= 100,
                Candidate.interview_completed_at.is_(None)
            ).all()
            
            for candidate in incomplete_100:
                candidate.interview_completed_at = now
                candidate.interview_status = 'completed'
                candidate.final_status = 'Interview Completed - Progress 100%'
                candidate.interview_ai_analysis_status = 'pending'
                logger.info(f"Completed interview at 100% progress for {candidate.name}")
            
            session.commit()
            
            # Clear cache after updates
            if stuck_interviews or incomplete_100:
                cache.clear()
            
        except Exception as e:
            logger.error(f"Auto-recovery error: {e}")
            session.rollback()
        finally:
            session.close()

@app.route('/api/interview/debug-status', methods=['GET'])
def debug_interview_status():
    """Debug endpoint to check interview statuses"""
    session = SessionLocal()
    try:
        # Get all interviews with various states
        all_interviews = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        results = {
            "total_scheduled": len(all_interviews),
            "breakdown": {
                "completed_with_score": 0,
                "completed_no_score": 0,
                "started_not_completed": 0,
                "scheduled_not_started": 0,
                "has_qa_data": 0
            },
            "details": []
        }
        
        for candidate in all_interviews:
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            
            status = "unknown"
            if candidate.interview_completed_at and candidate.interview_ai_score:
                status = "completed_with_score"
                results["breakdown"]["completed_with_score"] += 1
            elif candidate.interview_completed_at and not candidate.interview_ai_score:
                status = "completed_no_score"
                results["breakdown"]["completed_no_score"] += 1
            elif candidate.interview_started_at and not candidate.interview_completed_at:
                status = "started_not_completed"
                results["breakdown"]["started_not_completed"] += 1
            elif not candidate.interview_started_at:
                status = "scheduled_not_started"
                results["breakdown"]["scheduled_not_started"] += 1
            
            if qa_pairs:
                results["breakdown"]["has_qa_data"] += 1
            
            results["details"].append({
                "id": candidate.id,
                "name": candidate.name,
                "status": status,
                "scheduled": candidate.interview_scheduled,
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "ai_score": candidate.interview_ai_score,
                "ai_analysis_status": candidate.interview_ai_analysis_status,
                "auto_score_triggered": candidate.interview_auto_score_triggered,
                "qa_pairs_count": len(qa_pairs),
                "final_status": candidate.final_status
            })
        
        return jsonify(results), 200
        
    finally:
        session.close()

@app.route('/api/interview/fix-incomplete-interviews', methods=['POST'])
def fix_incomplete_interviews():
    """Fix interviews that have Q&A data but aren't marked as complete"""
    session = SessionLocal()
    fixed = 0
    
    try:
        # Find all scheduled interviews
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        for candidate in candidates:
            # Check if they have Q&A data
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            questions_asked = json.loads(candidate.interview_questions_asked or '[]')
            answers_given = json.loads(candidate.interview_answers_given or '[]')
            
            # If they have Q&A data but no completion timestamp, mark as complete
            if (qa_pairs or questions_asked) and not candidate.interview_completed_at:
                candidate.interview_completed_at = datetime.now()
                candidate.interview_status = 'completed'
                candidate.interview_progress_percentage = 100
                
                # Calculate duration if we have start time
                if candidate.interview_started_at:
                    duration = (candidate.interview_completed_at - candidate.interview_started_at).total_seconds()
                    candidate.interview_duration = int(duration)
                
                # Set for AI analysis
                candidate.interview_ai_analysis_status = 'pending'
                candidate.interview_auto_score_triggered = False
                candidate.final_status = 'Interview Completed - Pending Analysis'
                
                session.commit()
                
                # Now trigger the scoring
                trigger_auto_scoring(candidate.id)
                
                fixed += 1
                logger.info(f"Fixed and triggered analysis for {candidate.name} (ID: {candidate.id})")
        
        # Clear cache
        cache.delete_memoized(get_cached_candidates)
        
        return jsonify({
            "success": True,
            "fixed": fixed,
            "checked": len(candidates),
            "message": f"Fixed {fixed} incomplete interviews"
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error fixing incomplete interviews: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/force-score/<int:candidate_id>', methods=['POST'])
def force_score_candidate(candidate_id):
    """Force score a specific candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Mark as completed if not already
        if not candidate.interview_completed_at:
            candidate.interview_completed_at = datetime.now()
            candidate.interview_status = 'completed'
            candidate.interview_progress_percentage = 100
        
        # Reset analysis flags
        candidate.interview_ai_analysis_status = 'pending'
        candidate.interview_auto_score_triggered = False
        
        session.commit()
        
        # Trigger scoring
        trigger_auto_scoring(candidate_id)
        
        return jsonify({
            "success": True,
            "message": f"Triggered scoring for {candidate.name}",
            "candidate_id": candidate_id
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error forcing score: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

# Add this after your existing track_qa_enhanced function

@app.route('/api/interview/qa/get/<int:candidate_id>', methods=['GET'])
def get_interview_qa_data(candidate_id):
    """Get complete Q&A data for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Parse all Q&A data
        questions = json.loads(candidate.interview_questions_asked or '[]')
        answers = json.loads(candidate.interview_answers_given or '[]')
        qa_pairs = json.loads(getattr(candidate, 'interview_qa_pairs', '[]'))
        
        return jsonify({
            "success": True,
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email,
                "position": candidate.job_title
            },
            "qa_data": {
                "questions": questions,
                "answers": answers,
                "qa_pairs": qa_pairs,
                "total_questions": len(questions),
                "total_answers": len(answers),
                "completion_rate": f"{(len(answers) / len(questions) * 100) if questions else 0:.1f}%"
            },
            "transcript": candidate.interview_transcript,
            "analysis": {
                "status": candidate.interview_ai_analysis_status,
                "overall_score": candidate.interview_ai_score,
                "technical_score": candidate.interview_ai_technical_score,
                "communication_score": candidate.interview_ai_communication_score,
                "problem_solving_score": candidate.interview_ai_problem_solving_score,
                "cultural_fit_score": candidate.interview_ai_cultural_fit_score,
                "feedback": candidate.interview_ai_overall_feedback,
                "recommendation": candidate.interview_final_status
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting Q&A data: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/debug/candidate-fields/<int:candidate_id>', methods=['GET'])
def debug_candidate_fields(candidate_id):
    """Debug endpoint to see what fields a candidate has"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        # Get all attributes
        fields = {}
        for key in dir(candidate):
            if not key.startswith('_') and not callable(getattr(candidate, key)):
                try:
                    value = getattr(candidate, key)
                    fields[key] = str(type(value).__name__)
                except:
                    fields[key] = "error reading"
        
        return jsonify({
            "candidate_id": candidate_id,
            "available_fields": fields,
            "has_knowledge_base_id": hasattr(candidate, 'knowledge_base_id'),
            "has_interview_kb_id": hasattr(candidate, 'interview_kb_id')
        }), 200
        
    finally:
        session.close()

@app.route('/api/interview/check-analysis/<int:candidate_id>', methods=['GET'])
def check_analysis_status(candidate_id):
    """Check detailed analysis status for a candidate"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        return jsonify({
            "candidate": {
                "id": candidate.id,
                "name": candidate.name,
                "email": candidate.email
            },
            "interview": {
                "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
                "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None,
                "duration": candidate.interview_duration,
                "progress": candidate.interview_progress_percentage
            },
            "analysis": {
                "status": candidate.interview_ai_analysis_status,
                "triggered": candidate.interview_auto_score_triggered,
                "scores": {
                    "overall": candidate.interview_ai_score,
                    "technical": candidate.interview_ai_technical_score,
                    "communication": candidate.interview_ai_communication_score,
                    "problem_solving": candidate.interview_ai_problem_solving_score,
                    "cultural_fit": candidate.interview_ai_cultural_fit_score
                },
                "feedback": candidate.interview_ai_overall_feedback,
                "final_status": candidate.interview_final_status,
                "strengths": json.loads(candidate.interview_ai_strengths or '[]'),
                "weaknesses": json.loads(candidate.interview_ai_weaknesses or '[]')
            },
            "qa_data": {
                "qa_pairs": len(json.loads(candidate.interview_qa_pairs or '[]')),
                "questions_asked": len(json.loads(candidate.interview_questions_asked or '[]')),
                "answers_given": len(json.loads(candidate.interview_answers_given or '[]'))
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error checking analysis: {e}")
        return jsonify({"error": str(e)}), 500
    finally:
        session.close()

@app.route('/api/interview/recording/<int:candidate_id>', methods=['GET'])
def get_interview_recording_info(candidate_id):
    """Get interview recording information"""
    session = SessionLocal()
    try:
        candidate = session.query(Candidate).filter_by(id=candidate_id).first()
        if not candidate:
            return jsonify({"error": "Candidate not found"}), 404
        
        recording_info = {
            "recording_file": candidate.interview_recording_file,
            "recording_duration": candidate.interview_recording_duration,
            "recording_size": candidate.interview_recording_size,
            "recording_format": candidate.interview_recording_format,
            "recording_quality": candidate.interview_recording_quality,
            "recording_status": candidate.interview_recording_status,
            "session_id": candidate.interview_session_id,
            "started_at": candidate.interview_started_at.isoformat() if candidate.interview_started_at else None,
            "completed_at": candidate.interview_completed_at.isoformat() if candidate.interview_completed_at else None
        }
        
        return jsonify({
            "success": True,
            "candidate_id": candidate_id,
            "recording_info": recording_info
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting recording info: {e}")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        session.close()

@app.route('/api/routes', methods=['GET'])
def list_routes():
    """Debug endpoint to list all available routes"""
    routes = []
    for rule in app.url_map.iter_rules():
        if rule.endpoint != 'static':
            routes.append({
                'endpoint': rule.endpoint,
                'methods': list(rule.methods),
                'rule': str(rule)
            })
    return jsonify({
        "total_routes": len(routes),
        "routes": sorted(routes, key=lambda x: x['rule'])
    }), 200

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({
        "error": "Endpoint not found",
        "available_endpoints": [
            "/",
            "/api/jobs",
            "/api/candidates",
            "/api/run_full_pipeline",
            "/api/pipeline_status",
            "/api/recruitment-stats",
            "/secure-interview/<token>",
            "/health",
            "/api/routes"
        ]
    }), 404

@app.route('/api/interview/speech/track', methods=['POST', 'OPTIONS'])
@cross_origin()
def track_speech_segment():
    """Track speech recognition segments with confidence scores"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        segment = data.get('segment', {})
        
        if not session_id:
            return jsonify({"error": "session_id required"}), 400
        
        session = SessionLocal()
        try:
            # Find candidate
            candidate = session.query(Candidate).filter_by(
                interview_session_id=session_id
            ).first()
            
            if not candidate:
                # Try alternate lookup
                parts = session_id.split('_')
                if len(parts) >= 2 and parts[1].isdigit():
                    candidate = session.query(Candidate).filter_by(id=int(parts[1])).first()
            
            if not candidate:
                return jsonify({"error": "Session not found"}), 404
            
            # Parse existing speech data
            speech_segments = json.loads(getattr(candidate, 'interview_speech_segments', '[]'))
            
            # Add new segment
            segment_data = {
                'id': f"seg_{len(speech_segments)}_{int(time.time())}",
                'text': segment.get('text', ''),
                'confidence': segment.get('confidence', 0),
                'is_final': segment.get('is_final', False),
                'timestamp': segment.get('timestamp', datetime.now().isoformat()),
                'duration_ms': segment.get('duration', 0)
            }
            
            speech_segments.append(segment_data)
            
            # Store segments
            if hasattr(candidate, 'interview_speech_segments'):
                candidate.interview_speech_segments = json.dumps(speech_segments)
            
            # Update transcript
            if segment.get('is_final'):
                transcript = candidate.interview_transcript or ""
                transcript += f"\n[Candidate]: {segment.get('text', '')}\n"
                candidate.interview_transcript = transcript
            
            # Update last activity
            candidate.interview_last_activity = datetime.now()
            
            session.commit()
            
            return jsonify({
                "success": True,
                "segment_id": segment_data['id'],
                "total_segments": len(speech_segments)
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Speech tracking error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview/speech/complete-utterance', methods=['POST', 'OPTIONS'])
@cross_origin()
def track_complete_utterance():
    """Track a complete utterance (full answer)"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        data = request.json
        session_id = data.get('session_id')
        utterance = data.get('utterance', {})
        
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(
                interview_session_id=session_id
            ).first()
            
            if not candidate:
                return jsonify({"error": "Session not found"}), 404
            
            # Find the current unanswered question
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            
            for qa in reversed(qa_pairs):
                if qa.get('question') and not qa.get('answer'):
                    # Found unanswered question - add answer
                    qa['answer'] = utterance.get('text', '')
                    qa['answer_confidence'] = utterance.get('confidence', 0)
                    qa['answer_timestamp'] = utterance.get('timestamp')
                    qa['answer_duration_ms'] = utterance.get('duration')
                    break
            
            # Update database
            candidate.interview_qa_pairs = json.dumps(qa_pairs)
            
            # Update answer count
            answered = sum(1 for qa in qa_pairs if qa.get('answer'))
            candidate.interview_answered_questions = answered
            
            # Update progress
            if candidate.interview_total_questions > 0:
                progress = (answered / candidate.interview_total_questions) * 100
                candidate.interview_progress_percentage = progress
            
            session.commit()
            
            # Check for auto-completion
            if answered >= 10 or (candidate.interview_total_questions > 0 and 
                                  answered >= candidate.interview_total_questions):
                # Trigger completion
                executor.submit(check_and_complete_interview, candidate.id)
            
            return jsonify({
                "success": True,
                "answered_questions": answered,
                "progress": candidate.interview_progress_percentage
            }), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Utterance tracking error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/interview/migrate-qa-data', methods=['POST'])
def migrate_qa_data():
    """Migrate Q&A data from qa_pairs to questions/answers format"""
    session = SessionLocal()
    migrated = []
    
    try:
        candidates = session.query(Candidate).filter(
            Candidate.interview_scheduled == True
        ).all()
        
        for candidate in candidates:
            changes = []
            
            # Parse existing data
            qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            questions = json.loads(candidate.interview_questions_asked or '[]')
            answers = json.loads(candidate.interview_answers_given or '[]')
            
            # If qa_pairs exist but questions/answers are empty, migrate
            if qa_pairs and not questions:
                new_questions = []
                new_answers = []
                
                for qa in qa_pairs:
                    if qa.get('question'):
                        new_questions.append({
                            'id': qa.get('id'),
                            'text': qa.get('question'),
                            'timestamp': qa.get('timestamp'),
                            'order': qa.get('order', len(new_questions) + 1)
                        })
                        
                    if qa.get('answer'):
                        new_answers.append({
                            'id': qa.get('id'),
                            'text': qa.get('answer'),
                            'timestamp': qa.get('answered_at', qa.get('timestamp')),
                            'question_order': qa.get('order', len(new_answers) + 1)
                        })
                
                # Update the database
                candidate.interview_questions_asked = json.dumps(new_questions)
                candidate.interview_answers_given = json.dumps(new_answers)
                candidate.interview_total_questions = len(new_questions)
                candidate.interview_answered_questions = len(new_answers)
                
                changes.append(f"Migrated {len(new_questions)} questions and {len(new_answers)} answers")
            
            # Fix counts based on actual data
            actual_questions = max(
                len(qa_pairs),
                len(questions),
                candidate.interview_total_questions or 0
            )
            actual_answers = len([qa for qa in qa_pairs if qa.get('answer')]) or len(answers)
            
            if candidate.interview_total_questions != actual_questions:
                candidate.interview_total_questions = actual_questions
                changes.append(f"Fixed total questions: {actual_questions}")
            
            if candidate.interview_answered_questions != actual_answers:
                candidate.interview_answered_questions = actual_answers
                changes.append(f"Fixed answered questions: {actual_answers}")
            
            # Ensure completed interviews have completion timestamp
            if actual_questions > 0 and not candidate.interview_completed_at:
                if candidate.interview_started_at:
                    # Check if it's been more than 30 minutes
                    time_since = (datetime.now() - candidate.interview_started_at).total_seconds()
                    if time_since > 1800:  # 30 minutes
                        candidate.interview_completed_at = datetime.now()
                        candidate.interview_progress_percentage = 100
                        changes.append("Marked as completed (timeout)")
            
            if changes:
                migrated.append({
                    'id': candidate.id,
                    'name': candidate.name,
                    'changes': changes
                })
        
        session.commit()
        
        # Clear cache
        cache.clear()
        
        return jsonify({
            'success': True,
            'migrated_count': len(migrated),
            'details': migrated
        }), 200
        
    except Exception as e:
        session.rollback()
        logger.error(f"Error migrating Q&A data: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        session.close()

def run_bulk_scraping_with_monitoring():
    """Wrapper to run bulk scraping with monitoring"""
    start_time = time.time()
    
    try:
        logger.info("Starting bulk scraping for all pending assessments")
        
        # Import and run the bulk scraping function
        try:
            from testlify_results_scraper import scrape_all_pending_assessments
        except ImportError as e:
            logger.error(f"Failed to import scraper: {e}")
            notify_admin(
                "Scraper Import Error",
                f"Could not import results scraper: {str(e)}. Please ensure testlify_results_scraper.py is available."
            )
            return
        
        # Run the async scraping function
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            results_summary = loop.run_until_complete(scrape_all_pending_assessments())
        finally:
            loop.close()
        
        duration = time.time() - start_time
        total_candidates = sum(results_summary.values()) if isinstance(results_summary, dict) else 0
        
        logger.info(f"Bulk scraping completed in {duration:.2f} seconds. Processed {len(results_summary)} assessments, {total_candidates} candidates.")
        
        # Send success notification
        if isinstance(results_summary, dict):
            summary_text = "\n".join([f"- {assessment}: {count} candidates" for assessment, count in results_summary.items()])
        else:
            summary_text = f"Processed {total_candidates} total candidates"
            
        notify_admin(
            "Bulk Assessment Results Scraping Completed",
            f"Assessments processed: {len(results_summary) if isinstance(results_summary, dict) else 'Unknown'}\nTotal candidates: {total_candidates}\nDuration: {duration:.2f} seconds\n\nBreakdown:\n{summary_text}"
        )
        
    except Exception as e:
        duration = time.time() - start_time
        error_msg = f"Bulk scraping failed after {duration:.2f} seconds"
        logger.error(error_msg, exc_info=True)
        
        # Send failure notification
        notify_admin(
            "Bulk Assessment Results Scraping Failed",
            error_msg,
            error_details=traceback.format_exc()
        )

@app.route('/health', methods=['GET'])
def health_check():
    """Enhanced health check endpoint with system status"""
    health_status = {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "2.1.0",
        "checks": {},
        "performance": {
            "avg_response_time": f"{request_metrics['avg_response_time']:.3f}s",
            "total_requests": request_metrics['total_requests'],
            "slow_requests": request_metrics['slow_requests']
        }
    }
    
    # Check database
    try:
        session = SessionLocal()
        session.execute("SELECT 1")
        session.close()
        health_status["checks"]["database"] = "healthy"
    except Exception as e:
        health_status["checks"]["database"] = f"unhealthy: {str(e)}"
        health_status["status"] = "degraded"
    
    # Check cache
    try:
        cache.set('health_check', 'ok', timeout=1)
        if cache.get('health_check') == 'ok':
            health_status["checks"]["cache"] = "healthy"
        else:
            health_status["checks"]["cache"] = "unhealthy"
            health_status["status"] = "degraded"
    except Exception as e:
        health_status["checks"]["cache"] = f"unhealthy: {str(e)}"
        health_status["status"] = "degraded"
    
    # Check thread pool
    try:
        if hasattr(executor, '_threads'):
            active_threads = len([t for t in executor._threads if t.is_alive()])
            health_status["checks"]["thread_pool"] = f"healthy ({active_threads} active threads)"
        else:
            health_status["checks"]["thread_pool"] = "healthy"
    except Exception as e:
        health_status["checks"]["thread_pool"] = f"degraded: {str(e)}"
    
    return jsonify(health_status), 200 if health_status["status"] == "healthy" else 503

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Endpoint not found","available_endpoints":["/","/api/jobs","/api/candidates","/api/secure_interview/<token>","/heath"]}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"Internal server error: {error}")
    return jsonify({"error": "Internal server error"}), 500

@app.errorhandler(429)
def rate_limit_exceeded(error):
    return jsonify({"error": "Rate limit exceeded. Please try again later."}), 429

# Request logging
@app.before_request
def log_request_info():
    """Log incoming requests"""
    logger.info(f"ðŸŒ {request.method} {request.path} from {request.remote_addr}")

def test_routes():
    """Test if routes are properly registered"""
    print("ðŸ“‹ Registered Routes:")
    for rule in app.url_map.iter_rules():
        print(f"  {rule.methods} {rule.rule} -> {rule.endpoint}")

# backend.py - Add error recovery mechanisms

class InterviewErrorRecovery:
    """Handle interview system errors and recovery"""
    
    @staticmethod
    def recover_incomplete_interviews():
        """Recover and complete incomplete interviews"""
        session = SessionLocal()
        try:
            # Find stuck interviews (started > 2 hours ago, not completed)
            cutoff_time = datetime.now() - timedelta(hours=2)
            stuck_interviews = session.query(Candidate).filter(
                Candidate.interview_started_at < cutoff_time,
                Candidate.interview_completed_at.is_(None)
            ).all()
            
            for candidate in stuck_interviews:
                # Check last activity
                if candidate.interview_last_activity:
                    time_since = (datetime.now() - candidate.interview_last_activity).total_seconds()
                    if time_since > 3600:  # No activity for 1 hour
                        # Force complete
                        candidate.interview_completed_at = datetime.now()
                        candidate.interview_ai_analysis_status = 'pending'
                        
                        # Trigger scoring with what we have
                        trigger_auto_scoring(candidate.id)
                        
                        logger.warning(f"Force completed stuck interview for candidate {candidate.id}")
            
            session.commit()
            
        except Exception as e:
            logger.error(f"Error recovering interviews: {e}")
        finally:
            session.close()
    
    @staticmethod
    def validate_interview_data(candidate_id):
        """Validate and fix interview data integrity"""
        session = SessionLocal()
        try:
            candidate = session.query(Candidate).filter_by(id=candidate_id).first()
            if not candidate:
                return False
            
            issues_fixed = []
            
            # Fix Q&A pairs
            try:
                qa_pairs = json.loads(candidate.interview_qa_pairs or '[]')
            except:
                qa_pairs = []
                candidate.interview_qa_pairs = '[]'
                issues_fixed.append('Reset invalid Q&A data')
            
            # Fix counts
            if candidate.interview_total_questions != len(qa_pairs):
                candidate.interview_total_questions = len(qa_pairs)
                issues_fixed.append('Fixed question count')
            
            answered = sum(1 for qa in qa_pairs if qa.get('answer'))
            if candidate.interview_answered_questions != answered:
                candidate.interview_answered_questions = answered
                issues_fixed.append('Fixed answer count')
            
            # Fix progress
            if candidate.interview_total_questions > 0:
                progress = (answered / candidate.interview_total_questions) * 100
                if abs(candidate.interview_progress_percentage - progress) > 1:
                    candidate.interview_progress_percentage = progress
                    issues_fixed.append('Fixed progress percentage')
            
            # Fix status inconsistencies
            if candidate.interview_completed_at and not candidate.interview_started_at:
                candidate.interview_started_at = candidate.interview_completed_at - timedelta(minutes=30)
                issues_fixed.append('Fixed missing start time')
            
            if issues_fixed:
                session.commit()
                logger.info(f"Fixed issues for candidate {candidate_id}: {', '.join(issues_fixed)}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating interview data: {e}")
            return False
        finally:
            session.close()

# Add scheduled task to check and recover
@app.route('/api/interview/health-check', methods=['GET'])
def interview_health_check():
    """Check interview system health and recover if needed"""
    try:
        # Run recovery
        InterviewErrorRecovery.recover_incomplete_interviews()
        
        # Get system stats
        session = SessionLocal()
        try:
            stats = {
                'active_interviews': session.query(Candidate).filter(
                    Candidate.interview_started_at.isnot(None),
                    Candidate.interview_completed_at.is_(None)
                ).count(),
                'pending_scoring': session.query(Candidate).filter(
                    Candidate.interview_completed_at.isnot(None),
                    Candidate.interview_ai_score.is_(None)
                ).count(),
                'failed_scoring': session.query(Candidate).filter(
                    Candidate.interview_ai_analysis_status == 'failed'
                ).count(),
                'system_status': 'healthy'
            }
            
            return jsonify(stats), 200
            
        finally:
            session.close()
            
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return jsonify({"error": str(e), "system_status": "unhealthy"}), 500

@app.route('/api/cache/clear', methods=['POST'])
def clear_all_cache():
    """Clear all cached data"""
    try:
        cache.clear()
        return jsonify({"success": True, "message": "Cache cleared"}), 200
    except Exception as e:
        logger.error(f"Error clearing cache: {e}")
        return jsonify({"error": str(e)}), 500

# Add periodic health check
def start_health_monitoring():
    """Start periodic health monitoring"""
    def monitor():
        while True:
            time.sleep(300)  # Check every 5 minutes
            try:
                InterviewErrorRecovery.recover_incomplete_interviews()
            except Exception as e:
                logger.error(f"Health monitor error: {e}")
    
    thread = threading.Thread(target=monitor, daemon=True)
    thread.start()

# Add this improved completion handler to backend.py

class InterviewCompletionHandler:
    """Handles all interview completion logic with guaranteed database updates"""
    
    @staticmethod
    def complete_interview(token: str, trigger_source: str = "unknown") -> dict:
        """
        Complete an interview and ensure database is updated
        Returns: dict with success status and details
        """
        from datetime import datetime, timezone
        
        session = SessionLocal()
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                # Get candidate with lock to prevent concurrent updates
                candidate = session.query(Candidate).filter_by(
                    interview_token=token
                ).with_for_update().first()
                
                if not candidate:
                    return {"success": False, "error": "Candidate not found"}
                
                # Check if already completed
                if candidate.interview_completed_at:
                    logger.info(f"Interview already completed for {candidate.name}")
                    return {
                        "success": True,
                        "already_completed": True,
                        "completed_at": candidate.interview_completed_at.isoformat()
                    }
                
                # Set completion fields
                completion_time = datetime.now(timezone.utc)
                candidate.interview_completed_at = completion_time
                candidate.interview_status = 'completed'
                candidate.interview_progress_percentage = 100.0
                candidate.final_status = f'Interview Completed - {trigger_source}'
                
                # Calculate duration if started
                if candidate.interview_started_at:
                    # Handle timezone-aware comparison
                    start_time = candidate.interview_started_at
                    if start_time.tzinfo is None:
                        start_time = start_time.replace(tzinfo=timezone.utc)
                    duration = (completion_time - start_time).total_seconds()
                    candidate.interview_duration = int(duration)
                else:
                    candidate.interview_duration = 0
                
                # Set for AI analysis
                candidate.interview_ai_analysis_status = 'pending'
                candidate.interview_auto_score_triggered = False
                
                # Update Q&A stats if needed
                if not candidate.interview_total_questions:
                    candidate.interview_total_questions = 10  # Default
                if not candidate.interview_answered_questions:
                    candidate.interview_answered_questions = candidate.interview_total_questions
                
                # Force commit with explicit flush
                session.flush()
                session.commit()
                
                # Verify the update worked
                session.refresh(candidate)
                if candidate.interview_completed_at:
                    logger.info(f"Interview completed successfully for {candidate.name} at {completion_time}")
                    
                    # Trigger AI scoring in background
                    try:
                        executor.submit(trigger_auto_scoring, candidate.id)
                    except Exception as e:
                        logger.error(f"Failed to trigger scoring: {e}")
                    
                    return {
                        "success": True,
                        "completed_at": candidate.interview_completed_at.isoformat(),
                        "duration": candidate.interview_duration,
                        "trigger_source": trigger_source
                    }
                else:
                    raise Exception("Completion timestamp not saved")
                    
            except Exception as e:
                session.rollback()
                retry_count += 1
                logger.error(f"Attempt {retry_count} failed: {e}")
                
                if retry_count >= max_retries:
                    # Last resort: try direct SQL update
                    try:
                        session.close()
                        session = SessionLocal()
                        result = session.execute(
                            text("""
                                UPDATE candidates 
                                SET interview_completed_at = :completed_at,
                                    interview_status = 'completed',
                                    interview_progress_percentage = 100,
                                    final_status = :final_status,
                                    interview_ai_analysis_status = 'pending'
                                WHERE interview_token = :token
                            """),
                            {
                                "completed_at": datetime.now(),
                                "final_status": f"Interview Completed - {trigger_source}",
                                "token": token
                            }
                        )
                        session.commit()
                        
                        if result.rowcount > 0:
                            logger.info(f"Completed via direct SQL for token {token}")
                            return {"success": True, "method": "direct_sql"}
                    except Exception as sql_error:
                        logger.error(f"Direct SQL also failed: {sql_error}")
                    
                    return {"success": False, "error": str(e)}
                
                time.sleep(0.5)  # Brief delay before retry
            finally:
                if retry_count >= max_retries:
                    session.close()
        
        session.close()
        return {"success": False, "error": "Max retries exceeded"}

# Create global instance
completion_handler = InterviewCompletionHandler()

# Add this service to automatically complete interviews based on conditions

class AutomaticCompletionMonitor:
    """Monitor and automatically complete interviews"""
    
    def __init__(self):
        self.is_running = False
        self.check_interval = 30  # seconds
        self.thread = None
    
    def start(self):
        if self.is_running:
            return
        self.is_running = True
        self.thread = threading.Thread(target=self._monitor_loop, daemon=True)
        self.thread.start()
        logger.info("Automatic completion monitor started")
    
    def stop(self):
        self.is_running = False
        if self.thread:
            self.thread.join(timeout=5)
    
    def _monitor_loop(self):
        while self.is_running:
            try:
                self._check_all_interviews()
                time.sleep(self.check_interval)
            except Exception as e:
                logger.error(f"Monitor error: {e}")
                time.sleep(10)
    
    def _check_all_interviews(self):
        """Check all active interviews for completion conditions"""
        session = SessionLocal()
        try:
            # Find all active interviews
            active_interviews = session.query(Candidate).filter(
                Candidate.interview_started_at.isnot(None),
                Candidate.interview_completed_at.is_(None)
            ).all()
            
            for candidate in active_interviews:
                if self._should_complete(candidate):
                    logger.info(f"Auto-completing interview for {candidate.name}")
                    completion_handler.complete_interview(
                        candidate.interview_token,
                        "automatic_condition_met"
                    )
            
        except Exception as e:
            logger.error(f"Error checking interviews: {e}")
        finally:
            session.close()
    
    def _should_complete(self, candidate) -> bool:
        """Check if interview should be completed"""
        now = datetime.now()
        
        # Condition 1: Progress is 100%
        if candidate.interview_progress_percentage >= 100:
            return True
        
        # Condition 2: All questions answered
        if (candidate.interview_total_questions > 0 and 
            candidate.interview_answered_questions >= candidate.interview_total_questions):
            return True
        
        # Condition 3: Minimum threshold met (10 questions)
        if candidate.interview_answered_questions >= 10:
            return True
        
        # Condition 4: Interview running for over 45 minutes
        if candidate.interview_started_at:
            duration = (now - candidate.interview_started_at).total_seconds()
            if duration > 2700:  # 45 minutes
                return True
        
        # Condition 5: No activity for 15 minutes
        if candidate.interview_last_activity:
            inactive_time = (now - candidate.interview_last_activity).total_seconds()
            if inactive_time > 900 and candidate.interview_answered_questions >= 5:
                return True
        
        return False

# Create global instance
completion_monitor = AutomaticCompletionMonitor()

def process_pending_analyses():
    """Process any pending interview analyses"""
    session = SessionLocal()
    try:
        # Find all completed interviews without scores
        pending = session.query(Candidate).filter(
            Candidate.interview_completed_at.isnot(None),
            Candidate.interview_ai_score.is_(None),
            Candidate.interview_ai_analysis_status != 'processing'
        ).all()
        
        logger.info(f"Found {len(pending)} pending interview analyses")
        
        for candidate in pending:
            logger.info(f"Processing pending analysis for candidate {candidate.id}")
            trigger_auto_scoring(candidate.id)
            time.sleep(1)  # Small delay between processing
            
    except Exception as e:
        logger.error(f"Error processing pending analyses: {e}")
    finally:
        session.close()

# Add scheduled task to check every 5 minutes
def start_analysis_monitor():
    """Start monitoring for pending analyses"""
    def monitor():
        while True:
            time.sleep(300)  # Check every 5 minutes
            try:
                process_pending_analyses()
            except Exception as e:
                logger.error(f"Analysis monitor error: {e}")
    
    thread = threading.Thread(target=monitor, daemon=True)
    thread.start()

# Configure Flask-Mail
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', '587'))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'True').lower() == 'true'
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER')
mail = Mail(app)
app.register_blueprint(auth_bp)

# Cleanup on shutdown
import atexit

def cleanup():
    """Cleanup resources on shutdown"""
    logger.info("Shutting down TalentFlow AI Backend...")
    executor.shutdown(wait=True)
    
atexit.register(cleanup)

if __name__ == "__main__":
    print("Starting TalentFlow AI Backend (Optimized Version)...")
    print("Server running at http://127.0.0.1:5000")
    print("Logging to: logs/talentflow.log")
    print("Performance optimizations enabled")
    print("Caching enabled")
    print("Pipeline status tracking enabled")

    print("Starting Interview Automation System...")
    start_interview_automation()
    print("Interview automation running (checking every 30 minutes)")
    print("Starting interview analysis monitor...")
    start_analysis_monitor()

    print("Running conversation storage migration...")
    migrate_conversation_storage()
    print("Migration complete")


    with app.app_context():
        print("\n Registered Routes:")
        for rule in app.url_map.iter_rules():
            print(f"  {list(rule.methods)} {rule.rule} -> {rule.endpoint}")

    try:
        try: 
            print("\n Starting Interview Automation System...")
            start_interview_automation()
            print("Interview automation running (checking every 30 minutes)")

            recovery_thread = threading.Thread(target=interview_auto_recovery_system, daemon=True)
            recovery_thread.start()
            logger.info("Interview auto-recovery system started")
            
            completion_monitor.start()
            print("Automatic completion monitor started")

        except Exception as e:
            print(f" Interview automation not available: {e}")
           
            
        app.run(
            host='0.0.0.0',
            port=5000,
            debug=os.getenv('FLASK_ENV') == 'development',
            use_reloader=False,
            threaded=True
        )
    finally:
        print("Shutting down interview automation...")
        stop_interview_automation()
        completion_monitor.stop()